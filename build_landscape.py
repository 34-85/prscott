#!/usr/bin/env python3
"""Competitive Landscape one-pager + positioning map for The Scenic City Insider.
Reader-funded, no-ads membership vs. the ad/grant-funded Chattanooga field (verified June 2026)."""
import os, csv, matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY="1B3A4B"; TEAL="2A9D8F"; ORANGE="E76F51"; SAND="E9C46A"; SLATE="415A6B"
LIGHT="EEF3F4"; GREY="6B7B85"; GOLD="B8860B"; GREEN="2E7D52"; RUST="A6452F"
ASSETS="/home/user/prscott/assets"; os.makedirs(ASSETS,exist_ok=True)
plt.rcParams.update({"font.family":"DejaVu Sans","figure.dpi":160})

# (name, type, model, focus, overlap_with_scout, conf)
ROWS=[
 ("NOOGAtoday (6AM City)","Free daily email + web","Ads / sponsorship; free to reader","'Need-to-know' news, things-to-do, events; broad audience",
  "HIGH — same daily-newsletter habit & events hook. The scaled, professional incumbent.","High"),
 ("Chattanooga Times Free Press","Daily news org; print + digital","Subscription / paywall + ads","Region's paper of record; full local + events (ChattNow)",
  "MED — reader-paid like us, but legacy & retreating (weekday single-copy retail ended Jan 2026; Sunday print only).","High"),
 ("Chattynooga","Weekly email newsletter","Free; word-of-mouth + paid Meta ads; occasional creative sponsorship","'Porch-level' feel-good stories, neighbors, events; ~900 subs",
  "HIGH (tone) — closest in spirit (warm, curated, hyperlocal). Also stewards Roundabout NorthChatt. The peer to know.","High"),
 ("Roundabout / New_ Public","Local community app (AT Protocol)","Grant / philanthropy; free; revenue model undefined","Channels + events calendar + guides; suburban ring; closed beta",
  "LOW-MED — community feed, not editorial. Adjacent; possible channel/tip-line, not a substitute.","High"),
 ("CHA Guide","Weekly email + mobile app","Free; local-partner content","Upbeat 'positive news,' events, restaurants, experiences",
  "MED — curated discovery overlap; free/partner-funded.","Med"),
 ("HERE Chattanooga","Local news site + newsletter (HERE network)","Advertising + paid services / press releases","Local news, business, lifestyle; Hamilton County",
  "MED — ad-funded local news; events & business overlap.","Med"),
 ("The Pulse","Alt-weekly, now online-only (Brewer Media)","Ads; free","Arts, music, culture, dining; young professionals",
  "MED — culture/events calendar overlap; print suspended 2020.","High"),
 ("The Chattanoogan","Free web-only daily","Free; ad-supported (inferred)","Breaking news, obituaries, gov't, business; since 1999",
  "LOW — hard news/record, not curated lifestyle.","Med"),
 ("Chattanooga City Lifestyle","Glossy lifestyle mag/site (franchise)","Advertising (franchise)","Affluent lifestyle features",
  "LOW-MED — advertiser-driven by design; opposite of our bright line.","Med"),
 ("Visit Chattanooga / Nextdoor / FB groups","Tourism calendar + social platforms","Tourism-funded / free ad platforms","Events listings; neighbor chatter & recommendations",
  "MED (events) / LOW (editorial) — discovery & community substitutes, not finished editorial.","Med"),
]

# --- 2x2 positioning map ---
# x: 0=advertiser/grant-funded(free)  ->  1=reader-funded(membership/paid)
# y: 0=raw feed/community/listings    ->  1=curated editorial / finished product
PLACED=[
 # name, x, y, color, note
 ("The Scenic City Insider",0.86,0.90,ORANGE,"us"),
 ("Times Free Press",0.74,0.72,SLATE,"legacy/retreating"),
 ("NOOGAtoday",0.30,0.70,NAVY,""),
 ("Chattynooga",0.22,0.78,TEAL,"peer"),
 ("CHA Guide",0.26,0.55,GREEN,""),
 ("HERE Chattanooga",0.30,0.50,GREY,""),
 ("The Pulse",0.24,0.62,GOLD,""),
 ("The Chattanoogan",0.34,0.40,GREY,""),
 ("City Lifestyle",0.14,0.52,GOLD,""),
 ("Visit CHA calendar",0.12,0.28,SLATE,""),
 ("Roundabout (NorthChatt)",0.10,0.16,TEAL,"grant / beta"),
 ("Nextdoor / FB groups",0.16,0.10,GREY,""),
]
def positioning_map():
    fig,ax=plt.subplots(figsize=(9.6,6.0)); ax.set_xlim(0,1); ax.set_ylim(0,1)
    ax.axhline(0.5,color="#cfd8dc",lw=1.2,zorder=1); ax.axvline(0.5,color="#cfd8dc",lw=1.2,zorder=1)
    # white-space highlight (top-right)
    ax.add_patch(plt.Rectangle((0.5,0.5),0.5,0.5,facecolor="#"+ORANGE,alpha=0.06,zorder=0))
    ax.text(0.985,0.985,"WHITE SPACE\nreader-funded + curated",ha="right",va="top",fontsize=8,
            color="#"+ORANGE,fontweight="bold",style="italic")
    for name,x,y,c,note in PLACED:
        big = name=="The Scenic City Insider"
        ax.scatter([x],[y],s=240 if big else 90,color="#"+c,zorder=3,
                   edgecolor="white",linewidth=1.6 if big else 1.0,alpha=0.95)
        lbl=name+(f"\n({note})" if note and note not in("us",) else "")
        ax.annotate(lbl,(x,y),xytext=(0,11 if big else 8),textcoords="offset points",
                    ha="center",fontsize=7.6 if big else 6.6,
                    fontweight="bold" if big else "normal",color="#"+NAVY)
    ax.set_xticks([0.07,0.93]); ax.set_xticklabels(["Advertiser- / grant-funded\n(free to reader)","Reader-funded\n(membership / paid)"],fontsize=8,color="#"+SLATE)
    ax.set_yticks([0.07,0.93]); ax.set_yticklabels(["Raw feed /\ncommunity / listings","Curated editorial /\nfinished product"],fontsize=8,color="#"+SLATE,rotation=90,va="center")
    for s in ("top","right"): ax.spines[s].set_visible(False)
    for s in ("left","bottom"): ax.spines[s].set_color("#b0bec5")
    ax.set_title("Where the Insider sits — the Chattanooga local-info field (verified June 2026)",
                 fontweight="bold",color="#"+NAVY,fontsize=11)
    fig.tight_layout(); fig.savefig(f"{ASSETS}/fig_landscape.png",bbox_inches="tight"); plt.close(fig)
positioning_map(); print("map done")

# --- CSV ---
with open("/home/user/prscott/competitive_landscape.csv","w",newline="") as f:
    w=csv.writer(f); w.writerow(["Outlet","Type","Revenue model","Focus","Overlap with the Insider","Confidence"])
    for r in ROWS: w.writerow(r)
print("CSV rows:",len(ROWS))

# ============= DOCX =============
doc=Document(); sec=doc.sections[0]; sec.orientation=WD_ORIENT.LANDSCAPE
sec.page_width,sec.page_height=Inches(11),Inches(8.5)
for m in ("top_margin","bottom_margin","left_margin","right_margin"): setattr(sec,m,Inches(0.55))
nm=doc.styles["Normal"]; nm.font.name="Calibri"; nm.font.size=Pt(10); nm.paragraph_format.space_after=Pt(4)
def sca(r,h): r.font.color.rgb=RGBColor.from_string(h)
def shade(c,h):
    p=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd")
    s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); p.append(s)
def ctext(c,t,bold=False,color=None,size=8.2,white=False,align="left"):
    c.text=""; p=c.paragraphs[0]
    p.alignment={"left":WD_ALIGN_PARAGRAPH.LEFT,"center":WD_ALIGN_PARAGRAPH.CENTER}[align]
    r=p.add_run(t); r.font.size=Pt(size); r.font.bold=bold
    if white: r.font.color.rgb=RGBColor.from_string("FFFFFF")
    elif color: r.font.color.rgb=RGBColor.from_string(color)
def H(t,level=1,color=None):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(12 if level==1 else 8); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(t); r.font.bold=True
    if level==1:
        r.font.size=Pt(15); sca(r,color or NAVY)
        pPr=p._p.get_or_add_pPr(); b=OxmlElement("w:pBdr"); bt=OxmlElement("w:bottom")
        bt.set(qn("w:val"),"single"); bt.set(qn("w:sz"),"6"); bt.set(qn("w:space"),"2"); bt.set(qn("w:color"),TEAL)
        b.append(bt); pPr.append(b)
    else: r.font.size=Pt(12); sca(r,color or TEAL)
def P(t,italic=False,color=None,size=10.5):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(5)
    for i,ch in enumerate(t.split("**")):
        if not ch: continue
        r=p.add_run(ch); r.font.size=Pt(size); r.font.italic=italic
        if i%2==1: r.font.bold=True
        if color: sca(r,color)
def BUL(items,size=9.5):
    for it in items:
        p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(3)
        for i,ch in enumerate(it.split("**")):
            if not ch: continue
            r=p.add_run(ch); r.font.size=Pt(size)
            if i%2==1: r.font.bold=True
def FIG(path,cap,width=8.6):
    doc.add_picture(path,width=Inches(width)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    P(cap,italic=True,color=GREY,size=8.5); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
def BOX(title,body):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    t.columns[0].width=Inches(9.8); cell=t.rows[0].cells[0]; shade(cell,LIGHT)
    cell.text=""; p=cell.paragraphs[0]; r=p.add_run(title); r.font.bold=True; r.font.size=Pt(10.5); sca(r,ORANGE)
    for para in body:
        q=cell.add_paragraph(); q.paragraph_format.space_before=Pt(4)
        for i,ch in enumerate(para.split("**")):
            if not ch: continue
            rr=q.add_run(ch); rr.font.size=Pt(11); rr.font.italic=True; sca(rr,NAVY)
            if i%2==1: rr.font.bold=True
    # navy left border
    tcPr=cell._tc.get_or_add_tcPr(); borders=OxmlElement("w:tcBorders")
    lb=OxmlElement("w:left"); lb.set(qn("w:val"),"single"); lb.set(qn("w:sz"),"24"); lb.set(qn("w:space"),"4"); lb.set(qn("w:color"),ORANGE)
    borders.append(lb); tcPr.append(borders)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
def LANDTABLE(rows):
    cols=["Outlet","Type","Revenue model","Focus","Overlap with the Insider (our read)","Conf."]
    widths=[1.7,1.5,1.7,2.0,2.55,0.45]
    t=doc.add_table(rows=1,cols=len(cols)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style="Table Grid"
    for i,h in enumerate(cols):
        shade(t.rows[0].cells[i],NAVY); ctext(t.rows[0].cells[i],h,bold=True,white=True,size=8.3,align="left" if i<5 else "center")
    for r,row in enumerate(rows):
        cells=t.add_row().cells; base=LIGHT if r%2==0 else "FFFFFF"
        ctext(cells[0],row[0],bold=True,color=NAVY,size=8.1); ctext(cells[1],row[1],size=7.9)
        ctext(cells[2],row[2],size=7.9); ctext(cells[3],row[3],size=7.9)
        ctext(cells[4],row[4],size=7.9); ctext(cells[5],row[5],size=7.9,align="center")
        for c in cells: shade(c,base)
    for i,w in enumerate(widths):
        for row in t.rows: row.cells[i].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)

# Title
doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("THE SCENIC CITY SCOUT"); r.font.size=Pt(22); r.font.bold=True; sca(r,NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Competitive Landscape — and why reader-funded is open ground"); r.font.size=Pt(16); sca(r,TEAL)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("The Chattanooga local-info field, verified June 2026 — by format, revenue model, and overlap with us")
r.font.size=Pt(10.5); r.font.italic=True; sca(r,SLATE)

H("The one-sentence read",1)
P("Almost every local-info product in Chattanooga is **free to the reader and funded by advertisers, sponsors, or grants** — the lone reader-paid product is the legacy daily, which is actively retreating from print. **A reader-funded, no-ads, curated membership product is open ground.** Our risk isn't a price war (you can't out-free a sponsor- or foundation-subsidized rival); it's the *events-calendar and community-bulletin* functions, which free incumbents do well. We win on curation, voice, reporting, and inbox ownership — not on being cheaper.")
FIG(f"{ASSETS}/fig_landscape.png","Funding model (left↔right) by product type (feed↔curated). The top-right quadrant — reader-funded + curated — is essentially unoccupied.")

doc.add_page_break()
H("The field, outlet by outlet",1)
LANDTABLE(ROWS)
P("Sources & confidence per row in competitive_landscape.csv. 'Overlap' is our editorial judgment, not the outlet's claim.",italic=True,color=GREY,size=8.5)

H("What this means for us",1)
BUL([
 "**Don't compete on price — compete on the finished product.** Sponsor- and grant-funded rivals are free forever; we can't and shouldn't undercut. We sell curation, voice, trust, and a no-ads promise the ad-funded field structurally can't make.",
 "**Defend the events overlap with a point of view.** A free calendar is a commodity; 'the five things actually worth your Saturday, chosen by a neighbor' is not. Curate, don't list.",
 "**Treat Chattynooga as the peer to know, not the enemy.** Same warm, hyperlocal lane; it also stewards Roundabout NorthChatt. Coopetition, cross-promo, and shared respect beat a turf fight in a small town.",
 "**Use Roundabout as a channel, not a worry.** A trust-first community feed is a clean Insider tip-line and a place to seed content that drives signups — adjacent infrastructure, not a substitute for editorial.",
 "**Watch two triggers:** (1) anyone standing up a strong *metro* (not suburban) Roundabout, and (2) any free outlet moving into original *reporting* with a paid tier. Either narrows the white space.",
 "**Own the bright line publicly.** In a field where the reader is usually the product sold to advertisers, 'the reader is the customer — never the product' is both true and a differentiator. Say it often.",
])

doc.add_page_break()
H("Positioning paragraph for the beehiiv About page",1)
P("Drop-in copy that distinguishes a curated, reader-funded newsletter from a community feed and from ad-funded roundups — without naming competitors or going negative. Matches the existing About voice (a real local neighbor, not private equity).",italic=True,color=GREY,size=9)
BOX("Primary — for the About page (≈90 words):",[
 "A quick word on what the Insider is, and isn't. It isn't a feed to scroll, a community bulletin board, or a roundup paid for by the places it covers. There are good versions of all of those in town, and I read them. The Insider is something narrower: one carefully made issue a week, written by a neighbor, paid for by you. No ads, ever, and no sponsored sentences — because the day a business can buy its way in is the day you stop being able to trust what's here. My only job is to be worth your inbox.",
])
BOX("Short — for a footer, signup form, or 'why pay?' line (≈30 words):",[
 "No ads. No sponsorships. No sponsored sentences. The Insider is one curated issue a week, written by a neighbor and paid for by readers — so the reader is the customer, never the product.",
])
BOX("One-liner — tagline / social bio:",[
 "A curated weekly for Chattanooga — reader-funded, ad-free, written by a neighbor. The reader is the customer, never the product.",
])
P("Plain-text and HTML versions of these are saved alongside this doc (insider_positioning_about.txt / .html) for pasting into beehiiv.",italic=True,color=GREY,size=8.5)

out="/home/user/prscott/Competitive_Landscape.docx"
doc.save(out); print("saved",out)

# --- positioning copy as txt + html for beehiiv ---
PRIMARY=("A quick word on what the Insider is, and isn't. It isn't a feed to scroll, a community bulletin "
"board, or a roundup paid for by the places it covers. There are good versions of all of those in town, "
"and I read them. The Insider is something narrower: one carefully made issue a week, written by a neighbor, "
"paid for by you. No ads, ever, and no sponsored sentences — because the day a business can buy its way "
"in is the day you stop being able to trust what's here. My only job is to be worth your inbox.")
SHORT=("No ads. No sponsorships. No sponsored sentences. The Insider is one curated issue a week, written by a "
"neighbor and paid for by readers — so the reader is the customer, never the product.")
ONELINE=("A curated weekly for Chattanooga — reader-funded, ad-free, written by a neighbor. "
"The reader is the customer, never the product.")
with open("/home/user/prscott/insider_positioning_about.txt","w") as f:
    f.write("PRIMARY (About page, ~90 words)\n\n"+PRIMARY+"\n\n---\n\nSHORT (footer / why-pay, ~30 words)\n\n"+SHORT+"\n\n---\n\nONE-LINER (tagline / bio)\n\n"+ONELINE+"\n")
with open("/home/user/prscott/insider_positioning_about.html","w") as f:
    f.write('<!DOCTYPE html><html><head><meta charset="utf-8"></head>'
      '<body style="font-family:Georgia,serif;color:#1B3A4B;max-width:620px;margin:0 auto;padding:18px;">'
      '<h3 style="color:#E76F51;font-family:Arial,sans-serif;">What the Insider is &mdash; and isn’t</h3>'
      f'<p style="font-size:17px;line-height:1.6;">{PRIMARY}</p>'
      f'<hr style="border:none;border-top:1px solid #ddd;margin:18px 0;">'
      f'<p style="font-size:14px;color:#415A6B;"><strong>Short:</strong> {SHORT}</p>'
      f'<p style="font-size:14px;color:#415A6B;"><strong>Tagline:</strong> {ONELINE}</p>'
      '</body></html>')
print("positioning copy written")
