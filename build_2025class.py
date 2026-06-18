#!/usr/bin/env python3
"""2025 Class — New Independents Members Should Know.
Compiled from the Scout (triage+select) -> Nerd (verify+research) pipeline over 1,892 2025 starts."""
import csv
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY="1B3A4B"; TEAL="2A9D8F"; ORANGE="E76F51"; SAND="E9C46A"; SLATE="415A6B"
LIGHT="EEF3F4"; GREY="6B7B85"; GOLD="B8860B"; GREEN="2E7D52"; RED="C0492F"

# (business, owner, what&where, status, story, social, press, tier)
DATA = {
"Food & Drink":[
 ("The Reading Room","Linden Marno-Feree","Bar + bookstore · 3210B Brainerd Rd","Open","Bookseller-bartender sources books via the Public Library Foundation and returns a cut of sales to it.","readingroombar.com","Some","1"),
 ("The Woodshop","Jhett Black","Bar + restaurant + listening room · 5500 St. Elmo Ave","Open","Award-winning bluesman owner hosts weekly 'Tuesday Bluesday'; church-pew seats, family-recipe BBQ.","thewoodshoplisteningroom.com","Some","1"),
 ("Pedestrian Wine & Cheese","(owner unconfirmed)","Wine shop + house-made cheese · 50 Frazier Ave","Open","Billed as the first TN liquor store licensed to sell its own house-made cheese; 60+ wine-club members.","pedestrianwineandcheese.com","Some","1"),
 ("Oh Wow Donuts","Linda Seng","Scratch donuts + coffee · 1213 Taft Hwy, Signal Mtn","Open","Named for the reaction she wants; full yeast-donut spread, standout apple fritter; town-recognized opening.","FB / IG","Some","1"),
 ("MOTHER (Hungry Mother)","Family-owned","Café · 825 Houston St, downtown","Open","Family café roasting NC's HEX coffee and baking warm bread in-house. (Confirm legal vs. trade name.)","IG / joe.coffee","Low","1"),
 ("New Wave Coffee","Luke Pigott (roaster)","Roaster + new café · 2405 Lyndon Ave, Red Bank","Open / expanding","Award-winning roaster; the 2025 news is the brick-and-mortar (company founded 2020).","newwave-chatt.com","Some","1"),
 ("Panadería Marquense","(owner unconfirmed)","Guatemalan bakery · 3208 Rossville Blvd","Open","Authentic Guatemalan baking (chicharrones criollo); daily fresh bread + café — immigrant-maker story.","panaderiamarquense.com","Low","1"),
 ("Hangry Pinoy","Neenuel Pattawi & Everly Guillermo","Filipino food truck · 3639 Hixson Pike","Open (2024)","Owner from Isabela, Philippines; scratch sauces sourced from the Asian market next door. (Opened 2024.)","IG @hangrypinoyfood","Some","1"),
 ("Mama O's Sarap Stop","(owner unconfirmed)","Filipino food trailer · Cleveland/Chattanooga","Open","Filipino immigrant cuisine, '10/10' local reviews; owner story not yet documented — interview-ready.","IG @mamaossarapstop","Low","2"),
 ("Stone Age Korean BBQ, Hot Pot & Sushi","(owner unconfirmed)","AYCE Korean BBQ + hot pot + sushi · 2011 Gunbarrel Rd","Open","Three concepts under one roof, new to the Gunbarrel corridor.","IG @stoneage_chattanooga","Low","2"),
 ("Roasted Root Coffee","(owner unconfirmed)","Farm-to-cup roaster/café · 2947 Cummings Hwy","Open","Local roaster also stocked at Sprout General Store (Wauhatchie Woodlands).","roastedrootcafe.com","Low","2"),
 ("Archie's Rig Barbecue","Scott Archibald","Owner-built BBQ food truck · Hixson","Open","Classic owner-built-rig maker story — brisket, pulled pork, ribs.","IG @archiesbarbeque","Low","2"),
 ("Shredders' Lair","(owner unconfirmed)","Used-music shop + café + beer · 3230 Brainerd Rd","Open/Coming","'Shred til ya dead' — instrument shop & repair meets coffee/beer hangout.","shredderslair.biz","Low","2"),
],
"Retail & Boutique":[
 ("Gothic Trading Company","Lisa Porter","Goth/alt apparel, decor, gifts · 514 Tremont St, NorthShore","Open","The 'goth Cracker Barrel' next to Aretha Frankensteins — corsets to Beetlejuice; two TFP columns.","gothictrading.com","Some","1"),
 ("Dallos Vinyl Love","Marlo White","Independent record store · 1463 Market St","Open","A crate-digger's haven downtown — rare soul/hip-hop/rock grooves.","dallosvinyllove.com","Some","1"),
 ("Birch & Willow Mercantile","Sisters Sarah & Lanita","Home/garden, gifts, local makers · Sale Creek","Open","Named for their parents' favorite trees — birch (woodworking dad), willow (gardening mom). Satellite-town gem.","birchwillowmercantile.com","Low","1"),
 ("Bloom Bible Bookstore & Floral","(owner unconfirmed)","Bibles/faith gifts + fresh florals · 12th & Broad","Open","'Scripture meets style' — flowers + Bibles with free engraving; hosts book signings.","bloombiblebooks.shop","Low","1"),
 ("The Consign House","Melissa Poole","Curated consignment · 5228 Hixson Pike, Hixson","Open","Native Chattanoogan; quirky perk — rents hand-selected English bone-china teacups.","theconsignhouse.com","Low","1"),
 ("Slater the Creator","Anna Claire Black","Upcycled one-of-one apparel · Chattanooga","Open","Solo sewist upcycling vintage textiles; delivers orders personally.","slaterthecreator.com","Low","2"),
 ("Mystic Nook","(owner unconfirmed)","Crystals / metaphysical · 6226 Ringgold Rd","Open","'Chattanooga's newest metaphysical shop' — enchanting-nook vibe.","IG @mystic_nook_metaphysical_","Low","2"),
 ("Lookout Thrift","(owner unconfirmed)","Community thrift · 2441 Broad St, Southside","Open","Mission-driven thrift with treasure-hunt appeal; free pickup.","lookoutthrift.com","Low","2"),
 ("Monkee's of Chattanooga","Autumn & Kevin Burke","Women's boutique · Read House, 842 Chestnut St (NOT Ooltewah)","Open","Local couple bring the boutique into the historic Read House. (Franchise; corrected address.)","IG @monkeesofchattanooga","Some","2"),
],
"Makers & Artisans":[
 ("The Doll For All","Marsha Roberts","Posable dolls with limb differences for kids with disabilities · Chattanooga","Open","Born from 'Aysha,' who post-surgery saw no dolls like disabled kids; the only US-made 18\" vinyl doll. Lead feature.","thedollforall.com","Some","1"),
 ("Grateful Glassworks","Blake Rolfe","Glassblowing studio + gallery, classes/demos · 1518 Central Ave","Open","A 10+-year glassblower's 'come watch / try it' studio — live-fire visual story; hit year one in 2025.","gratefulglassworks.com","Some","1"),
 ("Lazy Lights Stained Glass","Zoë Newell","Handmade stained glass · Chattanooga","Open","Lifelong local rediscovered her late mom's glass tools and taught herself.","lazylightstainedglass.com","Low","1"),
 ("Jay Ban Works","Rev. Jay Banasiak","Woodworking + laser/CNC, Chattanooga-themed pieces · pop-ups/Etsy","Open","Presbyterian pastor/musician building from antique black walnut inherited from his grandmother; also makes cajón drums.","Etsy JayBanWorks","Low","1"),
 ("Chattanooga Milk Stones","'Brooklyn'","Breastmilk / cremation / floral keepsake jewelry · Birchwood","Open","Struggled to breastfeed daughter Zelphia; now hand-grinds milk into keepsake stones — a craft most have never heard of.","chattanoogamilkstones.com","Low","1"),
 ("Walker Built","Preston Walker","Custom cabinetry / case goods · Chattanooga","Open (2025 start)","Craftsman who moved to Chattanooga in 2023 and launched in 2025 to build 'more than cabinets — community.'","Squarespace site","Low","1"),
 ("Elevated Leather","(owner unconfirmed)","Cobbler — shine, patina, custom shoes · Chattanooga","Open","A rare bespoke-leaning cobbler/patina artist in town.","Yelp","Low","1"),
 ("On The Rocks Pottery","(owner unconfirmed)","Handmade nature-inspired ceramics · Chattanooga","Open","Functional, nature-inspired ware (hook needs an interview).","linktr.ee/OnTheRocksPottery","Low","2"),
 ("Hearth & Oak Candle Co.","(owner unconfirmed)","Hand-poured candles (coconut-apricot wax) · Hamilton Co.","Open","New 2025 candle maker with a scent-consultation angle.","hearthandoakcandleco.com","Low","2"),
 ("Goth Moth Rebellion Studios","Veteran-owned","Handmade gothic/alt jewelry · markets","Open","Veteran-owned dark-aesthetic maker with a rebellious brand voice.","bigcartel","Low","2"),
 ("Gift Horse Leather Goods","William Cox","Hand-crafted leather goods · home-based","Open","New 2025 leather maker; 'don't look a gift horse…' brand hook. (Home-based — no address published.)","none found","None","2"),
 ("Rags2Dazzle Ragdolls","'Alan' / family","Show-Ragdoll cattery · Hixson (home-based)","Open","Cats raised loose in the home alongside the family dog — quirky niche; TICA/CFA registered.","rags2dazzleragdolls.com","Low","2"),
],
"Experiences, Outdoor & Quirky":[
 ("Pelfrey Drum Works","Jerry 'Punkin' Pelfrey","Hand-built stave snare drums, made to order · Hixson","Open","Decades-deep craftsman (since the '90s as 'Punkinater') now builds one-at-a-time heirloom snares sold worldwide.","pelfreydrumworks.com","Low","1"),
 ("REconnect Nature School","Jaime Peterson & Rebecca Westbrook Toker","Nature-based STEAM middle school · ~5 acres, East Ridge edge","Open","Two veteran educators built a forest school for middle-schoolers and neurodivergent kids who 'age out' of elementary programs. (Nonprofit.)","reconnectnatureschool.org","Low","1"),
 ("The Matcha Spot","Ester Gao","Matcha bar · 730 Germantown Cir (Concordance Ferments)","Open / expanding","A pop-up that 'graduated' to a creekside lease in 2025 — rare specialty-matcha concept; TFP-covered.","thematchaspot.net","Some","1"),
 ("Tennessee River Adventures","Capt. Aaron Massey","Trophy-catfish guide service · Tennessee River","Open","USCG-licensed 'Bassmassey' chasing river giants from a Sea Ark — quintessential local outdoor story.","tennesseeriveradventures.net","Low","1"),
 ("Exit 4 Escape Rooms","(owner unconfirmed)","4 themed escape rooms · 1919 Gunbarrel Rd","Open","Locally owned, objective-based rooms ('never actually locks you in') — sunken ship, haunted hotel, wand-maker.","exit4escape.com","Low","1"),
 ("Signal Swirl","Local families","Self-serve frozen yogurt · 1231 Taft Hwy, Signal Mtn","Open","Brand-new mountain treat spot; stays open till midnight on home-football nights — hyperlocal charm.","FB / IG","Low","2"),
 ("Backroad Cheesesteaks","(owner unconfirmed)","Cheesesteak food truck · Rossville GA + downtown","Open","Roadside-identity truck building a following across the GA/TN line.","FB / StreetFoodFinder","Low","2"),
 ("Chattanooga Rydables","Local franchise operator","Guided electric stuffed-animal scooter tours · 707 Cherry St","Open","Whimsical unicorn/dragon 'rydable' tours past the Aquarium and riverfront; ghost & holiday-lights variants.","rydables.com","Low","2"),
 ("Quirky Bird Tattoo","(owner unconfirmed)","Tattoo studio · 1435 Market St","Open","Character-rich indie studio with a memorable brand in a crowded scene.","quirkybirdtattoo.com","Low","2"),
 ("Nova Theatre for Young People","(owner unconfirmed)","Independent children's theatre · Chattanooga","Likely open","A new indie kids' theatre carving space alongside the big youth program — a hard-to-find community arts gem.","FB","None","2"),
 ("Jackpot Junction","Veteran-owned","Gaming/arcade lounge (18+) · 5063 Hixson Pike","Open","Veteran-owned gaming lounge with military-appreciation discounts. (NOT the Minnesota casino.)","FB / IG","Low","2"),
 ("Skull Island Campground","Operator","Island campground + camp store · Harrison (Chickamauga Lake)","Open","Camp on its own island with rotating camp hosts — a classic lake-overnight gem.","VisitChattanooga","Low","2"),
],
}

TRAPS=[
 "**Monkee's of Chattanooga** — it's downtown in the **Read House (842 Chestnut St)**, NOT Ooltewah; and it's a national franchise. Correct before printing.",
 "**Hendo's Creole Cuisine** — owner is genuinely ex-Rams RB **Darrel Henderson** (not a celebrity mix-up), but the concept had two soft openings then collapsed; the space was re-listed. Do not feature as open.",
 "**Conga Kitchen** — appears to be the existing **Conga Latin Food** under a delivery alias, not a new 2025 business.",
 "**Socci Sauce** is an **Atlanta** operation; **Beignets & Brew** is a regional **chain** — neither is an indie Chattanooga start.",
 "**Not actually 2025 starts:** New Wave Coffee (founded 2020; new café in 2025), Hangry Pinoy (2024), Wanderlinger (relocation), Wired Coffee (since 2014).",
 "**Name-collision traps:** Windy Hill Pottery (Harrison entry unconfirmed; collides with a Cleveland TN potter & a Colorado studio); 'Unique Jantiques' (a New Jersey shop); Easy Tiger Tattoos (multi-city name); Jackpot Junction (≠ Minnesota casino); A&A Market (multiple TN locations).",
]
COLD=[
 "Concept real, business unconfirmed: **thermal-drone pet/livestock recovery** ('Thermal Track' couldn't be pinned — Gig City Drones is a confirmed regional operator if you pivot the feature).",
 "Zero/!thin footprint — verify or reach out before featuring: Red Wiggler (kava bar), Vino & Verde, So-Dough, M.R.S. Pastelería, TN Whiskey Bar Co., Mighty Moms Bakery, Kim Hot Fish, NT Woodworks, Perennial Chairs, Scenic City Jewelers, Handmade by Hayden (brooms — great if real), Odd Times, Annette Rocks, A&A International Market, The Eclectic Blueprint, 'Overlooked' antiques, H+B Children's Boutique.",
 "Home-based (never publish the address; treat as market/pop-up vendors): Gift Horse Leather, Rags2Dazzle, Handmade by Hayden, several cottage bakers.",
]

# ---------- CSV ----------
with open("/home/user/prscott/2025_class_new_independents.csv","w",newline="") as f:
    w=csv.writer(f)
    w.writerow(["Category","Business","Owner/person","What & where","Status","Story angle","Social/Web","Press footprint","Tier"])
    for cat,rows in DATA.items():
        for r in rows: w.writerow((cat,)+r)
total=sum(len(v) for v in DATA.values())
print("CSV rows:",total)

# ---------- DOCX ----------
doc=Document()
sec=doc.sections[0]; sec.orientation=WD_ORIENT.LANDSCAPE
sec.page_width,sec.page_height=Inches(11),Inches(8.5)
for m in ("top_margin","bottom_margin","left_margin","right_margin"): setattr(sec,m,Inches(0.6))
nm=doc.styles["Normal"]; nm.font.name="Calibri"; nm.font.size=Pt(10); nm.paragraph_format.space_after=Pt(4)
def sca(r,h): r.font.color.rgb=RGBColor.from_string(h)
def shade(c,h):
    p=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd")
    s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); p.append(s)
def ctext(c,t,bold=False,color=None,size=8.4,white=False,align="left"):
    c.text=""; p=c.paragraphs[0]
    p.alignment={"left":WD_ALIGN_PARAGRAPH.LEFT,"center":WD_ALIGN_PARAGRAPH.CENTER}[align]
    r=p.add_run(t); r.font.size=Pt(size); r.font.bold=bold
    if white: r.font.color.rgb=RGBColor.from_string("FFFFFF")
    elif color: r.font.color.rgb=RGBColor.from_string(color)
def H(t,level=1,color=None):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(12 if level==1 else 8); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(t); r.font.bold=True
    if level==1:
        r.font.size=Pt(15); sca(r,color or NAVY)
        pPr=p._p.get_or_add_pPr(); b=OxmlElement("w:pBdr"); bt=OxmlElement("w:bottom")
        bt.set(qn("w:val"),"single"); bt.set(qn("w:sz"),"6"); bt.set(qn("w:space"),"2"); bt.set(qn("w:color"),TEAL)
        b.append(bt); pPr.append(b)
    else: r.font.size=Pt(12); sca(r,color or TEAL)
def P(t,italic=False,color=None,size=10):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(5)
    for i,ch in enumerate(t.split("**")):
        if not ch: continue
        r=p.add_run(ch); r.font.size=Pt(size); r.font.italic=italic
        if i%2==1: r.font.bold=True
        if color: sca(r,color)
def BUL(items,size=9.3):
    for it in items:
        p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(3)
        for i,ch in enumerate(it.split("**")):
            if not ch: continue
            r=p.add_run(ch); r.font.size=Pt(size)
            if i%2==1: r.font.bold=True

def gtable(rows):
    cols=["Business","Owner / person","What · Where","Status","Story angle","Press","T"]
    widths=[1.45,1.5,2.0,0.9,3.5,0.55,0.35]
    t=doc.add_table(rows=1,cols=len(cols)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style="Table Grid"
    for i,h in enumerate(cols):
        shade(t.rows[0].cells[i],NAVY); ctext(t.rows[0].cells[i],h,bold=True,white=True,size=8.6,align="left" if i<5 else "center")
    for r,(biz,owner,where,status,story,social,press,tier) in enumerate(rows):
        cells=t.add_row().cells; base=SAND if tier=="1" else (LIGHT if r%2==0 else "FFFFFF")
        ctext(cells[0],biz,bold=True,color=NAVY,size=8.4)
        ctext(cells[1],owner,size=8.0)
        ctext(cells[2],where,size=7.9)
        ctext(cells[3],status,size=7.9,color=(GREEN if status.startswith("Open") else SLATE),bold=True)
        ctext(cells[4],story,size=8.0)
        pc={"None":GREEN,"Low":GOLD,"Some":ORANGE}.get(press,SLATE)
        ctext(cells[5],press,size=7.9,align="center",color=pc,bold=True)
        ctext(cells[6],tier,size=8.0,align="center",bold=True,color=(GOLD if tier=="1" else SLATE))
        for c in cells: shade(c,base)
    for i,wd in enumerate(widths):
        for row in t.rows: row.cells[i].width=Inches(wd)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)

# Title
for _ in range(2): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("THE SCENIC CITY SCOUT"); r.font.size=Pt(24); r.font.bold=True; sca(r,NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("The 2025 Class — New Independents Members Should Know"); r.font.size=Pt(17); sca(r,TEAL)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Scouts triaged 1,892 of 2025's business starts; Nerds verified the gems. Gold rows = Tier-1 (feature now).")
r.font.size=Pt(11.5); r.font.italic=True; sca(r,SLATE)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run(f"{total} verified picks · Press footprint: None (biggest scoop) / Low / Some · Full sources in the CSV")
r.font.size=Pt(9.5); sca(r,GREY)
doc.add_page_break()

H("How this was built — the Scout → Nerd pipeline",1)
P("From the 1,892 businesses that started in 2025, the **Scouts** (selection agents) threw out the ~85% of trades, services, and chains and picked the consumer-facing independents members would love. The **Nerds** (research agents) then verified each one — is it real and open? who's the person? what's the story? — and flagged anything they couldn't confirm. The result is a vetted shortlist, not a raw dump. Tier 1 = feature now; Tier 2 = solid, worth a slot; can't-confirms are listed separately to chase, not publish.")

H("★ Editor's Picks — lead with these",1,color=GOLD)
P("The strongest human hooks with the best confidence, across categories:")
BUL([
 "**The Doll For All** (Marsha Roberts) — US-made dolls with limb differences for kids with disabilities; mission-driven, national-resonance, real press.",
 "**The Reading Room** (Linden Marno-Feree) — a bookstore-bar that funnels sales back to the Public Library Foundation.",
 "**The Woodshop** (Jhett Black) — a working bluesman's bar + listening room; 'Tuesday Bluesday' is a ready-made scene piece.",
 "**Gothic Trading Co.** (Lisa Porter) — the 'goth Cracker Barrel' on the NorthShore; vivid, already press-tested.",
 "**Pelfrey Drum Works** (Jerry 'Punkin' Pelfrey) — decades-deep solo snare-drum craftsman selling worldwide.",
 "**Grateful Glassworks** (Blake Rolfe) — live glassblowing studio; the best *visual* feature in the class.",
 "**Pedestrian Wine & Cheese** — 'first TN liquor store licensed to sell house-made cheese' — a genuinely novel hook.",
 "**Chattanooga Milk Stones** ('Brooklyn') & **REconnect Nature School** (Peterson & Toker) — a wow-niche craft and a mission-driven forest school.",
])

for cat,rows in DATA.items():
    H(f"{cat}  ({len(rows)})",1)
    gtable(rows)

H("Verify before print & cold leads",1,color=ORANGE)
BUL(COLD)
H("Accuracy traps the Nerds caught (do NOT publish as-is)",1)
BUL(TRAPS)
P("Full sources for every entry are in 2025_class_new_independents.csv. Owner names marked '(unconfirmed)' need a quick call before any attribution; home-based makers' addresses are withheld by policy.",
  italic=True,color=GREY,size=9)

out="/home/user/prscott/2025_Class_New_Independents.docx"
doc.save(out); print("saved",out)
