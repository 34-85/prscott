#!/usr/bin/env python3
"""Brief Scout Media — Portfolio view. How Scenic City Insider, Encore, and Mise en Place
roll up: model contrast, combined revenue shape, shared infrastructure, and the launch sequence."""
import os, matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY="1B3A4B"; TEAL="2A9D8F"; CLARET="6E2A35"; AMBER="B5722E"; INK="23282E"; SLATE="5A6470"; GREY="8A8170"; LIGHT="EEF3F4"
ASSETS="/home/user/prscott/assets"; os.makedirs(ASSETS,exist_ok=True)
plt.rcParams.update({"font.family":"DejaVu Sans","figure.dpi":160})

# Illustrative combined revenue ($K) by year
YEARS=["Year 1","Year 2","Year 3"]
SCI=[0,50,150]; ENC=[65,360,950]; MEP=[30,90,160]
def chart():
    fig,ax=plt.subplots(figsize=(8.8,4.6))
    import numpy as np
    x=range(len(YEARS))
    b1=ax.bar(x,SCI,color="#"+TEAL,label="Scenic City Insider (reader-funded)")
    b2=ax.bar(x,ENC,bottom=SCI,color="#"+CLARET,label="Encore (reader-funded)")
    b3=ax.bar(x,[SCI[i]+ENC[i] for i in range(3)],width=0) # spacer (no-op)
    b3=ax.bar(x,MEP,bottom=[SCI[i]+ENC[i] for i in range(3)],color="#"+AMBER,label="Mise en Place (sponsored)")
    for i in range(3):
        tot=SCI[i]+ENC[i]+MEP[i]
        ax.text(i,tot+18,f"~${tot}K",ha="center",fontweight="bold",color="#"+INK,fontsize=10)
    ax.set_xticks(list(x)); ax.set_xticklabels(YEARS,fontsize=10)
    ax.set_ylabel("Illustrative net revenue ($K)",fontsize=9,color="#"+SLATE)
    ax.set_ylim(0,1300)
    ax.legend(loc="upper left",fontsize=8,frameon=False)
    for s in ("top","right"): ax.spines[s].set_visible(False)
    ax.set_title("Brief Scout Media — illustrative combined revenue (conservative)",fontweight="bold",color="#"+NAVY,fontsize=12)
    fig.tight_layout(); fig.savefig(f"{ASSETS}/fig_portfolio_rev.png",bbox_inches="tight"); plt.close(fig)
chart(); print("chart done")

doc=Document(); sec=doc.sections[0]
for m in ("top_margin","bottom_margin"): setattr(sec,m,Inches(0.7))
for m in ("left_margin","right_margin"): setattr(sec,m,Inches(0.85))
nm=doc.styles["Normal"]; nm.font.name="Georgia"; nm.font.size=Pt(10.3); nm.paragraph_format.space_after=Pt(6); nm.paragraph_format.line_spacing=1.16
def sca(r,h): r.font.color.rgb=RGBColor.from_string(h)
def shade(c,h):
    p=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); p.append(s)
def ctext(c,t,bold=False,color=None,size=8.6,white=False,align="left"):
    c.text=""; p=c.paragraphs[0]; p.alignment={"left":WD_ALIGN_PARAGRAPH.LEFT,"center":WD_ALIGN_PARAGRAPH.CENTER}[align]
    for i,ch in enumerate(t.split("**")):
        if not ch: continue
        r=p.add_run(ch); r.font.size=Pt(size); r.font.bold=bold or (i%2==1)
        if white: r.font.color.rgb=RGBColor.from_string("FFFFFF")
        elif color: r.font.color.rgb=RGBColor.from_string(color)
def H(t,level=1,color=None):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(12 if level==1 else 7); p.paragraph_format.space_after=Pt(3)
    r=p.add_run(t); r.font.bold=True
    if level==1:
        r.font.size=Pt(14.5); sca(r,color or NAVY)
        pPr=p._p.get_or_add_pPr(); b=OxmlElement("w:pBdr"); bt=OxmlElement("w:bottom")
        bt.set(qn("w:val"),"single"); bt.set(qn("w:sz"),"6"); bt.set(qn("w:space"),"2"); bt.set(qn("w:color"),TEAL); b.append(bt); pPr.append(b)
    else: r.font.size=Pt(11.5); sca(r,color or INK)
def P(t,italic=False,color=None,size=10.3):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(6)
    for i,ch in enumerate(t.split("**")):
        if not ch: continue
        r=p.add_run(ch); r.font.size=Pt(size); r.font.italic=italic
        if i%2==1: r.font.bold=True
        if color: sca(r,color)
def BUL(items,size=9.9):
    for it in items:
        p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(3)
        for i,ch in enumerate(it.split("**")):
            if not ch: continue
            r=p.add_run(ch); r.font.size=Pt(size)
            if i%2==1: r.font.bold=True
def FIG(path,cap,width=7.0):
    doc.add_picture(path,width=Inches(width)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    P(cap,italic=True,color=GREY,size=8.5); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
def TABLE(headers,rows,widths,fontsize=8.5,hdr=NAVY):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style="Table Grid"
    for i,h in enumerate(headers):
        shade(t.rows[0].cells[i],hdr); ctext(t.rows[0].cells[i],h,bold=True,white=True,size=fontsize+0.2)
    for r,row in enumerate(rows):
        cells=t.add_row().cells; base=LIGHT if r%2==0 else "FFFFFF"
        for i,val in enumerate(row):
            ctext(cells[i],val,size=fontsize,color=INK if i==0 else None,bold=(i==0)); shade(cells[i],base)
    for i,w in enumerate(widths):
        for row in t.rows: row.cells[i].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)

doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("BRIEF SCOUT MEDIA"); r.font.size=Pt(27); r.font.bold=True; sca(r,NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("The Portfolio — a town, a generation, and a table"); r.font.size=Pt(15); r.font.italic=True; sca(r,TEAL)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("How Scenic City Insider, Encore & Mise en Place roll up  ·  all figures illustrative & conservative")
r.font.size=Pt(9); r.font.italic=True; sca(r,GREY)
doc.add_paragraph()

H("One company, three audiences",1)
P("Brief Scout Media isn't one product — it's a small house of titles on shared infrastructure, each serving a different audience with a different model. The discipline that holds it together is the **bright line**: the reader-funded titles keep their souls ad-free, and the sponsored title is openly, honestly sponsored. Same values, different engines.")
TABLE(["Title","Audience","Model","The promise"],
[["Scenic City Insider","A town — Chattanooga","Reader-funded, no ads","Know the city deeply; the reader is the customer"],
 ["Encore","A generation — accomplished 55–75, national","Reader-funded, no ads","Build what's next; members are the only customer"],
 ["Mise en Place","A table — Chattanooga food & drink","Openly sponsored + events","Eat & drink well; sponsors labeled, picks never for sale"]],
[1.55,2.4,1.7,2.4],fontsize=8.5)
P("Two reader-funded titles for trust and depth; one sponsored title for cash flow and joy. The sponsored engine can even help fund the patient ones.",italic=True,color=SLATE,size=9.8)

H("The combined revenue shape",1)
FIG(f"{ASSETS}/fig_portfolio_rev.png","Illustrative, conservative — Encore is the scale driver, MeP the early local cash flow, SCI the patient premium soul.")
TABLE(["Year","SCI","Encore","MeP","Combined (illustrative)"],
[["Year 1","~$0 (build)","~$65K","~$30K","~$95K"],
 ["Year 2","~$50K","~$360K","~$90K","~$500K"],
 ["Year 3","~$150K","~$950K","~$160K","~$1.26M"]],
[1.0,1.1,1.2,1.1,1.9],fontsize=8.6)
P("Read it as shape, not forecast. Each title's revenue lives in its own model doc; this simply stacks them. **Encore carries the ceiling** (national + affluent + events), **MeP delivers early local cash and joy**, **SCI compounds slowly toward its 200×$1,000 premium core.**",size=9.8,color=SLATE)

doc.add_page_break()
H("Shared infrastructure — why a house beats three startups",1)
BUL([
 "**One production system** — the HITL + AI workflow and the Scouts-and-Nerds research pipeline serve all three titles.",
 "**One stack** — beehiiv, the brand toolkit, the events playbook, the membership plumbing — built once, reused.",
 "**Cross-pollination** — SCI and MeP share Chattanooga ground (cross-promote freely); Encore brings the national, affluent reach. A reader can belong to more than one.",
 "**One founder voice & credibility** — the retired-CEO-and-author standing underwrites all three; the certifications (WSET/bourbon) feed MeP specifically.",
 "**Shared events muscle** — the dinner/retreat/supper-club operating model transfers across titles.",
])

H("The real constraint: founder bandwidth — so sequence it",1)
P("You cannot run three weeklies, three events calendars, a credential ladder, and a book at once, solo. The portfolio is a *destination*, reached in order. A sensible sequence:")
TABLE(["Phase","Lead with","Why","Add when ready"],
[["Now → Yr 1","Scenic City Insider (the soul)","It's the flagship and the trust engine; patient by design (little Y1 revenue)","Open founding-member waitlists for all three"],
 ["Yr 1 → 2","Mise en Place (cash + joy)","Fastest to revenue (events + sponsors), ties to your passion & certs, shares Chattanooga ground with SCI","Stand up the Supper Club; begin SCI membership"],
 ["Yr 2 → 3","Encore (the ceiling)","Most personal and the biggest revenue ceiling; needs the network and your time","Cohorts, retreats, the Founders' Circle"]],
[1.0,1.7,2.7,2.0],fontsize=8.4)
P("Alternative: if energy pulls you to MeP first (passion + fast cash), lead there and let SCI build alongside — the order is yours; the point is *one at a time gets the founder's full heart*.",italic=True,color=SLATE,size=9.8)

H("What the house adds up to",1)
P("Three honest titles, two engines, one set of values: **a town, a generation, and a table.** Conservatively, the portfolio reaches roughly **$1M+ in combined revenue by Year 3** — but the truer measure is that each title is something you'd be proud to make, funded by the people it serves (or by clearly-labeled sponsors who never buy a word). Built once as a house, on shared rails, it's far more achievable than three separate startups — and it's unmistakably yours.")
P("Each title's detail lives in its own docs: the business plan & manifesto (SCI), the concept brief & revenue model (Encore), and the one-pager, samples & revenue model (MeP).",italic=True,color=GREY,size=9)

out="/home/user/prscott/Brief_Scout_Media_Portfolio.docx"
doc.save(out); print("saved",out)
