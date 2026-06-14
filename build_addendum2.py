#!/usr/bin/env python3
"""Second Strategic Addendum: the narrow community strategy (celebrate local businesses/makers/
builders), assessment, SWOT, optimizations, and layered monetization."""
import os
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.ticker import FuncFormatter
from matplotlib.patches import FancyArrowPatch, Rectangle, FancyBboxPatch
import numpy as np

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY="1B3A4B"; TEAL="2A9D8F"; ORANGE="E76F51"; SAND="E9C46A"; SLATE="415A6B"
GREEN="2E7D52"; RED="C0492F"; LIGHT="EEF3F4"; GREY="6B7B85"
ASSETS="/home/user/prscott/assets"; os.makedirs(ASSETS,exist_ok=True)
plt.rcParams.update({"font.family":"DejaVu Sans","font.size":11,"axes.edgecolor":"#"+SLATE,
 "axes.labelcolor":"#"+NAVY,"axes.titlecolor":"#"+NAVY,"axes.titleweight":"bold",
 "xtick.color":"#"+SLATE,"ytick.color":"#"+SLATE,"axes.grid":True,"grid.color":"#DDE5E8",
 "grid.linewidth":0.8,"figure.dpi":160})
def usd(x,p=None): return f"${x/1000:.0f}K" if abs(x)>=1000 else f"${x:.0f}"

# ---- focused-model illustrative Year-3 mix ----
reader_mem=110000; events=110000; products=50000; biz=40000
total=reader_mem+events+products+biz
reader_funded=reader_mem+events+products

# ===== CHART 1: niche vs general =====
def chart_niche():
    fig,(a1,a2)=plt.subplots(1,2,figsize=(8.4,4.0))
    a1.bar(["General\ninterest","Niche /\ncommunity"],[21.5,32],color=["#"+SLATE,"#"+TEAL])
    for i,v in enumerate([21.5,32]): a1.text(i,v+0.6,f"{v:.0f}%",ha="center",fontweight="bold",color="#"+NAVY)
    a1.set_ylabel("Avg open rate"); a1.set_ylim(0,40); a1.set_title("Open Rate")
    a2.bar(["Audience\ngrowth","Revenue\nper sub"],[1.27,2.3],color=["#"+SAND,"#"+ORANGE])
    a2.axhline(1.0,ls="--",color="#"+SLATE,lw=1)
    a2.text(1.5,1.03,"general-interest baseline = 1.0×",ha="right",fontsize=7.5,color="#"+SLATE)
    for i,v in enumerate([1.27,2.3]): a2.text(i,v+0.05,f"{v:.2f}×",ha="center",fontweight="bold",color="#"+NAVY)
    a2.set_ylim(0,2.7); a2.set_title("Niche Advantage (relative)")
    fig.suptitle("Going Narrow Outperforms Going Broad",fontweight="bold",color="#"+NAVY,y=1.02)
    fig.tight_layout(); fig.savefig(f"{ASSETS}/fig2_niche.png",bbox_inches="tight"); plt.close(fig)

# ===== CHART 2: flywheel =====
def chart_flywheel():
    fig,ax=plt.subplots(figsize=(7.0,6.0)); ax.set_xlim(-1.5,1.5); ax.set_ylim(-1.5,1.5)
    ax.axis("off"); ax.set_aspect("equal")
    nodes=[
        ("Celebrate local\nbusinesses & makers\n(content people love)", 0, 1.05, "#"+NAVY),
        ("Engaged community\ngrows (trust +\nword of mouth)", 1.05, -0.35, "#"+TEAL),
        ("They gather & buy\n(events, markets,\nproducts, membership)", 0, -1.15, "#"+ORANGE),
        ("Businesses participate\n(booths, partners,\ncollabs) → more stories", -1.05, -0.35, "#"+SAND),
    ]
    for txt,x,y,c in nodes:
        box=FancyBboxPatch((x-0.5,y-0.32),1.0,0.64,boxstyle="round,pad=0.03,rounding_size=0.08",
                           fc=c,ec="white",lw=2)
        ax.add_patch(box)
        ax.text(x,y,txt,ha="center",va="center",color="white",fontsize=8.5,fontweight="bold")
    # circular arrows between consecutive nodes
    pts=[(0,1.05),(1.05,-0.35),(0,-1.15),(-1.05,-0.35)]
    for i in range(4):
        x1,y1=pts[i]; x2,y2=pts[(i+1)%4]
        arr=FancyArrowPatch((x1*0.62,y1*0.62),(x2*0.62,y2*0.62),
                            connectionstyle="arc3,rad=0.32",arrowstyle="-|>",
                            mutation_scale=20,lw=2.2,color="#"+SLATE)
        ax.add_patch(arr)
    ax.text(0,0,"THE\nCOMMUNITY\nFLYWHEEL",ha="center",va="center",fontsize=11,
            fontweight="bold",color="#"+NAVY)
    fig.tight_layout(); fig.savefig(f"{ASSETS}/fig2_flywheel.png",bbox_inches="tight"); plt.close(fig)

# ===== CHART 3: bright line =====
def chart_brightline():
    fig,ax=plt.subplots(figsize=(8.6,4.4)); ax.axis("off"); ax.set_xlim(0,10); ax.set_ylim(0,10)
    ax.add_patch(Rectangle((0.2,0.4),4.5,9.2,fc="#EAF3EE",ec="#"+GREEN,lw=2))
    ax.add_patch(Rectangle((5.3,0.4),4.5,9.2,fc="#FBECE8",ec="#"+RED,lw=2))
    ax.text(2.45,9.0,"ALLOWED",ha="center",fontweight="bold",color="#"+GREEN,fontsize=14)
    ax.text(2.45,8.45,"buy a booth, a listing, a ticket,\na job post, or a product",
            ha="center",color="#"+GREEN,fontsize=8.5,style="italic")
    ax.text(7.55,9.0,"BANNED",ha="center",fontweight="bold",color="#"+RED,fontsize=14)
    ax.text(7.55,8.45,"buy a sentence, a profile,\nor favorable framing",
            ha="center",color="#"+RED,fontsize=8.5,style="italic")
    allowed=["Vendor / booth fees at events","Paid (labeled) directory listing",
             "Member discount-card participation","Event tickets / experience partner",
             "Wholesale merch collabs","Job-board posts & affiliate links"]
    banned=["Paid 'profiles' dressed as articles","Sponsored newsletter placement",
            "'Featured business' written by staff","Advertorial / native ads",
            "Any quid-pro-quo for coverage","Pay-to-play review or ranking"]
    for i,t in enumerate(allowed):
        ax.text(0.5,7.6-i*1.15,"✓ "+t,color="#"+GREEN,fontsize=9.2,va="center")
    for i,t in enumerate(banned):
        ax.text(5.6,7.6-i*1.15,"✗ "+t,color="#"+RED,fontsize=9.2,va="center")
    ax.text(5.0,0.05,'"Money can buy a booth, a listing, a ticket, a job post, or a product — never a sentence."',
            ha="center",fontsize=9,fontweight="bold",color="#"+NAVY)
    fig.tight_layout(); fig.savefig(f"{ASSETS}/fig2_brightline.png",bbox_inches="tight"); plt.close(fig)

# ===== CHART 4: SWOT quadrant =====
def chart_swot():
    fig,ax=plt.subplots(figsize=(8.6,6.2)); ax.axis("off"); ax.set_xlim(0,10); ax.set_ylim(0,10)
    quads=[(0.2,5.1,"#"+TEAL,"STRENGTHS",[
        "Reader-aligned trust (no ads) — rare moat","Self-replenishing, high-engagement beat",
        "Convenable, commercially active community","Authentic local voice AI can't copy",
        "Events & commerce native to the niche"]),
        (5.1,5.1,"#"+NAVY,"WEAKNESSES",[
        "Single-metro audience ceiling","Solo-founder capacity / burnout risk",
        "Unproven local willingness-to-pay","Must resist pay-to-play pressure",
        "Many small relationships to manage"]),
        (0.2,0.2,"#"+SAND,"OPPORTUNITIES",[
        "Local-news vacuum + shop-local (78%)","Own the 'local doers' beat no one holds",
        "Maker markets = high-integrity revenue","B2B partner program (directory, card)",
        "Transplant inflow wants discovery"]),
        (5.1,0.2,"#"+ORANGE,"THREATS",[
        "Incumbent adds paid/community tier","Downturn squeezes small businesses",
        "Over-reliance on few anchors","Trust collapse if independence slips",
        "Platform / key-person dependency"])]
    for x,y,c,title,items in quads:
        ax.add_patch(FancyBboxPatch((x,y),4.7,4.7,boxstyle="round,pad=0.02,rounding_size=0.1",
                                    fc=c,ec="white",lw=2,alpha=0.93))
        ax.text(x+2.35,y+4.35,title,ha="center",color="white",fontweight="bold",fontsize=12)
        for i,t in enumerate(items):
            ax.text(x+0.2,y+3.7-i*0.72,"• "+t,color="white",fontsize=8.6,va="center")
    fig.suptitle("SWOT — The Narrow Community Strategy",fontweight="bold",color="#"+NAVY,y=0.99)
    fig.tight_layout(); fig.savefig(f"{ASSETS}/fig2_swot.png",bbox_inches="tight"); plt.close(fig)

# ===== CHART 5: revenue layers =====
def chart_layers():
    fig,ax=plt.subplots(figsize=(8.6,2.8))
    segs=[("Reader membership",reader_mem,"#"+NAVY),("Events",events,"#"+TEAL),
          ("Products / maker collabs",products,"#"+ORANGE),("Business-partner program",biz,"#"+SAND)]
    left=0
    for lab,val,c in segs:
        ax.barh(0,val,left=left,color=c,height=0.5,edgecolor="white")
        ax.text(left+val/2,0,f"{lab}\n{usd(val)}",ha="center",va="center",color="white",
                fontsize=8.5,fontweight="bold")
        left+=val
    ax.set_xlim(0,total); ax.set_ylim(-0.5,0.6); ax.set_yticks([])
    ax.xaxis.set_major_formatter(FuncFormatter(usd)); ax.grid(False)
    ax.annotate("",xy=(reader_funded,0.42),xytext=(0,0.42),
                arrowprops=dict(arrowstyle="<->",color="#"+GREEN,lw=1.6))
    ax.text(reader_funded/2,0.5,f"Reader-funded ≈ {reader_funded/total*100:.0f}% (protects trust)",
            ha="center",color="#"+GREEN,fontsize=8.5,fontweight="bold")
    ax.annotate("",xy=(total,-0.4),xytext=(reader_funded,-0.4),
                arrowprops=dict(arrowstyle="<->",color="#"+SLATE,lw=1.6))
    ax.text((reader_funded+total)/2,-0.48,f"Business-facing ≈ {biz/total*100:.0f}%",
            ha="center",color="#"+SLATE,fontsize=8.5)
    ax.set_title(f"Illustrative Focused-Model Year-3 Revenue ({usd(total)})")
    fig.tight_layout(); fig.savefig(f"{ASSETS}/fig2_layers.png",bbox_inches="tight"); plt.close(fig)

for fn in (chart_niche,chart_flywheel,chart_brightline,chart_swot,chart_layers): fn()
print("charts done")

# ============================================================================
doc=Document()
normal=doc.styles["Normal"]; normal.font.name="Calibri"; normal.font.size=Pt(11)
normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.12
def sc(r,h): r.font.color.rgb=RGBColor.from_string(h)
def shade(cell,h):
    p=cell._tc.get_or_add_tcPr(); s=OxmlElement("w:shd")
    s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); p.append(s)
def ctext(cell,t,bold=False,color=None,size=10,align="left",white=False):
    cell.text=""; p=cell.paragraphs[0]
    p.alignment={"left":WD_ALIGN_PARAGRAPH.LEFT,"center":WD_ALIGN_PARAGRAPH.CENTER,
                 "right":WD_ALIGN_PARAGRAPH.RIGHT}[align]
    r=p.add_run(t); r.font.size=Pt(size); r.font.bold=bold
    if white: r.font.color.rgb=RGBColor.from_string("FFFFFF")
    elif color: r.font.color.rgb=RGBColor.from_string(color)
def H(t,level=1):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(14 if level==1 else 10)
    p.paragraph_format.space_after=Pt(4); r=p.add_run(t); r.font.bold=True
    if level==1:
        r.font.size=Pt(17); sc(r,NAVY)
        pPr=p._p.get_or_add_pPr(); b=OxmlElement("w:pBdr"); bt=OxmlElement("w:bottom")
        bt.set(qn("w:val"),"single"); bt.set(qn("w:sz"),"6"); bt.set(qn("w:space"),"2"); bt.set(qn("w:color"),TEAL)
        b.append(bt); pPr.append(b)
    elif level==2: r.font.size=Pt(13.5); sc(r,TEAL)
    else: r.font.size=Pt(11.5); sc(r,SLATE)
def P(t,italic=False,color=None,size=11,align="left",space_after=6):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(space_after)
    p.alignment={"left":WD_ALIGN_PARAGRAPH.LEFT,"center":WD_ALIGN_PARAGRAPH.CENTER,
                 "right":WD_ALIGN_PARAGRAPH.RIGHT,"just":WD_ALIGN_PARAGRAPH.JUSTIFY}[align]
    for i,ch in enumerate(t.split("**")):
        if not ch: continue
        r=p.add_run(ch); r.font.size=Pt(size); r.font.italic=italic
        if i%2==1: r.font.bold=True
        if color: sc(r,color)
    return p
def BUL(items):
    for it in items:
        p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(3)
        for i,ch in enumerate(it.split("**")):
            if not ch: continue
            r=p.add_run(ch); r.font.size=Pt(10.5)
            if i%2==1: r.font.bold=True
def TABLE(headers,rows,widths=None,fontsize=9,caption=None,first_col_bold=False,header_fill=NAVY):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style="Table Grid"
    for i,h in enumerate(headers):
        shade(t.rows[0].cells[i],header_fill)
        ctext(t.rows[0].cells[i],h,bold=True,white=True,size=fontsize,align="left" if i==0 else "center")
    for r,row in enumerate(rows):
        cells=t.add_row().cells; fill=LIGHT if r%2==0 else "FFFFFF"
        for i,v in enumerate(row):
            shade(cells[i],fill)
            ctext(cells[i],str(v),bold=(first_col_bold and i==0),size=fontsize,
                  align="left" if i==0 else "center",color=NAVY if (first_col_bold and i==0) else None)
    if widths:
        for i,w in enumerate(widths):
            for row in t.rows: row.cells[i].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    if caption: P(caption,italic=True,color=GREY,size=8.5,space_after=10)
def FIG(path,caption,width=6.4):
    doc.add_picture(path,width=Inches(width)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    P(caption,italic=True,color=GREY,size=8.5,align="center",space_after=12)
for s in doc.sections:
    s.top_margin=Inches(0.9); s.bottom_margin=Inches(0.9); s.left_margin=Inches(0.9); s.right_margin=Inches(0.9)

# ---- title ----
for _ in range(3): doc.add_paragraph()
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=t.add_run("BRIEF SCOUT MEDIA"); r.font.size=Pt(32); r.font.bold=True; sc(r,NAVY)
s=doc.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=s.add_run("Second Strategic Addendum"); r.font.size=Pt(20); sc(r,TEAL)
s=doc.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=s.add_run("The Narrow Community Strategy: Championing Chattanooga's Businesses, Makers & Builders")
r.font.size=Pt(13); r.font.italic=True; sc(r,SLATE)
s=doc.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=s.add_run("Assessment · Scope Definition · SWOT · Optimizations · Layered Monetization")
r.font.size=Pt(11.5); r.font.italic=True; sc(r,SLATE)
for _ in range(2): doc.add_paragraph()
s=doc.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=s.add_run("Companion working draft to the Brief Scout Media Business Plan  ·  June 2026")
r.font.size=Pt(10.5); sc(r,GREY)
doc.add_page_break()

# ---- 1 thesis ----
H("1. The Strategic Shift: Go Narrow to Build Community",1)
P("The original plan framed the Scenic City Scout as a broad \"what matters + what to do\" local "
  "brief. This addendum tests a sharper hypothesis the founder raised: **the narrower the focus, the "
  "stronger the community** — and specifically, that anchoring the Scout to **championing "
  "Chattanooga's local businesses, restaurants, bars, startups, makers, and builders** is the "
  "fastest route to a defensible, monetizable community. The research is emphatic: it is.",align="just")
P("Niche beats broad on every metric that matters. Per beehiiv's platform data, clearly-defined "
  "niche newsletters grow audiences **~27% faster** and earn **~2.3× the revenue** of "
  "general-interest titles, and trust-heavy niches routinely post **30%+ open rates** versus a "
  "~21.5% industry average. Crucially, the **\"what's new / what's open / what closed\"** local "
  "business-and-food beat is consistently the highest-engagement content in local media — it is the "
  "engine behind Eater's flagship \"Heatmap,\" Axios Local's 41%-open food columns, and 6AM City's "
  "growth to 2M+ subscribers (built, tellingly, by *avoiding* politics and crime and celebrating "
  "local). A narrow, positive, local-business identity is not a limitation — it is the product.",
  align="just")
FIG(f"{ASSETS}/fig2_niche.png",
    "Figure B1. Niche/community newsletters outperform general-interest on open rate, growth, and revenue per subscriber. Source: beehiiv, Letterhead, Kit (2024–2026).")

# ---- 2 scope ----
H("2. What This Focus Means — and What It Doesn't",1)
P("\"Championing local businesses, restaurants, bars, startups, makers, and builders\" is a wide "
  "enough field to never run out of material — there is a constant stream of openings, closings, "
  "new makers, expansions, and milestones — but it must be **defined by what it excludes** as much "
  "as what it includes. Discipline about scope is what creates the identity and the trust.",align="just")
TABLE(["This IS the focus (in scope)","This is NOT the focus (out of scope)"],
[
 ["The people who build the local economy: independent shops, restaurants, bars, breweries, makers, artisans, founders, builders/developers, chefs","Hard news, crime, politics, disasters, and partisan controversy (the things 6AM City deliberately avoids)"],
 ["\"What's new / opening / closing / expanding\" — the renewable engagement beat","Generic, undifferentiated event listings or a calendar dump anyone can copy"],
 ["Origin stories, the craft, the people behind the counter, how they built it","National/chain news or syndicated content with no local hand"],
 ["Honest, independent recommendations and context (incl. when something disappoints)","Boosterism / chamber-brochure cheerleading that readers can't trust"],
 ["Helping readers discover, support, and spend locally; helping makers find customers","\"Pay-to-play\" coverage, paid profiles, advertorial, or PR reprints (see Section 6)"],
 ["Convening the community in person (markets, meet-the-maker nights, mixers)","Being an advertising vehicle dressed up as editorial"],
],widths=[3.15,3.15],fontsize=9,first_col_bold=False,
caption="Table B1. The scope is defined by a clear in/out boundary. The exclusions are what make the inclusions trustworthy.")
P("**The most important boundary** is the last one. The instant readers suspect that a business "
  "appears because it *paid* to appear, the entire premise — an independent guide you can trust — "
  "collapses. Celebrating local businesses editorially and monetizing local businesses commercially "
  "must be kept rigorously separate (Section 6).")

# ---- 3 assessment ----
H("3. Will It Work? — The Assessment",1)
P("**Verdict: yes, with high confidence — provided the trust boundary holds and the founder manages "
  "the single-market ceiling by expanding into adjacent verticals over time.** Four evidence streams "
  "converge:",align="just")
BUL([
 "**The format is proven and monetizable at small scale.** Side Dish with Schniper celebrates the local "
 "Colorado Springs food scene and earns **~$70K/yr at just ~3,000 subscribers**; The Food Section runs a "
 "**72% open rate** and **$50K+/yr** on a regional-food niche. These are near-exact templates for the model.",
 "**The audience is unusually engaged and commercially active.** \"Shop local\" is a real, money-backed "
 "movement: independents recirculate far more locally ($0.68 vs $0.43 per dollar), **78% of consumers say "
 "they'd pay more to support small businesses**, and Small Business Saturday moves ~$22B in a single day. "
 "This community shows up, buys, and cross-promotes (word-of-mouth is small business's #1 channel).",
 "**Chattanooga over-indexes for exactly this.** The Chattanooga Market draws **~6,000 visitors every "
 "weekend** with 100–300+ maker/food vendors; the Southside/Main Street corridor is a dense independent "
 "dining/brewery scene; Chattanooga Whiskey is a flagship local-maker success; and the startup ecosystem "
 "(CO.LAB, Brickyard, Dynamo, the nation's 3rd-largest INCubator, ~$1.8–2B in exits) supplies a steady "
 "founder/builder storyline. The subject matter and the audience are both already here.",
 "**The lane is open.** NOOGAtoday is broad and ad-funded; the Times Free Press is a paywalled legacy "
 "daily; The Pulse's events vehicle is gone. No one owns the \"champion of local doers\" beat as a "
 "reader-funded, community-convening brand.",
])
P("**The principal risk is the audience ceiling** of a single mid-sized metro narrowed further to a "
  "niche-of-a-niche. The research-backed answer is the sequence already in Addendum 1: **start narrow "
  "to win the community, then expand into adjacent verticals** (food edition, maker edition) — never "
  "broaden back into general news, which operators warn is \"a recipe for disaster.\"")
FIG(f"{ASSETS}/fig2_flywheel.png",
    "Figure B2. The community flywheel: celebrating local businesses earns an engaged community, which gathers and buys, which gives businesses reasons to participate, which generates more stories.",width=5.0)

# ---- 4 SWOT ----
H("4. SWOT Analysis",1)
FIG(f"{ASSETS}/fig2_swot.png","Figure B3. SWOT for the narrow community strategy.",width=6.6)
TABLE(["Quadrant","Detail & strategic implication"],
[
 ["Strengths","No-ads trust is a genuine moat; the local-business beat is self-replenishing and high-engagement; the community is convenable and commercially active; events and commerce are native to the niche, not bolted on."],
 ["Weaknesses","A single metro caps audience; a solo founder has finite capacity; local willingness-to-pay is unproven; the model must continuously resist pay-to-play pressure from the very businesses it covers."],
 ["Opportunities","A local-news vacuum plus strong shop-local sentiment; an uncontested \"local doers\" position; high-integrity revenue from maker markets and a B2B partner program; a transplant inflow hungry for local discovery; clear vertical-expansion paths."],
 ["Threats","An incumbent (6AM City) could add a paid/community tier; a downturn squeezes small-business participation; over-reliance on a few anchor businesses/events; and — the existential one — a trust collapse if editorial independence ever slips."],
],widths=[1.3,5.0],fontsize=9,first_col_bold=True,
caption="Table B2. SWOT detail. The strategy's biggest strength (trust) and biggest threat (losing it) are two sides of the same coin.")

# ---- 5 optimizations ----
H("5. Optimizations & Recommended Changes",1)
P("The research suggests several concrete refinements to sharpen the strategy and de-risk it:")
TABLE(["Recommendation","Why (research-backed)"],
[
 ["Sharpen positioning to one sentence: “Chattanooga's champion of the people who build it.”",
  "Clear niche identity drives ~27% faster growth and higher conversion; specificity is the asset."],
 ["Make “what's new / opening / closing” the recurring spine of every issue",
  "It is consistently the highest-engagement local content (Eater Heatmap; Axios food columns at 41% open)."],
 ["Publish an Editorial Independence Policy and market “we can't be bought” as a feature",
  "<2% of readers spot native ads and those who do trust the outlet less; turn the constraint into the brand."],
 ["Build the business/partner program as a separate, clearly-labeled product, walled from editorial",
  "Protects trust (the whole asset); mirrors A Little Beacon Blog's labeled directory and INN independence standards."],
 ["Anchor events to the existing maker economy (partner with the Chattanooga Market; meet-the-maker nights)",
  "Leverages a ready 6,000/weekend audience; booth/vendor fees are the highest-integrity business revenue."],
 ["Use member-sourced tips, nominations, and spotlights to scale content and deepen belonging",
  "Community-led growth is “extraordinarily difficult to replicate” — it is the moat."],
 ["Segment the list (foodies / founders / makers) for tailored sends",
  "Lifts open rates and free→paid conversion; supports later vertical spin-offs."],
 ["Keep reader revenue the majority of the total; cap business-facing revenue",
  "Defector (~95% subscriber) and Hell Gate (ads minor) show reader-dominance is what protects independence."],
],widths=[2.7,3.6],fontsize=8.8,first_col_bold=False,
caption="Table B3. Research-backed optimizations.")

# ---- 6 monetization layering ----
H("6. Layering In the Revenue Streams",1)
P("The narrow focus does not replace the three-pillar model from the primary plan — it "
  "**supercharges** it, because a passionate business-and-maker community is exactly the audience "
  "that joins, attends, and buys. It also unlocks a *new, optional* business-facing line. Everything "
  "hangs on one governing rule:",align="just")
P("“Money can buy a booth, a listing, a ticket, a job post, or a product — it can never buy a "
  "sentence.”", color=NAVY, size=12.5, align="center", space_after=8)
FIG(f"{ASSETS}/fig2_brightline.png",
    "Figure B4. The bright line that keeps the model honest: businesses can pay for venues, exposure surfaces, and services — never for editorial coverage or framing.",width=6.7)
TABLE(["Stream","Who pays","For what (allowed)","Indicative pricing"],
[
 ["Reader membership (engine)","Readers","Members-only content, member event pricing, discount card, community","$8–$20/mo; keep majority of revenue"],
 ["Events","Attendees + makers","Tickets; vendor/booth fees at markets & meet-the-maker nights","$25–$50 ticket; $25–$300+ booth"],
 ["Products","Buyers","Merch; wholesale maker collabs; paid city/maker guides","Goods margin; guides near-100%"],
 ["Business/partner program (new)","Businesses","Labeled directory/marketplace listing; discount-card participation; first access to booths","~$30–$40/mo small-biz; premium tiers higher"],
 ["Utility commerce","Businesses / readers","Local job board; affiliate links in clearly-marked resource sections","$100–$200/job post; affiliate %"],
],widths=[1.5,1.2,2.4,1.2],fontsize=8.6,first_col_bold=True,
caption="Table B4. How each revenue stream layers onto the community focus — all editorially inert.")
P("**The new B2B layer, done safely.** A separate “business member / partner” program "
  "(distinct from reader membership) can sell a labeled local-business directory, discount-card "
  "participation, and priority access to event booths. Comparable local outlets price such programs "
  "from chamber-like ~$30–$40/month up to premium feature tiers (A Little Beacon Blog runs to "
  "$1,250/mo). The non-negotiable: it is a clearly-marked marketing product, partners get **zero** "
  "editorial consideration, and a published independence policy makes that promise enforceable.")
FIG(f"{ASSETS}/fig2_layers.png",
    "Figure B5. Illustrative focused-model Year-3 revenue. A business-partner line is added, but reader-funded streams stay ~87% of the total to protect independence.",width=6.9)
P(f"In this illustrative focused build, Year-3 revenue of ~{usd(total)} keeps reader-funded streams "
  f"(membership + events + products) at ~{reader_funded/total*100:.0f}% — preserving the trust that "
  f"makes the whole thing work — while a disciplined, walled business-partner program adds ~"
  f"{usd(biz)} of incremental, on-brand revenue. This is additive to, not a replacement for, the "
  f"core plan's projections.")

# ---- 7 deltas ----
H("7. Recommended Changes to the Core Plan",1)
BUL([
 "**Reposition the Scout** from broad general-interest to the explicit champion of Chattanooga's local businesses, makers, builders, and founders (narrower identity, stronger community).",
 "**Adopt the “what's new/open/closed” beat** as the editorial spine and add member-sourced spotlights.",
 "**Add a fourth, optional revenue line** — a walled business/partner program — capped so reader revenue stays dominant.",
 "**Re-anchor events** to the maker economy (Chattanooga Market partnership; meet-the-maker nights; booth fees).",
 "**Publish an Editorial Independence Policy** as both governance and marketing — the “we can't be bought” promise.",
 "**Plan vertical expansion** (food edition, maker edition) as the answer to the single-market ceiling — narrow first, then widen.",
])

# ---- 8 bottom line ----
H("8. Bottom Line",1)
P("Going narrower is the right call. A newsletter that becomes Chattanooga's trusted champion of its "
  "independent businesses, makers, and builders will out-engage and out-monetize a broad general "
  "brief, because it owns a passionate, convenable, commercially active community that no ad-funded "
  "incumbent can authentically serve. The same focus that draws the community also feeds the "
  "revenue: members join the cause, the community gathers and buys at events and markets, makers "
  "collaborate on products, and a carefully-walled business-partner program adds B2B revenue — all "
  "without ever selling a sentence. Guard the trust boundary, lean on the maker economy, and expand "
  "into verticals when the base is strong, and the narrow strategy becomes the deepest moat Brief "
  "Scout Media can build.",align="just")

out="/home/user/prscott/Brief_Scout_Media_Addendum_2_Community_Strategy.docx"
doc.save(out); print("saved",out)
