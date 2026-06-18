#!/usr/bin/env python3
"""Build the Brief Scout Media founder business plan as a Word document with charts."""
import os
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.ticker import FuncFormatter
import numpy as np

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ----------------------------------------------------------------------------
# Brand palette
# ----------------------------------------------------------------------------
NAVY   = "1B3A4B"   # deep slate-navy (primary)
TEAL   = "2A9D8F"   # scenic teal (secondary)
ORANGE = "E76F51"   # warm accent
SAND   = "E9C46A"   # sand/gold
SLATE  = "415A6B"   # muted slate
LIGHT  = "EEF3F4"   # light fill for zebra rows
GREY   = "6B7B85"

PALETTE = ["#"+NAVY, "#"+TEAL, "#"+ORANGE, "#"+SAND, "#"+SLATE, "#7FB5AE"]

ASSETS = "/home/user/prscott/assets"
os.makedirs(ASSETS, exist_ok=True)

plt.rcParams.update({
    "font.family": "DejaVu Sans",
    "font.size": 11,
    "axes.edgecolor": "#"+SLATE,
    "axes.labelcolor": "#"+NAVY,
    "axes.titlecolor": "#"+NAVY,
    "axes.titleweight": "bold",
    "xtick.color": "#"+SLATE,
    "ytick.color": "#"+SLATE,
    "axes.grid": True,
    "grid.color": "#DDE5E8",
    "grid.linewidth": 0.8,
    "figure.dpi": 160,
})

def usd(x, pos=None):
    if abs(x) >= 1000:
        return f"${x/1000:.0f}K"
    return f"${x:.0f}"

# ============================================================================
# FINANCIAL MODEL
# ============================================================================
years = ["Year 1", "Year 2", "Year 3"]

# Revenue by stream
membership = [16000, 54000, 120000]
events     = [9000, 40000, 105000]
products   = [6000, 22000, 55000]
total_rev  = [m+e+p for m, e, p in zip(membership, events, products)]   # 31k,116k,280k

# Audience
free_eoy    = [7000, 16000, 28000]
members_eoy = [350, 850, 1700]

# Operating expenses (excludes owner's draw)
opex = [10500, 34100, 83500]
owner_draw = [t - o for t, o in zip(total_rev, opex)]  # 20.5k, 81.9k, 196.5k

# Discovery-engine-adjusted (illustrative upside for Years 2-3; Year 1 held at base)
eng_membership = [membership[0], 60000, 135000]   # +intel-perk lift to conversion/retention
eng_events     = [events[0], 44000, 115000]       # easier, better-sourced events
eng_products   = [products[0], 24000, 60000]
eng_total      = [m+e+p for m, e, p in zip(eng_membership, eng_events, eng_products)]
eng_opex       = [opex[0], 30000, 76000]          # lower content/sourcing cost
eng_draw       = [t - o for t, o in zip(eng_total, eng_opex)]

# ----------------------------------------------------------------------------
# CHART 1 — Market funnel (TAM / SAM / SOM)
# ----------------------------------------------------------------------------
def chart_funnel():
    fig, ax = plt.subplots(figsize=(7.2, 4.3))
    labels = ["TAM\nMSA adults (18+)", "SAM\nEngaged, can-pay residents",
              "Free list (3-yr target)", "SOM\nPaying members (3-yr)"]
    vals = [470000, 55000, 28000, 1700]
    widths = [1.0, 0.62, 0.40, 0.20]
    colors = ["#"+NAVY, "#"+TEAL, "#"+SAND, "#"+ORANGE]
    y = np.arange(len(vals))[::-1]
    for yi, w, c, lab, v in zip(y, widths, colors, labels, vals):
        ax.barh(yi, w, height=0.72, color=c, left=(1-w)/2)
        ax.text(0.5, yi, f"{lab}\n{v:,.0f}", ha="center", va="center",
                color="white", fontsize=10, fontweight="bold")
    ax.set_xlim(0, 1); ax.set_ylim(-0.6, len(vals)-0.4)
    ax.axis("off")
    ax.set_title("Chattanooga Addressable Market — Funnel", pad=12)
    fig.tight_layout()
    fig.savefig(f"{ASSETS}/fig_funnel.png", bbox_inches="tight")
    plt.close(fig)

# ----------------------------------------------------------------------------
# CHART 2 — Revenue by stream (stacked) + total line
# ----------------------------------------------------------------------------
def chart_revenue_stack():
    fig, ax = plt.subplots(figsize=(7.2, 4.3))
    x = np.arange(len(years))
    b1 = ax.bar(x, membership, color="#"+NAVY, label="Membership")
    b2 = ax.bar(x, events, bottom=membership, color="#"+TEAL, label="Events")
    bot2 = [m+e for m, e in zip(membership, events)]
    b3 = ax.bar(x, products, bottom=bot2, color="#"+ORANGE, label="Products & commerce")
    for xi, t in zip(x, total_rev):
        ax.text(xi, t+6000, usd(t), ha="center", fontweight="bold", color="#"+NAVY)
    ax.set_xticks(x); ax.set_xticklabels(years)
    ax.yaxis.set_major_formatter(FuncFormatter(usd))
    ax.set_ylim(0, max(total_rev)*1.18)
    ax.set_title("Projected Revenue by Stream (No Sponsorship)")
    ax.legend(frameon=False, loc="upper left")
    fig.tight_layout()
    fig.savefig(f"{ASSETS}/fig_revenue_stack.png", bbox_inches="tight")
    plt.close(fig)

# ----------------------------------------------------------------------------
# CHART 3 — Year-3 revenue mix donut
# ----------------------------------------------------------------------------
def chart_mix():
    fig, ax = plt.subplots(figsize=(5.6, 4.3))
    vals = [membership[2], events[2], products[2]]
    labs = ["Membership", "Events", "Products & commerce"]
    cols = ["#"+NAVY, "#"+TEAL, "#"+ORANGE]
    w, t, a = ax.pie(vals, colors=cols, startangle=90,
                     wedgeprops=dict(width=0.42, edgecolor="white"),
                     autopct=lambda p: f"{p:.0f}%", pctdistance=0.78,
                     textprops=dict(color="white", fontweight="bold"))
    ax.legend(w, labs, frameon=False, loc="center", bbox_to_anchor=(0.5, -0.08), ncol=1)
    ax.set_title("Year-3 Revenue Mix")
    ax.text(0, 0, f"{usd(total_rev[2])}", ha="center", va="center",
            fontsize=15, fontweight="bold", color="#"+NAVY)
    fig.tight_layout()
    fig.savefig(f"{ASSETS}/fig_mix.png", bbox_inches="tight")
    plt.close(fig)

# ----------------------------------------------------------------------------
# CHART 4 — Audience growth (free list bars + members line)
# ----------------------------------------------------------------------------
def chart_audience():
    fig, ax = plt.subplots(figsize=(7.2, 4.3))
    x = np.arange(len(years))
    ax.bar(x, free_eoy, color="#"+TEAL, width=0.55, label="Free subscribers (EoY)")
    for xi, v in zip(x, free_eoy):
        ax.text(xi, v+700, f"{v:,}", ha="center", color="#"+NAVY, fontweight="bold")
    ax.set_xticks(x); ax.set_xticklabels(years)
    ax.set_ylabel("Free subscribers")
    ax.set_ylim(0, max(free_eoy)*1.2)
    ax2 = ax.twinx()
    ax2.plot(x, members_eoy, color="#"+ORANGE, marker="o", lw=2.5, label="Paying members (EoY)")
    for xi, v in zip(x, members_eoy):
        ax2.text(xi, v+60, f"{v:,}", ha="center", color="#"+ORANGE, fontweight="bold")
    ax2.set_ylabel("Paying members", color="#"+ORANGE)
    ax2.set_ylim(0, max(members_eoy)*1.35)
    ax2.grid(False)
    ax.set_title("Audience Growth: Free List & Paying Members")
    l1, lab1 = ax.get_legend_handles_labels()
    l2, lab2 = ax2.get_legend_handles_labels()
    ax.legend(l1+l2, lab1+lab2, frameon=False, loc="upper left")
    fig.tight_layout()
    fig.savefig(f"{ASSETS}/fig_audience.png", bbox_inches="tight")
    plt.close(fig)

# ----------------------------------------------------------------------------
# CHART 5 — Comparable operators: revenue vs list size (small-list, big-revenue)
# ----------------------------------------------------------------------------
def chart_comparables():
    fig, ax = plt.subplots(figsize=(7.4, 4.5))
    # (name, subs, annual_rev, model)
    data = [
        ("Catskill Crew", 18000, 620000, "reader/events/product"),
        ("Charlotte Ledger", 8900, 162000, "membership"),
        ("Racket (Mpls)", 11000, 322000, "membership"),
        ("Side Dish (CO Spgs)", 2000, 72000, "sponsor"),
        ("Naptown Scoop", 20000, 250000, "ads"),
        ("Hell Gate (NYC)", 12000, 840000, "membership"),
    ]
    for i, (name, subs, rev, model) in enumerate(data):
        c = "#"+ORANGE if model in ("reader/events/product", "membership") else "#"+SLATE
        ax.scatter(subs, rev, s=140, color=c, zorder=3, edgecolor="white", linewidth=1.2)
        ax.annotate(name, (subs, rev), textcoords="offset points",
                    xytext=(8, 6), fontsize=9, color="#"+NAVY)
    ax.scatter([], [], color="#"+ORANGE, label="Reader-funded / membership")
    ax.scatter([], [], color="#"+SLATE, label="Ad / sponsor funded")
    ax.set_xlabel("Subscribers (free + paid)")
    ax.set_ylabel("Estimated annual revenue")
    ax.yaxis.set_major_formatter(FuncFormatter(usd))
    ax.xaxis.set_major_formatter(FuncFormatter(lambda v, p: f"{v/1000:.0f}K"))
    ax.set_title("Comparable Local/Independent Newsletters\nSmall lists can produce real revenue")
    ax.legend(frameon=False, loc="upper left")
    fig.tight_layout()
    fig.savefig(f"{ASSETS}/fig_comparables.png", bbox_inches="tight")
    plt.close(fig)

# ----------------------------------------------------------------------------
# CHART 6 — Events as % of revenue at major media co's (validates pillar)
# ----------------------------------------------------------------------------
def chart_events_share():
    fig, ax = plt.subplots(figsize=(7.2, 4.3))
    names = ["TIME\n(2026E)", "Semafor\n(2025)", "Atlantic", "Forbes", "FT Live"]
    share = [50, 50, 25, 23, 7]
    bars = ax.bar(names, share, color=["#"+NAVY, "#"+TEAL, "#"+ORANGE, "#"+SAND, "#"+SLATE])
    for b, s in zip(bars, share):
        ax.text(b.get_x()+b.get_width()/2, s+1, f"{s}%", ha="center",
                fontweight="bold", color="#"+NAVY)
    ax.set_ylabel("Events as % of revenue")
    ax.set_ylim(0, 60)
    ax.set_title("Events Are a Proven Media Revenue Pillar")
    fig.tight_layout()
    fig.savefig(f"{ASSETS}/fig_events_share.png", bbox_inches="tight")
    plt.close(fig)

# ----------------------------------------------------------------------------
# CHART 7 — P&L: revenue vs opex vs owner draw + break-even path
# ----------------------------------------------------------------------------
def chart_pnl():
    fig, ax = plt.subplots(figsize=(7.2, 4.3))
    x = np.arange(len(years)); w = 0.38
    ax.bar(x-w/2, total_rev, w, color="#"+TEAL, label="Total revenue")
    ax.bar(x+w/2, opex, w, color="#"+SLATE, label="Operating expenses")
    ax.plot(x, owner_draw, color="#"+ORANGE, marker="o", lw=2.5,
            label="Owner's draw / profit")
    for xi, v in zip(x, owner_draw):
        ax.text(xi, v+7000, usd(v), ha="center", color="#"+ORANGE, fontweight="bold")
    ax.axhline(60000, ls="--", color="#"+NAVY, lw=1.2)
    ax.text(2.35, 62000, "Full-time\nincome line", color="#"+NAVY, fontsize=8, va="bottom", ha="right")
    ax.set_xticks(x); ax.set_xticklabels(years)
    ax.yaxis.set_major_formatter(FuncFormatter(usd))
    ax.set_ylim(0, max(total_rev)*1.15)
    ax.set_title("Path to a Full-Time Income")
    ax.legend(frameon=False, loc="upper left")
    fig.tight_layout()
    fig.savefig(f"{ASSETS}/fig_pnl.png", bbox_inches="tight")
    plt.close(fig)

# ----------------------------------------------------------------------------
# CHART 8 — Local news decline (opportunity context)
# ----------------------------------------------------------------------------
def chart_newsdecline():
    fig, ax = plt.subplots(figsize=(7.2, 3.8))
    cats = ["Papers lost\nsince 2005", "Counties w/ 1\nor 0 outlets", "Net new digital\nsites (1 yr)"]
    vals = [3300, 1771, 81]
    bars = ax.bar(cats, vals, color=["#"+ORANGE, "#"+SLATE, "#"+TEAL])
    for b, v in zip(bars, vals):
        ax.text(b.get_x()+b.get_width()/2, v*1.02, f"{v:,}", ha="center",
                fontweight="bold", color="#"+NAVY)
    ax.set_yscale("log")
    ax.set_title("The Local-News Vacuum (Medill, 2024) — and the Digital Opening")
    ax.set_ylabel("Count (log scale)")
    fig.tight_layout()
    fig.savefig(f"{ASSETS}/fig_newsdecline.png", bbox_inches="tight")
    plt.close(fig)

def chart_engine():
    fig, ax = plt.subplots(figsize=(7.2, 4.3))
    cats = ["Y2\nRevenue", "Y3\nRevenue", "Y2 Owner's\ndraw", "Y3 Owner's\ndraw"]
    base = [total_rev[1], total_rev[2], owner_draw[1], owner_draw[2]]
    eng  = [eng_total[1], eng_total[2], eng_draw[1], eng_draw[2]]
    x = np.arange(len(cats)); w = 0.38
    ax.bar(x-w/2, base, w, color="#"+SLATE, label="Base plan (conservative)")
    ax.bar(x+w/2, eng, w, color="#"+TEAL, label="Discovery-engine adjusted")
    for xi, (b, e) in enumerate(zip(base, eng)):
        ax.text(xi-w/2, b+4000, usd(b), ha="center", fontsize=7.5, color="#"+NAVY)
        ax.text(xi+w/2, e+4000, usd(e), ha="center", fontsize=7.5, color="#"+NAVY, fontweight="bold")
    ax.set_xticks(x); ax.set_xticklabels(cats)
    ax.yaxis.set_major_formatter(FuncFormatter(usd)); ax.set_ylim(0, max(eng)*1.18)
    ax.set_title("The Engine's Effect: Lower Cost, Stronger Membership")
    ax.legend(frameon=False, loc="upper left")
    fig.tight_layout(); fig.savefig(f"{ASSETS}/fig_engine.png", bbox_inches="tight"); plt.close(fig)

for fn in (chart_funnel, chart_revenue_stack, chart_mix, chart_audience,
           chart_comparables, chart_events_share, chart_pnl, chart_newsdecline,
           chart_engine):
    fn()
print("charts done")

# ============================================================================
# DOCUMENT
# ============================================================================
doc = Document()

# base styles
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12

def set_run_color(run, hexc):
    run.font.color.rgb = RGBColor.from_string(hexc)

def shade_cell(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexc)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=None, size=10, align="left", white=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if white:
        run.font.color.rgb = RGBColor.from_string("FFFFFF")
    elif color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(17); set_run_color(run, NAVY)
        # bottom border
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "2"); bottom.set(qn("w:color"), TEAL)
        pbdr.append(bottom); pPr.append(pbdr)
    elif level == 2:
        run.font.size = Pt(13.5); set_run_color(run, TEAL)
    else:
        run.font.size = Pt(11.5); set_run_color(run, SLATE)
    return p

def add_para(text, italic=False, color=None, size=11, align="left", space_after=6):
    """Supports **bold** inline markers."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT, "just": WD_ALIGN_PARAGRAPH.JUSTIFY}[align]
    for i, chunk in enumerate(text.split("**")):
        if not chunk:
            continue
        run = p.add_run(chunk)
        run.font.size = Pt(size)
        run.font.italic = italic
        if i % 2 == 1:
            run.font.bold = True
        if color:
            set_run_color(run, color)
    return p

def add_bullets(items, style="List Bullet"):
    for it in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(3)
        for i, chunk in enumerate(it.split("**")):
            if not chunk:
                continue
            run = p.add_run(chunk)
            run.font.size = Pt(10.5)
            if i % 2 == 1:
                run.font.bold = True

def add_table(headers, rows, widths=None, header_fill=NAVY, fontsize=9.5, caption=None,
              first_col_bold=False):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hdr[i], header_fill)
        set_cell_text(hdr[i], h, bold=True, white=True, size=fontsize,
                      align="left" if i == 0 else "center")
    for r, row in enumerate(rows):
        cells = table.add_row().cells
        fill = LIGHT if r % 2 == 0 else "FFFFFF"
        for i, val in enumerate(row):
            shade_cell(cells[i], fill)
            set_cell_text(cells[i], str(val),
                          bold=(first_col_bold and i == 0),
                          size=fontsize,
                          align="left" if i == 0 else "center",
                          color=NAVY if (first_col_bold and i == 0) else None)
    if widths:
        for i, w in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    if caption:
        cap = add_para(caption, italic=True, color=GREY, size=8.5, space_after=10)
    return table

def add_figure(path, caption, width=6.3):
    doc.add_picture(path, width=Inches(width))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = add_para(caption, italic=True, color=GREY, size=8.5, align="center", space_after=12)

def hr():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "CCCCCC")
    pbdr.append(bottom); pPr.append(pbdr)

# ----- page margins -----
for s in doc.sections:
    s.top_margin = Inches(0.9); s.bottom_margin = Inches(0.9)
    s.left_margin = Inches(0.9); s.right_margin = Inches(0.9)

# ============================ TITLE PAGE ===================================
for _ in range(3):
    doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("BRIEF SCOUT MEDIA"); r.font.size = Pt(34); r.font.bold = True; set_run_color(r, NAVY)
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Founder Strategy & Business Plan"); r.font.size = Pt(17); set_run_color(r, TEAL)
doc.add_paragraph()
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Building a reader-funded local media & events company in Chattanooga, Tennessee")
r.font.size = Pt(12.5); r.font.italic = True; set_run_color(r, SLATE)
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Flagship product: The Scenic City Scout — a weekly digital newsletter")
r.font.size = Pt(12.5); r.font.italic = True; set_run_color(r, SLATE)
for _ in range(6):
    doc.add_paragraph()
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("No advertising. No sponsorship. Funded by members, products, and events.")
r.font.size = Pt(12); r.font.bold = True; set_run_color(r, ORANGE)
for _ in range(4):
    doc.add_paragraph()
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Confidential Working Document  ·  June 2026"); r.font.size = Pt(10.5); set_run_color(r, GREY)

doc.add_page_break()

# ============================ EXEC SUMMARY =================================
add_heading("1. Executive Summary", 1)
add_para(
    "**Brief Scout Media** is a new local media and events company forming in Chattanooga, "
    "Tennessee. Its first product is **The Scenic City Scout**, a weekly digital newsletter that "
    "tells residents what matters and what to do in Chattanooga. The defining strategic choice is "
    "what the Scout will **not** do: it will not sell advertising or sponsorship inside the "
    "newsletter. Instead, Brief Scout Media monetizes its audience directly through three owned "
    "revenue pillars — **paid membership, product and commerce sales, and live events** (both "
    "proprietary and joint-venture).", align="just")
add_para(
    "This is a deliberate bet against the dominant local-newsletter playbook. The incumbent in "
    "Chattanooga, NOOGAtoday (a 6AM City property), is free and advertising-funded. The legacy "
    "daily, the Chattanooga Times Free Press, is a paywalled, mostly-digital newspaper that has "
    "publicly acknowledged operating at a loss and raised its price to roughly $39/month. Between a "
    "free ad-product and an expensive legacy paywall sits an **open lane**: a premium, ad-light, "
    "reader-funded local brand with a genuine events franchise. No Axios Local edition and no "
    "established paid independent newsletter currently occupies it.", align="just")
add_para("**Why this can work — the evidence:**")
add_bullets([
    "**Email-first local media demand is proven here.** NOOGAtoday reports a large local free list "
    "(26,500+ as far back as 2020) and the Times Free Press lists 60,000+ newsletter recipients — "
    "tens of thousands of Chattanoogans already read local email daily.",
    "**Small lists can produce real income without ads.** Catskill Crew, a hyperlocal newsletter in "
    "rural New York, earns roughly $52,000–$57,000 per month at under 20,000 subscribers, at ~85% "
    "margin, with zero advertising, across 11 revenue streams (events, merch, a discount card, even "
    "a region-themed board game). Charlotte Ledger reached ~$150,000–$175,000/yr on a membership "
    "model; Racket (Minneapolis) reported $322,000.",
    "**Events are a validated media revenue pillar.** Events are on pace to be ~50% of TIME's "
    "revenue and ~50% of Semafor's; the share of publishers earning meaningful event revenue jumped "
    "from 47% to 70% in a single year.",
    "**The macro tailwind is strong.** The U.S. has lost ~3,300 newspapers since 2005 and ~55 "
    "million Americans now have limited local news — while reader-funded digital outlets are the "
    "fastest-growing bright spot.",
    "**The market is growing and affluent enough.** The Chattanooga metro (~595,000 people) is "
    "growing at more than twice the national rate, fueled by higher-income transplants from "
    "California and Illinois — a steady stream of newcomers who need exactly what the Scout offers.",
])
add_para(
    "**The financial arc.** Modeled conservatively with no advertising, Brief Scout Media reaches "
    "roughly **$31K in Year 1** (validation / side-income), **~$116K in Year 2** (approaching a "
    "full-time income), and **~$280K in Year 3** (a full-time founder salary plus a small "
    "contractor team and reinvestment). By Year 3 revenue is diversified across membership (~43%), "
    "events (~38%), and products (~19%) — the multi-stream structure that research repeatedly ties "
    "to local-media sustainability. Startup capital required is modest: roughly **$5,000**. These are "
    "deliberately conservative base-case figures; the proprietary **discovery engine** (Section 4.5) "
    "improves the Year-2/3 outlook by cutting content cost and lifting membership.",
    align="just")
add_figure(f"{ASSETS}/fig_revenue_stack.png",
           "Figure 1. Projected three-year revenue, built entirely from members, products, and events — no sponsorship.")

# ============================ COMPANY OVERVIEW =============================
add_heading("2. Company Overview", 1)
add_heading("2.1 Mission & vision", 2)
add_para(
    "**Mission:** Help people love living in Chattanooga by being the most trusted, most useful, and "
    "most loved guide to the city — independent of advertisers.")
add_para(
    "**Vision:** Build Brief Scout Media into a durable, locally-owned media company whose audience "
    "relationship is so strong that readers happily pay for it — and whose events become fixtures on "
    "the city's calendar. The Scenic City Scout is product one; the company is designed so the same "
    "reader-funded playbook can extend to new verticals (e.g., a dedicated food/dining edition, a "
    "newcomer's guide, a kids/family edition) and, eventually, new Southeastern cities.")
add_heading("2.2 The flagship product — The Scenic City Scout", 2)
add_para(
    "The Scenic City Scout is a **weekly** digital newsletter delivering a tightly curated brief on "
    "Chattanooga: what happened, what's changing, and — critically — **what to do this week**. "
    "\"Scenic City\" is Chattanooga's well-known civic nickname; the \"Scout\" voice is a trusted "
    "local insider who does the legwork so readers don't have to. The weekly cadence is a "
    "deliberate, sustainable choice: research shows over-sending accelerates churn and that "
    "daily-publishing burnout is a leading cause of solo-operator failure. A high-quality weekly "
    "protects editorial standards and founder stamina while a daily competitor (NOOGAtoday) saturates "
    "the inbox.")
add_heading("2.3 The core thesis — why no sponsorship", 2)
add_para(
    "Most local newsletters are free and sell ads. That model works, but it has structural costs: it "
    "requires large scale to matter (6AM City targets ~50,000 subscribers per market just to be "
    "profitable), it sells the reader's attention to third parties, and it creates an unavoidable "
    "tension between what's good for the reader and what's good for the advertiser. Brief Scout Media "
    "rejects that trade. By taking sponsorship off the table entirely, the Scout makes the reader — "
    "not an advertiser — the customer. Every incentive points the same direction: make something "
    "people love enough to pay for, sell them products they actually want, and gather them at events "
    "they're glad they attended.", align="just")
add_para("This is not idealism without precedent — it is the model behind some of the most resilient "
         "independent media of the last five years (Defector: 87% reader-funded, ~$4.6M/yr; 404 "
         "Media: profitable within ~6 months; Hell Gate, Racket, Aftermath; The Colorado Sun). "
         "Brief Scout Media applies that reader-funded ethic to a single mid-sized city.")

bsm_facts = [
    ["Company", "Brief Scout Media (LLC to be formed)"],
    ["Headquarters", "Chattanooga, Tennessee (Hamilton County)"],
    ["First product", "The Scenic City Scout — weekly digital newsletter"],
    ["Stage", "Pre-launch concept"],
    ["Revenue model", "Membership + products/commerce + events (no ads/sponsorship)"],
    ["Editorial cadence", "Weekly (free tier) + members-only extras"],
    ["Startup capital", "~$5,000"],
    ["3-yr revenue target", "~$280,000 with diversified, reader-funded streams"],
]
add_table(["Item", "Detail"], bsm_facts, widths=[1.7, 4.6], first_col_bold=True,
          caption="Table 1. Brief Scout Media at a glance.")

# ============================ MARKET OPPORTUNITY ==========================
add_heading("3. The Market Opportunity", 1)
add_heading("3.1 A local-news vacuum and a newsletter renaissance", 2)
add_para(
    "The collapse of legacy local news is well documented. Northwestern's Medill \"State of Local "
    "News 2024\" report counts **~3,300 newspapers lost since 2005** (more than a third of all U.S. "
    "papers), **127 closures in the prior 12 months** (~2.5 per week), **1,563 counties with only "
    "one** remaining news source plus 208 with none, and roughly **55 million Americans** with "
    "limited or no local news. The bright spot: a net **+81 standalone digital news sites** in a "
    "single year — the fastest-growing replacement for shuttered print, concentrated in metros like "
    "Chattanooga.", align="just")
add_figure(f"{ASSETS}/fig_newsdecline.png",
           "Figure 2. Legacy local news is collapsing while digital-native outlets grow — the structural opening Brief Scout Media is built for. Source: Medill State of Local News 2024.")
add_para(
    "Email is the medium of choice for this rebuild. Newsletters are owned audience — a direct, "
    "portable relationship not subject to a social platform's algorithm — which is precisely why "
    "operators from Morning Brew to Axios to thousands of indies have made the inbox their home. "
    "Typical newsletter open rates run 20–40% (far above most digital media), and the email list is "
    "consistently cited as the single most defensible asset a small media company can own.")

add_heading("3.2 The reader-revenue shift", 2)
add_para(
    "Across the industry, **reader revenue is now the top investment priority** for publishers, and "
    "the most durable independent outlets are funded primarily by their audiences rather than "
    "advertisers. The frameworks are mature: the Lenfest Institute and American Press Institute "
    "stress diversified reader-centric revenue and that \"retention is the new acquisition\"; the "
    "Membership Puzzle Project distinguishes transactional **subscription** from mission-driven "
    "**membership**, where people \"join the cause.\" Brief Scout Media is explicitly a membership "
    "business, not merely a paywall.")

add_heading("3.3 The Chattanooga market", 2)
add_para(
    "Chattanooga is an unusually good fit for a reader-funded local brand: large enough to support "
    "real revenue, small enough to dominate, growing fast, increasingly affluent, and digitally "
    "wired.")
mkt = [
    ["City of Chattanooga population", "~194,000 (2025 est.)", "U.S. Census / Times Free Press"],
    ["Chattanooga MSA (6-county TN-GA)", "~595,000 (2025)", "U.S. Census / UTC CRER"],
    ["MSA households (est.)", "~235,000", "Derived from Census"],
    ["Metro growth, 2020–2025", "+31,000 people; >2× U.S. rate", "UTC CRER (2026)"],
    ["Top in-migration origins", "California (#1), Illinois (#2)", "UTC CRER (2026)"],
    ["Median household income (county)", "$76,183", "ACS 2024 / FRED"],
    ["Median household income (city)", "~$64,500", "ACS 2024"],
    ["Median age (city)", "~36", "U.S. Census"],
    ["Tourism — annual visitors", "~11.1 million (2024)", "Chattanooga Tourism Co."],
    ["Tourism — visitor spending", "$1.806 billion (2024, record)", "Chattanooga Tourism Co."],
    ["Local events serviced / attendees", "253 events / 299,000+ (2024)", "Chattanooga Tourism Co."],
    ["EPB \"Gig City\" fiber customers", "~122,000 (~70% of homes)", "EPB 2024 Annual Report"],
    ["Unemployment (MSA)", "~3.4% (Dec 2025)", "BLS / FRED"],
    ["State income tax", "None (Tennessee)", "State of Tennessee"],
]
add_table(["Indicator", "Value", "Source"], mkt, widths=[2.7, 2.2, 2.0], first_col_bold=True,
          caption="Table 2. Chattanooga market snapshot.")
add_para(
    "Three features stand out for this business. First, **growth and migration**: the metro is "
    "adding higher-income transplants faster than almost any peer, and newcomers are the ideal "
    "audience for a \"make sense of your new city\" product. Second, a **deep events and tourism "
    "economy**: $1.8B in visitor spending and 253 serviced events a year mean a thick calendar to "
    "curate and a population primed to attend. Third, **\"Gig City\" digital readiness**: as the "
    "first U.S. city with citywide gigabit fiber (~70% household penetration), Chattanooga has an "
    "unusually connected, online-native population.")
add_figure(f"{ASSETS}/fig_funnel.png",
           "Figure 3. Top-down market sizing. The realistic three-year target captures a small slice of the engaged, able-to-pay audience.")

add_heading("3.4 Market sizing — TAM / SAM / SOM", 2)
add_bullets([
    "**TAM — ~470,000 MSA adults.** The total universe of potential readers across the six-county "
    "metro.",
    "**SAM — ~55,000 engaged, able-to-pay residents.** Benchmarked against demonstrated local "
    "email engagement (NOOGAtoday's local list plus the Times Free Press's 60,000+ newsletter "
    "recipients) and an income filter (the ~80,000–105,000 MSA households above ~$75K). This is the "
    "realistically reachable, willing-to-engage audience.",
    "**SOM — ~28,000 free subscribers and ~1,700 paying members within three years.** A small, "
    "credible share of the SAM, consistent with comparable operators. Paid conversion is modeled at "
    "the realistic 3–6% of an engaged free list (median ~3%; up to 10% only with a strong funnel).",
])

add_heading("3.5 Competitive landscape", 2)
comp = [
    ["NOOGAtoday (6AM City)", "Free daily newsletter", "Advertising / sponsorship",
     "Large free local list; daily habit", "No paid tier; ad-driven; inbox saturation"],
    ["Times Free Press", "Legacy daily, mostly digital", "Hard paywall ~$39/mo",
     "Newsroom, brand, archives", "Expensive; \"operating at a loss\"; legacy feel"],
    ["The Chattanooga Pulse", "Alt-weekly (arts/events)", "Was print ads",
     "Events/arts heritage", "Print suspended since 2020; diminished"],
    ["Chattanoogan.com", "Free web news daily", "Display ads",
     "Breaking/community news, free", "Utilitarian; not curated; no events"],
    ["ChattaMatters", "Civic explainer", "Grant / city-funded",
     "Civic depth", "Narrow scope; not consumer/events"],
    ["Brief Scout Media", "Weekly curated brief + events", "Members + products + events",
     "Reader-aligned; premium; events franchise", "New entrant; must earn trust"],
]
add_table(["Player", "What it is", "How it's funded", "Strengths", "Gaps / vulnerabilities"],
          comp, widths=[1.3, 1.3, 1.2, 1.3, 1.5], fontsize=8.5, first_col_bold=True,
          caption="Table 3. Competitive landscape. The reader-funded, events-anchored, premium lane is uncontested.")
add_para(
    "**The whitespace is clear.** The strongest newsletter (NOOGAtoday) is free and ad-funded; the "
    "only paid product (Times Free Press) is an expensive legacy paywall; the events/arts vehicle "
    "(The Pulse) effectively vacated the field in 2020. There is no Axios Local edition and no "
    "established paid independent Chattanooga newsletter. Brief Scout Media can be the **premium, "
    "reader-funded, events-driven** alternative — a position incumbents are structurally unable to "
    "copy without cannibalizing their own model.")

# ============================ BUSINESS MODEL ==============================
add_heading("4. The Business Model", 1)
add_para(
    "Brief Scout Media runs on three reinforcing revenue pillars, all owned and all reader-aligned. "
    "The newsletter is the **engine** that builds the audience and trust; membership, products, and "
    "events are the **monetization**. Research is unambiguous that diversification is what separates "
    "sustainable local media from fragile single-stream operators — outlets with **3+ revenue "
    "streams** are markedly more likely to be stable, and a formal business plan alone correlated "
    "with a 137% revenue increase in one LION Publishers study.", align="just")

add_heading("4.1 Pillar 1 — Membership", 2)
add_para(
    "A paid membership (target **$8/month or $80/year**, within the proven $5–$15/mo · $50–$120/yr "
    "band for local independent media) unlocks members-only content, early event access and member "
    "pricing, the digital archive, and the sense of belonging to the cause. Modeled conversion is a "
    "conservative 3–6% of the free list; at $80–$100/yr with ~3–4% monthly churn (typical for "
    "engaged local lists), gross member lifetime value lands around **$200–$275**.")

add_heading("4.2 Pillar 2 — Products & commerce", 2)
add_para(
    "High-margin owned products turn brand affection into revenue without touching advertising: "
    "Chattanooga-pride **merchandise** (print-on-demand to test, bulk once proven — ~$17/unit vs "
    "~$42 on-demand), **digital city guides** (newcomer guide, best-of guides — near-100% margin "
    "after creation), a **\"Scout Card\" local discount program** (recurring, partner-funded value "
    "for members), a **job board / classifieds** line ($200+/listing), and signature specialty "
    "items. Catskill Crew's region-themed board game (sold out in 17 days) is the proof that a "
    "beloved local brand can sell almost anything to its people.")

add_heading("4.3 Pillar 3 — Events", 2)
add_para(
    "Events are the highest-leverage pillar and the strongest moat. Two structures:")
add_bullets([
    "**Proprietary events** — ticketed gatherings the Scout owns end-to-end: member mixers, a "
    "\"new to Chattanooga\" newcomer night, themed tasting walks, and eventually an annual signature "
    "festival. Local consumer events cluster at a **$30–$50 ticket** sweet spot, with **$40–$75 of "
    "additional per-attendee on-site spend**.",
    "**Joint-venture events** — co-produced with breweries, restaurants, venues, and festivals on a "
    "**revenue-share** basis. The Scout brings the audience and promotion; the partner brings the "
    "space and operations; risk and cost are shared. This is the capital-light way to run frequent "
    "events without owning infrastructure.",
])
add_para(
    "The pillar is well-validated at scale — events are tracking toward ~50% of revenue at TIME and "
    "Semafor, and the share of publishers earning meaningful event income leapt from 47% to 70% in "
    "one year — and Chattanooga's 253-event, $1.8B visitor economy is fertile ground.")
add_figure(f"{ASSETS}/fig_events_share.png",
           "Figure 4. Events have become a core revenue pillar for media companies of every size. Sources: company disclosures via Adweek, Digiday, A Media Operator (2024–2026).")

add_heading("4.4 Proof points — comparable operators", 2)
add_para(
    "Brief Scout Media's model is assembled from tactics already working at named operators. The "
    "key insight: **a small, deeply engaged list can out-earn a large, shallow one** when monetized "
    "through membership, products, and events rather than ad impressions.")
cmp_rows = [
    ["Catskill Crew", "Catskills, NY", "<20,000", "~$620K/yr (~$52–57K/mo)",
     "Reader/events/products (no ads), ~85% margin, 11 streams"],
    ["Charlotte Ledger", "Charlotte, NC", "2,000 paid / 6,900 free", "~$150–175K/yr",
     "Membership ($9/mo, $99/yr)"],
    ["Racket", "Minneapolis, MN", "~3–4K paid", "~$322K/yr",
     "Worker-owned, reader-funded"],
    ["Hell Gate", "New York, NY", "~9,000 paid", "~$70K/mo MRR",
     "Membership co-op (~⅔ from subs)"],
    ["Defector", "National (sports)", "~42,000", "~$4.6M/yr",
     "87% reader-funded"],
    ["Naptown Scoop", "Annapolis, MD", "~18–21K", "~$200–300K/yr",
     "Advertising ($70 CPM) — contrast model"],
    ["NOOGAtoday", "Chattanooga, TN", "~26.5K+ (2020)", "n/d",
     "Free, advertising — local incumbent"],
]
add_table(["Operator", "Market", "Subscribers", "Est. revenue", "Model / notes"],
          cmp_rows, widths=[1.2, 1.2, 1.2, 1.2, 1.8], fontsize=8.5, first_col_bold=True,
          caption="Table 4. Comparable local & independent newsletters. Sources: beehiiv, Nieman Lab, Inbox Collective, Side Hustle Nation, company reports.")
add_figure(f"{ASSETS}/fig_comparables.png",
           "Figure 5. Revenue vs. list size for comparable operators. Reader-funded models (orange) extract far more value per subscriber than ad models.")

add_heading("4.5 The Discovery Engine — the asset that powers it all", 2)
add_para(
    "Brief Scout Media is not merely a newsletter that scouts; it is a **local discovery engine** "
    "that happens to publish a newsletter. The company systematically mines public signals — the "
    "county's weekly new-business licenses, a satellite-town beat from LaFayette to Cleveland to the "
    "Sequatchie Valley, the maker markets and the food incubator — and uses AI to find and "
    "pre-research candidates, with a human always verifying, visiting, building the relationship, and "
    "writing the story. The newsletter is the most visible **output** of the engine, not the whole "
    "business.", align="just")
add_para("**The revenue model does not change** — Brief Scout Media stays reader-funded (membership "
         "+ products + events, no advertising). But the engine upgrades three things:")
add_bullets([
    "**A deeper moat.** The promise becomes structural, not rhetorical: *we find the people building "
    "this city before anyone else, and we tell the truth about them.* Proprietary discovery is a "
    "fourth moat layer on top of trust, taste, and community — and the no-ads stance keeps it credible.",
    "**A lower content cost.** The expensive part of local media — finding and vetting what's worth "
    "covering — is largely automated, so a solo founder or tiny team can sustain newsroom-depth "
    "coverage. This makes the Year-1 'default-alive' economics safer and makes new verticals and new "
    "markets far cheaper to launch: the engine is the reusable asset carried into a food edition or a "
    "second city.",
    "**A self-feeding flywheel.** Events, products, and partner leads fall *out* of the pipeline "
    "rather than being separate lifts — Meet-the-Maker nights, a food-incubator showcase, maker "
    "guides, and a warm list of every new business the week it opens (celebrated editorially for "
    "free; the business-partner program kept walled and separate).",
])
add_para(
    "**A new membership perk.** The engine creates value that didn't exist before and that members "
    "will pay for: **first word on new openings** — the Scout's 'Just Licensed / Coming Soon' intel, "
    "delivered to members before the city knows. A concrete, unique reason to upgrade that lifts both "
    "conversion and retention.")
add_para(
    "**The one rule.** The engine *sources*; humans *own the trust*. AI finds and pre-vets; a real "
    "person verifies, visits, and writes in the Scout's voice. The moment it reads like an automated "
    "directory, the moat collapses — so the engine stays an internal weapon, never the product.")
add_para(
    "**Engine-adjusted outlook (Years 2–3).** Holding Year 1 as the conservative base case, the "
    "engine's two clearest effects — lower content/sourcing cost and the membership intel perk — "
    "improve the later years (with events and products also easier to source). The figures below are "
    "illustrative upside, not a replacement for the conservative base model in Section 9; notably, the "
    f"engine-adjusted Year-3 total (~{usd(eng_total[2])}) lands close to the focused-model estimate in "
    "Addendum 2.", align="just")
eng_rows = [
    ["Membership", usd(membership[1]), usd(eng_membership[1]), usd(membership[2]), usd(eng_membership[2])],
    ["Events", usd(events[1]), usd(eng_events[1]), usd(events[2]), usd(eng_events[2])],
    ["Products & commerce", usd(products[1]), usd(eng_products[1]), usd(products[2]), usd(eng_products[2])],
    ["Total revenue", usd(total_rev[1]), usd(eng_total[1]), usd(total_rev[2]), usd(eng_total[2])],
    ["Operating expenses", usd(opex[1]), usd(eng_opex[1]), usd(opex[2]), usd(eng_opex[2])],
    ["Owner's draw / profit", usd(owner_draw[1]), usd(eng_draw[1]), usd(owner_draw[2]), usd(eng_draw[2])],
]
add_table(["Line item", "Year 2 — base", "Year 2 — engine", "Year 3 — base", "Year 3 — engine"],
          eng_rows, widths=[2.2, 1.3, 1.3, 1.3, 1.3], fontsize=9.5, first_col_bold=True,
          caption="Table 4b. Discovery-engine-adjusted outlook (illustrative). Year 1 unchanged; "
                  "Section 9 remains the conservative base case.")
add_figure(f"{ASSETS}/fig_engine.png",
           "Figure 5b. The engine's two clearest effects — lower content cost and a stronger membership intel perk — lift owner's draw in Years 2–3.")

# ============================ PRODUCT STRATEGY ============================
add_heading("5. Product & Editorial Strategy", 1)
add_heading("5.1 The weekly Scout", 2)
add_para(
    "Each issue is skimmable in five minutes and genuinely useful: a short read on the week's most "
    "important local development, a curated **\"what to do this week\"** events rundown, a local "
    "recommendation (a place, a person, a dish), and one practical civic or insider note. The voice "
    "is warm, opinionated, and unmistakably local. Quality and consistency are the product — the "
    "free issue is the funnel for everything else.")
add_heading("5.2 Membership tiers", 2)
tiers = [
    ["Free", "$0", "Weekly Scout newsletter; public event listings",
     "Top-of-funnel; builds the list and the habit"],
    ["Member", "$8/mo or $80/yr", "Members-only deep dives & archives; first word on new openings "
     "(the Scout's 'Just Licensed' intel); member event pricing; Scout Card discounts; community access",
     "Core reader-revenue tier"],
    ["Founding Member", "$150/yr", "All Member benefits + signature merch, name recognition, "
     "free annual signature-event ticket", "Higher ARPU; superfans who fund the mission"],
]
add_table(["Tier", "Price", "What's included", "Role"], tiers,
          widths=[1.1, 1.2, 2.6, 1.4], fontsize=9, first_col_bold=True,
          caption="Table 5. Membership tiers. Pricing sits within the proven local-media band and mirrors Charlotte Ledger / 404 Media / Defector.")
add_heading("5.3 The product ladder", 2)
add_para(
    "The business is designed as an ascending ladder of commitment: a reader discovers the free "
    "Scout → becomes a habitual reader → attends a free or low-cost event → buys a guide or merch → "
    "converts to membership → upgrades to founding member and attends signature events. Each rung "
    "deepens the relationship and raises lifetime value, and every rung is reader-aligned.")

# ============================ REVENUE DEEP DIVE ===========================
add_heading("6. Revenue Model — Deep Dive", 1)
add_heading("6.1 Membership economics", 2)
add_para(
    "The table below shows annual membership revenue at a mature free-list scale of ~26,000, across "
    "a realistic grid of conversion rates and prices (net of ~3% payment processing). The base case "
    "— ~5% conversion at $80–$100 — produces a six-figure membership line on its own.")
# membership sensitivity grid
free_base = 26000
conv = [0.03, 0.05, 0.07]
price = [80, 100, 120]
grid_rows = []
for c in conv:
    members = int(free_base * c)
    row = [f"{int(c*100)}%  ({members:,} members)"]
    for pr in price:
        rev = members * pr * 0.97
        row.append(f"${rev/1000:.0f}K")
    grid_rows.append(row)
add_table(["Conversion (on 26K list)", "$80/yr", "$100/yr", "$120/yr"], grid_rows,
          widths=[2.6, 1.2, 1.2, 1.2], fontsize=9.5, first_col_bold=True,
          caption="Table 6. Annual membership-revenue sensitivity (net of payment fees).")

add_heading("6.2 Event economics", 2)
add_para(
    "Events range from low-cost community gatherings to a signature annual festival. Illustrative "
    "unit economics for representative event types:")
ev = [
    ["Member mixer (JV bar/venue)", "60", "Free / $5", "$5–10", "$300–600", "Retention + funnel"],
    ["Tasting walk / sip & stroll", "150", "$40", "$15–20", "$3,500–4,500", "Ticket + partner share"],
    ["Newcomer night (quarterly)", "100", "$25", "$10", "$1,800–2,500", "Onboards transplants"],
    ["Signature festival (annual)", "1,500", "$35", "$50+", "$30,000–40,000", "Brand + sponsorship-free gate/commerce"],
]
add_table(["Event type", "Typical attend.", "Ticket", "On-site spend", "Est. net to Scout", "Strategic role"],
          ev, widths=[1.5, 1.0, 0.9, 1.0, 1.2, 1.4], fontsize=8.5, first_col_bold=True,
          caption="Table 7. Illustrative event unit economics. Ticket and on-site figures benchmarked to local consumer-event data (Eventbrite, Ticket Fairy, regional festivals).")
add_para(
    "Joint-venture structures keep events capital-light: the partner supplies the venue, alcohol "
    "license, and operations; the Scout supplies the audience and promotion; the two share gate "
    "and/or food-and-beverage revenue. Note that because the Scout takes no in-newsletter "
    "sponsorship, **event sponsorship is treated separately** — it is permissible to have a brewery "
    "co-host an event, since that is a partnership around an experience, not a sale of the reader's "
    "inbox attention. (If the founder prefers absolute purity, events can run on tickets and "
    "commerce alone.)")

add_heading("6.3 Products & commerce", 2)
add_para(
    "Merchandise begins print-on-demand (zero inventory risk) and shifts proven designs to bulk "
    "screen-print (~$17/unit vs ~$42 on-demand) once demand is validated. Digital guides carry "
    "near-100% margins after creation. The Scout Card discount program is a recurring, "
    "partner-funded member perk that costs the Scout little and deepens member value. A job "
    "board/classifieds line adds modest, passive revenue ($200+/30-day listing).")

# ============================ GTM / GROWTH ================================
add_heading("7. Go-to-Market & Growth Strategy", 1)
add_para(
    "**Critical constraint:** geo-restricted local newsletters generally cannot use the paid "
    "acquisition networks (beehiiv Boosts, SparkLoop) that national newsletters rely on, because "
    "those deliver out-of-market subscribers. Growth must therefore come from **organic, referral, "
    "and partnership** channels — which research confirms are both the most effective and the most "
    "trusted for hyperlocal (42% of operators rank word-of-mouth #1; referral programs add ~17% to "
    "growth).", align="just")
add_bullets([
    "**Referral engine.** A built-in referral program (milestone rewards: stickers → merch → free "
    "membership month) turns readers into recruiters — the top growth lever for comparable operators.",
    "**Local partnerships.** Co-promotion with breweries, coffee shops, gyms, neighborhood "
    "associations, the startup ecosystem (CO.LAB, the Edney/Innovation District), realtors, and "
    "relocation services — especially powerful given heavy in-migration.",
    "**Events as acquisition.** Every event is a list-building machine; attendees join on the spot.",
    "**Organic social & content.** Short-form local video and Instagram/Facebook (where ~60% of "
    "younger residents discover events) drive inbound signups; the Scout's POV makes it shareable.",
    "**Newcomer funnel.** A free \"New to Chattanooga\" guide as a signup magnet targets the "
    "fast-growing transplant population precisely when they're hungry for local guidance.",
])
add_para(
    "Target trajectory (organic-led): **~7,000 free subscribers by end of Year 1**, ~16,000 by "
    "Year 2, ~28,000 by Year 3 — well within demonstrated local benchmarks (Mirror Indy grew its "
    "list +608% in year one; Catskill Crew went 0→10,000+ in under a year).")
add_figure(f"{ASSETS}/fig_audience.png",
           "Figure 6. Modeled audience growth. The free list is the funnel; paying members are the conservative 3–6% who convert.")

# ============================ OPERATIONS =================================
add_heading("8. Operations & Technology", 1)
add_para(
    "The operating philosophy is lean: a single founder-operator at launch, low fixed costs, and a "
    "platform that does not tax subscription revenue. Contractors (a part-time editor, an events/ops "
    "helper) are added as revenue allows, primarily in Years 2–3.")
add_heading("8.1 Platform selection", 2)
add_para(
    "The recommended platform is **beehiiv** (or self-hosted **Ghost**): both charge **0% on "
    "subscription revenue**, unlike Substack's 10% cut. On a $10 subscription, beehiiv/Ghost net "
    "~$9.41 after Stripe vs. ~$8.41 on Substack — a ~12% revenue difference that compounds at scale.")
plat = [
    ["beehiiv", "Free to 2,500; Scale ~$43–49/mo", "0%", "Built-in paid subs, referral engine, web"],
    ["Ghost (Pro)", "Publisher ~$29/mo (annual)", "0%", "Own-your-stack, memberships, 0% fee"],
    ["Substack", "$0 platform fee", "10% + Stripe", "Easiest start; discovery network"],
    ["Kit (ConvertKit)", "Free to 10K; Creator ~$33/mo", "n/a (ESP)", "Automation-heavy; no native paywall"],
    ["Mailchimp", "From ~$13/mo", "n/a", "General ESP; weak for paid membership"],
]
add_table(["Platform", "Indicative price", "Cut of sub revenue", "Notes"], plat,
          widths=[1.3, 2.0, 1.3, 1.9], fontsize=9, first_col_bold=True,
          caption="Table 8. Platform comparison. beehiiv and Ghost avoid taxing subscription revenue. Prices are 2026 indicative; verify on vendor pages.")
add_heading("8.2 Cost structure & startup costs", 2)
startup = [
    ["LLC formation + basic legal", "$800"],
    ["Brand identity, logo & website setup", "$1,500"],
    ["Platform & tools (3-month runway)", "$200"],
    ["Initial merch / design + launch event", "$1,500"],
    ["Contingency", "$1,000"],
    ["Total startup capital", "~$5,000"],
]
add_table(["Pre-launch item", "Cost"], startup, widths=[3.8, 1.6], fontsize=9.5,
          first_col_bold=True,
          caption="Table 9. Startup capital requirement (lean, solo). Consistent with the $1,500–$5,000 range for solo local newsletters.")
add_para(
    "Ongoing tooling is remarkably cheap — roughly **$60–$130/month** early (ESP + Canva + domain) "
    "— so operating break-even on tools alone is just **~6 members or one small event**. The "
    "dominant \"cost\" of this business is the founder's time, which is why the financial plan "
    "treats owner's compensation as the key milestone rather than tooling break-even.")

# ============================ FINANCIAL PLAN =============================
add_heading("9. Financial Plan", 1)
add_heading("9.1 Key assumptions", 2)
add_bullets([
    "**No advertising or sponsorship revenue** in any year (the core thesis).",
    "Free list: 7,000 / 16,000 / 28,000 (end of Years 1–3), organic-led.",
    "Paying members: 350 / 850 / 1,700 (EoY) — a conservative 3–6% of the free list.",
    "Blended membership ARPU ~$90/yr net of payment fees; ~3–4% monthly churn.",
    "Events scale 3 → 7 → 12 per year, adding a signature festival in Year 3.",
    "Products scale with brand: merch, guides, Scout Card, specialty items.",
    "All event and product figures shown **net** of direct costs (COGS, venue/JV share).",
])
add_heading("9.2 Three-year revenue build", 2)
rev_rows = [
    ["Membership", usd(membership[0]), usd(membership[1]), usd(membership[2])],
    ["Events", usd(events[0]), usd(events[1]), usd(events[2])],
    ["Products & commerce", usd(products[0]), usd(products[1]), usd(products[2])],
    ["Total revenue", usd(total_rev[0]), usd(total_rev[1]), usd(total_rev[2])],
]
add_table(["Revenue stream", "Year 1", "Year 2", "Year 3"], rev_rows,
          widths=[2.4, 1.3, 1.3, 1.3], fontsize=10, first_col_bold=True,
          caption="Table 10. Revenue build by stream.")
add_figure(f"{ASSETS}/fig_mix.png",
           "Figure 7. Year-3 revenue mix — diversified across all three reader-funded pillars.")

add_heading("9.3 Profit & loss summary", 2)
pnl_rows = [
    ["Total revenue", usd(total_rev[0]), usd(total_rev[1]), usd(total_rev[2])],
    ["Operating expenses", usd(opex[0]), usd(opex[1]), usd(opex[2])],
    ["  Tools / platform", "$1,800", "$2,600", "$3,500"],
    ["  Marketing / growth", "$3,000", "$7,000", "$14,000"],
    ["  Legal / admin / accounting", "$2,200", "$3,500", "$5,000"],
    ["  Contract content / team", "$2,000", "$18,000", "$55,000"],
    ["  Misc / contingency", "$1,500", "$3,000", "$6,000"],
    ["Owner's draw / net profit", usd(owner_draw[0]), usd(owner_draw[1]), usd(owner_draw[2])],
]
add_table(["Line item", "Year 1", "Year 2", "Year 3"], pnl_rows,
          widths=[2.6, 1.2, 1.2, 1.2], fontsize=9.5, first_col_bold=True,
          caption="Table 11. Simplified P&L. Sub-lines of operating expenses are indented.")
add_para(
    "**Interpretation.** Year 1 is validation: a ~$20K owner's draw alongside a likely day job. "
    "Year 2 approaches a full-time income (~$82K) as membership compounds and events scale. Year 3 "
    "supports a full founder salary plus a part-time contractor team (already inside operating "
    "expenses) with profit to reinvest — roughly the ~18-month-to-profit arc 6AM City reports for "
    "new markets, achieved here without advertising. An **engine-adjusted view** (Section 4.5) raises "
    f"Year-2 owner's draw toward ~{usd(eng_draw[1])} and Year-3 toward ~{usd(eng_draw[2])} as content "
    "costs fall and the membership intel perk lifts conversion.")
add_figure(f"{ASSETS}/fig_pnl.png",
           "Figure 8. Revenue, operating expenses, and owner's draw. The founder crosses a full-time income line during Year 2.")

add_heading("9.4 Scenario analysis", 2)
scen = [
    ["Conservative", "20,000", "1,100", "$180K", "Slower list growth; events ramp delayed"],
    ["Base", "28,000", "1,700", "$280K", "Plan assumptions as modeled"],
    ["Optimistic", "38,000", "2,600", "$420K", "Strong referral flywheel; festival overperforms"],
]
add_table(["Scenario", "Y3 free list", "Y3 members", "Y3 revenue", "Drivers"], scen,
          widths=[1.2, 1.2, 1.1, 1.1, 1.8], fontsize=9, first_col_bold=True,
          caption="Table 12. Three-year (Year-3) scenarios. Even the conservative case clears a full-time income.")

# ============================ RISKS =====================================
add_heading("10. Risks & Mitigations", 1)
risks = [
    ["Slow monetization / long ramp", "High", "Local media takes ~18+ months to profit; capitalize "
     "for a long runway, keep costs near-zero, validate willingness-to-pay early with a founding-member presale."],
    ["Founder burnout", "High", "Weekly (not daily) cadence; templatized production; bring on a "
     "part-time editor in Year 2; protect a sustainable workload."],
    ["Unproven local willingness-to-pay", "Medium", "No paid independent newsletter exists here yet "
     "— run a paid founding-member presale before launch to prove demand; lean on events/products if conversion lags."],
    ["Incumbent response (NOOGAtoday)", "Medium", "Compete on quality, curation, community and "
     "events, not volume; the incumbent's ad model can't easily add a premium paid tier without conflict."],
    ["Platform / deliverability dependency", "Medium", "Own the email list; choose a portable "
     "platform (beehiiv/Ghost); monitor inbox placement; never rely on a single social channel."],
    ["Over-reliance on one stream", "Medium", "Three pillars by design; target no single stream "
     ">50% of revenue (mirrors Catskill Crew's 11-stream resilience)."],
    ["Affordability ceiling (city poverty ~19%)", "Low", "Free tier keeps the brand universally "
     "accessible; paid tiers target the affluent, growing transplant segment."],
]
add_table(["Risk", "Severity", "Mitigation"], risks, widths=[1.9, 0.9, 3.5], fontsize=9,
          first_col_bold=True, caption="Table 13. Risk register.")

# ============================ ROADMAP ===================================
add_heading("11. Roadmap & Milestones", 1)
road = [
    ["Pre-launch (Mo. 0–2)", "Form LLC; build brand & site; choose beehiiv; line up 5–10 launch "
     "partners; open founding-member presale to validate demand."],
    ["Launch (Mo. 3)", "Publish first weekly Scout; activate referral program; host launch event."],
    ["Mo. 3–12", "Grow to ~7,000 free / ~350 members; run 3 events; ship first guide + merch drop."],
    ["Year 2", "Scale to ~16,000 free / ~850 members; 7 events incl. quarterly ticketed; add "
     "part-time editor; launch Scout Card."],
    ["Year 3", "~28,000 free / ~1,700 members; 12 events incl. signature festival; full-time "
     "founder + contractor team; evaluate a second edition/vertical."],
]
add_table(["Phase", "Key milestones"], road, widths=[1.7, 4.6], fontsize=9.5, first_col_bold=True,
          caption="Table 14. Phased roadmap.")

# ============================ CONCLUSION ================================
add_heading("12. Conclusion", 1)
add_para(
    "Brief Scout Media is a focused bet on a simple, evidence-backed idea: in a city losing legacy "
    "news but gaining engaged, affluent newcomers, a beloved weekly newsletter can be funded "
    "**directly by the people who read it** — through membership, products, and events — rather than "
    "by advertisers. The Chattanooga market is large, growing, wired, and event-rich; the "
    "competitive lane for a premium reader-funded brand is wide open; comparable operators from "
    "Catskill Crew to Charlotte Ledger to Defector prove the model pays; and the capital required is "
    "small. The plan is conservative by design and still reaches a full-time founder income within "
    "roughly 18–24 months, on three diversified, mission-aligned revenue streams. The path is "
    "demanding — local media rewards patience and consistency — but the structure is sound and the "
    "timing is right.", align="just")

# ============================ SOURCES ==================================
doc.add_page_break()
add_heading("Appendix — Selected Sources", 1)
add_para("This plan synthesizes a multi-source research effort. Key references by theme:",
         color=SLATE, size=10)

def src_block(title, items):
    add_para(title, color=TEAL, size=11)
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(it); run.font.size = Pt(8.5); set_run_color(run, SLATE)

src_block("Local news & newsletter industry", [
    "Northwestern Medill, State of Local News 2024 — localnewsinitiative.northwestern.edu",
    "Press Gazette, 6AM City / Axios Local coverage (2024) — pressgazette.co.uk",
    "Newsletter Operator; HubSpot State of Newsletters 2025; ClickMinded Newsletter Statistics 2026",
    "beehiiv — Catskill Crew case study; Side Hustle Nation — Naptown Scoop ($200K) ",
])
src_block("Reader-revenue & membership economics", [
    "Lenfest Institute — Business Models for Local News; American Press Institute — subscriber retention",
    "Membership Puzzle Project — The Membership Guide; Piano — Subscription Performance Benchmarks 2024",
    "Substack conversion analyses (Simon Owens; yana-g-y); Nieman Lab — Defector, Hell Gate, 404 Media, Racket",
    "Editor & Publisher — Charlotte Ledger, Cityside/Berkeleyside; Colorado Sun",
])
src_block("Events as media revenue", [
    "Adweek / Digiday — TIME, Semafor, Bloomberg, Atlantic, Forbes events revenue (2024–2026)",
    "A Media Operator — Financial Times / FT Live 2024 results; Eventbrite TRNDS Report; Ticket Fairy festival economics",
])
src_block("Platforms, products & costs", [
    "beehiiv, Substack, Ghost, Kit, Mailchimp pricing (2026); Stripe / Wise fee guides",
    "TailorBrands / LLCUniversity — LLC formation costs; print-on-demand vs bulk merch economics",
])
src_block("Chattanooga market", [
    "U.S. Census Bureau QuickFacts; UTC Center for Regional Economic Research white papers (2026)",
    "Chattanooga Tourism Co. 2024 Annual Report; Tennessee Dept. of Tourist Development (2024)",
    "EPB 2024 Annual Report (Gig City fiber); Times Free Press; Chattanooga Chamber major-employers list",
    "Times Free Press, NOOGAtoday/6AM City, The Chattanooga Pulse, Chattanoogan.com, ChattaMatters",
])
add_para(
    "Note: several figures are operator anecdotes or industry estimates (flagged as such in the "
    "underlying research) rather than audited statistics; Chattanooga subscriber counts for "
    "competitors are self-reported. Verify load-bearing local figures (Times Free Press circulation, "
    "current competitor lists, Census household/income tables) directly before external use.",
    italic=True, color=GREY, size=8.5)

out = "/home/user/prscott/Brief_Scout_Media_Business_Plan.docx"
doc.save(out)
print("saved", out)
