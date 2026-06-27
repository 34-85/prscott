#!/usr/bin/env python3
"""Founder's manifesto — The Scenic City Insider (internal / investor working draft).
Rebranded from 'Scout'; folds in the Council-governs-the-Card framing and Year-1 milestone
numbers from the cash-flow model. Outputs .docx + .md."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY="1B3A4B"; TEAL="2A9D8F"; ORANGE="E76F51"; SLATE="415A6B"; GREY="6B7B85"; LIGHT="EEF3F4"

# ---- content (single source of truth; also written to .md) ----
TITLE="The Scenic City Insider"
SUB="A founder's note — internal / investor working draft"

SECTIONS=[
 ("Why I started it",[
  "I started The Scenic City Insider because I wanted to know Chattanooga more deeply — and to help other people do the same.",
  "Like a lot of us, I found myself living in a city I love while sensing there were still extraordinary people, places, and stories hidden in plain sight. A hundred-year-old café I had never walked into. A maker quietly building something remarkable three towns over. A local institution everyone assumes will always be there — until one day it isn't.",
  "The Insider is how I explore this city with intention, and share what's worth finding.",
 ]),
 ("The mission",[
  "Every Wednesday, the Insider publishes one carefully crafted issue built around discovery, curiosity, and trust. The mission is simple: help people know Chattanooga more fully and love it harder. That means less noise, less content for content's sake, and more thoughtful curation of the businesses, builders, restaurants, shops, and hidden corners that make this place what it is.",
 ]),
 ("Reader-funded by design",[
  "The Insider is reader-funded on purpose. No advertising, no sponsorship, no paid editorial placement. Coverage is never for sale. That independence isn't a side detail — it is the business, because trust is the business. If the Insider recommends something, it's because it genuinely earned the recommendation.",
 ]),
 ("The one place a business pays — and who guards it",[
  "There is exactly one place a local business can pay: the **Insider Card**, a member-discount program businesses opt into. Even that has a guardrail. The **Insider Council** — a voluntary member advisory committee — vets the Card and the wider model, so the people who fund the Insider are the same people who decide it stays true to the mission. A business can sponsor a discount or take a booth; it can never buy a word of coverage. The members hold the line, not the founder.",
 ]),
 ("What it becomes",[
  "Over time, the Insider grows into more than a weekly publication. It becomes a membership community built around shared discovery, real connection, and a belief that the best cities are made by the people who choose to engage with them. The goal isn't simply to tell you where to go. It's to help all of us feel more rooted in the place we call home.",
 ]),
 ("How I get there",[
  "Here is how I move from a free weekly to a member-funded community without ever crossing that line:",
 ]),
]

# (year label, heading, body)
ROADMAP=[
 ("Year 1","Audience, trust, and systems — no monetization.",
  "This is the year I drive subscribers, earn trust, and build the machine that delivers value every week — and, honestly, the year I'll make mistakes, ship a typo, and get something wrong now and then. The goalposts that make everything after it credible: roughly **7,000+ engaged subscribers** by year-end, a **45–55% open rate** (well above the industry norm), and a **founding-member waitlist** that proves people will pay before I ever ask them to."),
 ("Year 2","Membership and governance.",
  "I introduce **paid membership**, which grants access: meetups, members-only content, dinners, products, and the **Insider Card** discount program. I also stand up the **Insider Council** — the voluntary advisory committee that keeps the model honest and helps shape product and monetization while protecting the member-funded promise. Early targets: convert roughly **4–6% of the list to paid** (on the order of **350+ founding members**), hold **monthly churn under 4%**, and build toward **$3,000–$8,000 in recurring monthly revenue**."),
 ("Year 3+","Product and extension.",
  "With a trusted audience and a proven membership base, I expand what membership is worth: new products, guides, events, and research (**Insider Insights**) — and, once the playbook is documented and repeatable, the option to extend into new verticals or markets."),
]

# ---------- DOCX ----------
doc=Document(); sec=doc.sections[0]
for m in ("top_margin","bottom_margin"): setattr(sec,m,Inches(0.9))
for m in ("left_margin","right_margin"): setattr(sec,m,Inches(1.1))
nm=doc.styles["Normal"]; nm.font.name="Georgia"; nm.font.size=Pt(11.5); nm.paragraph_format.space_after=Pt(8); nm.paragraph_format.line_spacing=1.25
def sca(r,h): r.font.color.rgb=RGBColor.from_string(h)
def runs(p,text,size=11.5,italic=False,color=None,bold=False):
    for i,ch in enumerate(text.split("**")):
        if not ch: continue
        r=p.add_run(ch); r.font.size=Pt(size); r.font.italic=italic; r.font.bold=bold or (i%2==1)
        if color: sca(r,color)

# Title block
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run(TITLE); r.font.size=Pt(26); r.font.bold=True; sca(r,NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run(SUB); r.font.size=Pt(12); r.font.italic=True; sca(r,TEAL)
# rule
p=doc.add_paragraph(); pPr=p._p.get_or_add_pPr(); b=OxmlElement("w:pBdr"); bt=OxmlElement("w:bottom")
bt.set(qn("w:val"),"single"); bt.set(qn("w:sz"),"6"); bt.set(qn("w:space"),"6"); bt.set(qn("w:color"),ORANGE)
b.append(bt); pPr.append(b)
doc.add_paragraph()

for head,paras in SECTIONS:
    h=doc.add_paragraph(); h.paragraph_format.space_before=Pt(10); h.paragraph_format.space_after=Pt(3)
    r=h.add_run(head); r.font.bold=True; r.font.size=Pt(13); r.font.name="Georgia"; sca(r,ORANGE)
    for para in paras:
        pp=doc.add_paragraph(); runs(pp,para)

# Roadmap as styled blocks
for yr,head,body in ROADMAP:
    t=doc.add_table(rows=1,cols=2); t.allow_autofit=False
    t.columns[0].width=Inches(1.0); t.columns[1].width=Inches(5.3)
    c0=t.rows[0].cells[0]; c0.text=""
    pr=c0.paragraphs[0]; rr=pr.add_run(yr); rr.font.bold=True; rr.font.size=Pt(12); sca(rr,TEAL); rr.font.name="Georgia"
    # shade year cell
    tcPr=c0._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),LIGHT); tcPr.append(s)
    c1=t.rows[0].cells[1]; c1.text=""
    ph=c1.paragraphs[0]; rh=ph.add_run(head); rh.font.bold=True; rh.font.size=Pt(11.5); sca(rh,NAVY); rh.font.name="Georgia"
    pb=c1.add_paragraph(); runs(pb,body,size=11)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)

# footer note
doc.add_paragraph()
fp=doc.add_paragraph(); runs(fp,"Working draft — figures (subscriber, open-rate, conversion, churn, MRR targets) are drawn from the Year-1 cash-flow model and the pre-funding milestone ladder in the strategic addendum.",size=9,italic=True,color=GREY)

out="/home/user/prscott/Scenic_City_Insider_Founder_Manifesto.docx"
doc.save(out); print("saved",out)

# ---------- MARKDOWN ----------
md=[f"# {TITLE}", f"*{SUB}*", ""]
for head,paras in SECTIONS:
    md.append(f"## {head}")
    for para in paras: md.append(para.replace("**","**")); md.append("")
md.append("## The roadmap")
for yr,head,body in ROADMAP:
    md.append(f"**{yr} — {head}** {body}"); md.append("")
md.append("---")
md.append("*Working draft — figures drawn from the Year-1 cash-flow model and the pre-funding milestone ladder in the strategic addendum.*")
with open("/home/user/prscott/Scenic_City_Insider_Founder_Manifesto.md","w") as f:
    f.write("\n".join(md))
print("saved manifesto .md")
