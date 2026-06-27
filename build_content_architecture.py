#!/usr/bin/env python3
"""Brief Scout Media — Content Architecture: section catalog, weekly template,
12-week rotation, membership tiers, and the content->revenue map."""
import os, matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY="1B3A4B"; TEAL="2A9D8F"; ORANGE="E76F51"; SAND="E9C46A"; SLATE="415A6B"
LIGHT="EEF3F4"; GREY="6B7B85"; GOLD="B8860B"; GREEN="2E7D52"
ASSETS="/home/user/prscott/assets"; os.makedirs(ASSETS,exist_ok=True)
plt.rcParams.update({"font.family":"DejaVu Sans","figure.dpi":160})

# ---------- charts ----------
def chart_portfolio():
    fig,ax=plt.subplots(figsize=(8.6,4.6)); ax.axis("off"); ax.set_xlim(0,10); ax.set_ylim(0,10)
    left=[("DISCOVERY","What's New · Hidden Gems ·\nWho Made This? · The Dig","#"+NAVY,7.6),
          ("RELATIONSHIP","The Regular · Building's Biography ·\nThe Lens · Open Door · Almanac · Sis?","#"+TEAL,4.6),
          ("PREMIUM","Flagship Profile ·\nGetaway Field Guide","#"+ORANGE,1.6)]
    right=[("MEMBERSHIP",8.0),("PRODUCTS",4.9),("EVENTS",1.8)]
    for title,body,c,y in left:
        ax.add_patch(FancyBboxPatch((0.3,y),4.2,1.9,boxstyle="round,pad=0.04,rounding_size=0.12",fc=c,ec="white",lw=2))
        ax.text(2.4,y+1.45,title,ha="center",color="white",fontweight="bold",fontsize=11)
        ax.text(2.4,y+0.65,body,ha="center",color="white",fontsize=7.6)
    for title,y in right:
        ax.add_patch(FancyBboxPatch((6.6,y),3.1,1.4,boxstyle="round,pad=0.04,rounding_size=0.12",fc="#"+SAND,ec="white",lw=2))
        ax.text(8.15,y+0.7,title,ha="center",color="#"+NAVY,fontweight="bold",fontsize=11)
    for _,_,_,y in left:
        for _,ry in right:
            ax.add_patch(FancyArrowPatch((4.6,y+0.95),(6.55,ry+0.7),arrowstyle="-",lw=0.8,color="#B9C6CC",alpha=0.8))
    ax.text(5.0,9.6,"Every section feeds the reader-funded model — no ads",ha="center",color="#"+SLATE,fontsize=9,style="italic")
    fig.tight_layout(); fig.savefig(f"{ASSETS}/fig_portfolio.png",bbox_inches="tight"); plt.close(fig)

def chart_tiers():
    fig,ax=plt.subplots(figsize=(8.4,4.0)); ax.axis("off"); ax.set_xlim(0,10); ax.set_ylim(0,10)
    steps=[("FREE","$0","The weekly Insider +\npublic sections","#"+SLATE,0.3,2.6,3.0),
           ("MEMBER","$8/mo · $80/yr","Full archive · members-first\n'Just Licensed' intel · The Dig\ninsider · member event pricing\n· Insider Card · community","#"+TEAL,3.5,5.0,3.0),
           ("FOUNDING SCOUT","$150/yr","All Member + seasonal print\nGetaway Field Guide (mailed) +\nannual print Almanac + a\nsignature-event ticket","#"+ORANGE,6.7,7.6,3.0)]
    for name,price,body,c,x,h,w in steps:
        ax.add_patch(FancyBboxPatch((x,0.4),w,h,boxstyle="round,pad=0.03,rounding_size=0.1",fc=c,ec="white",lw=2))
        ax.text(x+w/2,0.4+h-0.5,name,ha="center",color="white",fontweight="bold",fontsize=12)
        ax.text(x+w/2,0.4+h-1.05,price,ha="center",color="white",fontsize=10,style="italic")
        ax.text(x+w/2,0.4+h/2-0.7,body,ha="center",va="center",color="white",fontsize=7.8)
    ax.text(5.0,9.5,"The ladder: each tier adds value the section catalog produces",ha="center",color="#"+NAVY,fontweight="bold",fontsize=11)
    fig.tight_layout(); fig.savefig(f"{ASSETS}/fig_tiers.png",bbox_inches="tight"); plt.close(fig)

chart_portfolio(); chart_tiers(); print("charts done")

# ---------- data ----------
CATALOG=[
 ("The Insider's Note","Anchor","Every issue","Founder's voice — the human handshake an aggregator can't fake","Affinity & referrals"),
 ("What's New Around Town","Anchor · Discovery","Weekly","Openings/closings + members-first 'Just Licensed' intel from the engine","Membership upgrade driver; new-opening crawl"),
 ("You Told Us","Anchor","Weekly","The reader mailbag — community loop made visible","Word-of-mouth growth"),
 ("The Standing Invitation","Anchor · Discovery","Weekly","One evergreen thing to go do, with a POV — the anti-calendar","Member archive; '52 Assignments' deck"),
 ("Hidden Gems","Discovery","Weekly–biweekly","The under-covered, across the wider territory — press footprint of none","Hidden-Gems map; road-trip event"),
 ("Who Made This? (Meet a Maker)","Discovery","Weekly","Everyday things traced to the local hands behind them","'Made Here' guide; maker market; collab merch"),
 ("The Dig (What's Going In There?)","Discovery","Monthly","The curiosity beat + development watch — license-engine powered","Insider-tier driver"),
 ("The Regular / Last Call","Relationship","Biweekly","People, not businesses — and reverent goodbyes when one closes","Portrait book; tribute events"),
 ("The Building's Biography","Relationship","Monthly","The story of the address — who was there before","Building-bio prints; ticketed history crawl"),
 ("The Lens","Relationship","Monthly","The city reframed by one need (Free / Newcomer / After Dark…)","Themed guide & map products"),
 ("The Open Door","Relationship","Monthly","Access & inclusion service journalism almost no one does","Accessible directory; walled partner program"),
 ("Good Dog","Relationship","Monthly","The city through a dog's eyes","Dog map; 'Yappy Hour' event; walled pet-partner"),
 ("What's Up, Sis?","Relationship","Monthly–bimonthly","Chattanooga's sister cities via the people here (no politics)","Supper-club events; 'Eat the World' passport; org partnership"),
 ("The Scenic City Almanac","Relationship","Seasonal (+annual print)","The rituals that make a year here — traditions, not events","Annual printed Almanac; seasonal socials"),
 ("The Flagship Profile","Premium","Monthly (long-form)","The prestige 'Artisans of Craft' profile — editorial, never bought","Anchors premium tier; annual 'Builders' book"),
 ("The Getaway Field Guide","Premium","Monthly + quarterly","Where a Chattanoogan goes to leave town — curated, not a brochure","Print perk (premium magnet); Insider-led trips; DMO underwriting"),
]

WEEKLY=[
 ("Every issue opens with","The Insider's Note → What's New Around Town"),
 ("Every issue includes","The Standing Invitation (one thing to go do)"),
 ("Every issue closes with","You Told Us (mailbag) → the no-ads promise + membership invite"),
 ("Plus each week","2 rotating signature pillars (see the 12-week rotation)"),
 ("Monthly","The Flagship long-form (linked from the issue, lives on the site/premium)"),
]

ROT=[
 ("1","Hidden Gems","The Open Door",""),
 ("2","Who Made This?","The Building's Biography","Flagship profile drops"),
 ("3","Hidden Gems","The Lens — 'Free Chattanooga'",""),
 ("4","Who Made This?","Good Dog",""),
 ("5","Hidden Gems","The Dig — 'what's going in there?'",""),
 ("6","Who Made This?","What's Up, Sis? — Accra/Ghana",""),
 ("7","Hidden Gems","The Regular",""),
 ("8","Who Made This?","The Almanac — the season turns","Getaway: Worth the Drive — Knoxville"),
 ("9","Hidden Gems","The Building's Biography","Flagship profile drops"),
 ("10","Who Made This?","The Lens — 'Newcomer's First Week'",""),
 ("11","Hidden Gems","The Open Door",""),
 ("12","Who Made This?","The Regular","Getaway: The Long Weekend — Asheville (quarterly)"),
]

TIERS=[
 ("Free","$0","The weekly Insider newsletter and public sections. Top-of-funnel; builds the list and the habit."),
 ("Member","$8/mo or $80/yr","Full members-only archive; members-first 'Just Licensed' intel; The Dig insider notes; the Standing-Invitation assignment archive; member event pricing; the Insider Card; community access."),
 ("Founding Insider","$150/yr","Everything in Member, plus the seasonal print Getaway Field Guide mailed to your door, the annual printed Scenic City Almanac, a free signature-event ticket, the Flagship in print, and name recognition."),
]

REV=[
 ("What's New / The Dig","—","New-opening crawl","Members-first intel (upgrade driver)"),
 ("Hidden Gems","Hidden-Gems map/guide","Gem-hunting road trip","—"),
 ("Who Made This?","'Made Here' field guide; maker merch","Maker market","Maker collabs"),
 ("The Standing Invitation","'52 Assignments' deck","—","Member archive"),
 ("The Regular / Last Call","'Regulars' portrait book","Tribute / legacy nights","—"),
 ("The Building's Biography","Building-bio prints","Ticketed history crawl","JV with featured venues"),
 ("The Lens","Themed guides & maps (Free/Date/Newcomer)","—","Newcomer guide as signup magnet"),
 ("The Open Door","'Accessible Chattanooga' directory","—","Walled business-partner listings"),
 ("Good Dog","Dog-friendly map","'Yappy Hour'","Walled pet-business partners"),
 ("What's Up, Sis?","'Eat the World' passport","Sister-city supper club","Sister Cities org partnership"),
 ("The Scenic City Almanac","Annual printed Almanac (giftable)","Seasonal socials (strawberry, leaf-turn)","—"),
 ("The Flagship Profile","Annual 'Builders' book","—","Anchors the premium tier"),
 ("The Getaway Field Guide","Seasonal print guide (premium perk + sales)","Insider-led group trips","DMO / destination underwriting (walled)"),
]

# ---------- doc ----------
doc=Document(); sec=doc.sections[0]; sec.orientation=WD_ORIENT.LANDSCAPE
sec.page_width,sec.page_height=Inches(11),Inches(8.5)
for m in ("top_margin","bottom_margin","left_margin","right_margin"): setattr(sec,m,Inches(0.65))
nm=doc.styles["Normal"]; nm.font.name="Calibri"; nm.font.size=Pt(10); nm.paragraph_format.space_after=Pt(4)
def sca(r,h): r.font.color.rgb=RGBColor.from_string(h)
def shade(c,h):
    p=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd")
    s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); p.append(s)
def ctext(c,t,bold=False,color=None,size=8.6,white=False,align="left"):
    c.text=""; p=c.paragraphs[0]
    p.alignment={"left":WD_ALIGN_PARAGRAPH.LEFT,"center":WD_ALIGN_PARAGRAPH.CENTER}[align]
    r=p.add_run(t); r.font.size=Pt(size); r.font.bold=bold
    if white: r.font.color.rgb=RGBColor.from_string("FFFFFF")
    elif color: r.font.color.rgb=RGBColor.from_string(color)
def H(t,level=1,color=None):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(12 if level==1 else 8); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(t); r.font.bold=True
    if level==1:
        r.font.size=Pt(16); sca(r,color or NAVY)
        pPr=p._p.get_or_add_pPr(); b=OxmlElement("w:pBdr"); bt=OxmlElement("w:bottom")
        bt.set(qn("w:val"),"single"); bt.set(qn("w:sz"),"6"); bt.set(qn("w:space"),"2"); bt.set(qn("w:color"),TEAL)
        b.append(bt); pPr.append(b)
    else: r.font.size=Pt(12.5); sca(r,color or TEAL)
def P(t,italic=False,color=None,size=10.5):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(5)
    for i,ch in enumerate(t.split("**")):
        if not ch: continue
        r=p.add_run(ch); r.font.size=Pt(size); r.font.italic=italic
        if i%2==1: r.font.bold=True
        if color: sca(r,color)
def FIG(path,cap,width=8.4):
    doc.add_picture(path,width=Inches(width)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    P(cap,italic=True,color=GREY,size=8.5); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
def TABLE(headers,rows,widths,fontsize=8.6,first_bold=True,caption=None):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style="Table Grid"
    for i,h in enumerate(headers):
        shade(t.rows[0].cells[i],NAVY); ctext(t.rows[0].cells[i],h,bold=True,white=True,size=fontsize)
    for r,row in enumerate(rows):
        cells=t.add_row().cells; base=LIGHT if r%2==0 else "FFFFFF"
        for i,v in enumerate(row):
            shade(cells[i],base)
            ctext(cells[i],str(v),bold=(first_bold and i==0),size=fontsize,color=NAVY if (first_bold and i==0) else None)
    for i,w in enumerate(widths):
        for row in t.rows: row.cells[i].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    if caption: P(caption,italic=True,color=GREY,size=8.3)

# title
for _ in range(2): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("THE SCENIC CITY SCOUT"); r.font.size=Pt(25); r.font.bold=True; sca(r,NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Content Architecture"); r.font.size=Pt(19); sca(r,TEAL)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("The section catalog, weekly issue template, a 12-week rotation, the membership tiers, and how every section feeds the business")
r.font.size=Pt(11.5); r.font.italic=True; sca(r,SLATE)
doc.add_page_break()

H("The principle: no commodity content",1)
P("The free incumbent owns the commodity lane — the generic event calendar, weather, traffic, national headlines, paid 'best-of' lists. **We run none of it.** Every section in this architecture is a point of view only the Insider has, which is the reason a membership is worth paying for. The portfolio runs on three layers — **Discovery** (acquisition), **Relationship** (retention), and a **Premium** layer — and each section throws off a product, an event, or a membership perk, so the content *is* the reader-funded business model.")
FIG(f"{ASSETS}/fig_portfolio.png","Figure 1. The content portfolio in three layers, every one feeding the ad-free model.",8.6)

H("Editorial standards — The Once Rule & the Nerd's veto",1)
P("These are non-negotiable, because repetition and filler make independent editorial look like advertising or automation — the two things that would destroy the brand's whole premise and the founder's personal credibility.")
def _bul(items,size=10):
    for it in items:
        p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(3)
        for i,ch in enumerate(it.split("**")):
            if not ch: continue
            r=p.add_run(ch); r.font.size=Pt(size)
            if i%2==1: r.font.bold=True
_bul([
 "**The Once Rule.** A business or location appears **once per issue — never twice.** Across a month, **no more than twice, and twice is rare** — reserved for genuine news (a closing, then later a reopening) and always from a *fresh* angle, never a repeat. Recurring *threads* are welcome; recurring *mentions of the same name* are not.",
 "**Freshness across issues.** Each issue's cast of businesses should be distinct from the last issue's. If a place was the Hidden Gem last week, it is not the 'now open' note this week.",
 "**The Nerd's veto.** The research layer does more than check facts — it **kills copy that reads forced, lazy, padded, or 'paid-for,' and anything that smells like AI slop** (recycled names, generic filler, the same story told twice). If it wouldn't survive a sharp local reader's eye, it doesn't run.",
 "**The test before publish:** would a skeptical Chattanoogan think this was written by a person who actually went there — or by a machine padding a template? Only the first ships.",
])

H("The section catalog",1)
TABLE(["Section","Layer","Frequency","The point of view","Primary monetization"],
      CATALOG, widths=[1.9,1.4,1.25,3.1,2.6], fontsize=8.4,
      caption="Sixteen sections; any single weekly issue runs the anchors + 2–3 rotating pillars.")

doc.add_page_break()
H("The weekly issue template",1)
TABLE(["Slot","What runs"],WEEKLY,widths=[2.3,7.0],fontsize=10,
      caption="A five-minute read: warm open, what's new, two rotating features, one thing to go do, the mailbag, and the membership invite.")
H("A 12-week sample rotation",1)
P("Illustrative placement — monthly and seasonal pillars recur on roughly the cadence in the catalog. Anchors (Insider's Note, What's New, The Standing Invitation, You Told Us) run **every** week and aren't repeated below.")
TABLE(["Wk","Rotating feature A (discovery)","Rotating feature B (pillar)","Monthly / premium drop"],
      ROT, widths=[0.6,3.0,3.3,3.0], fontsize=8.8, first_bold=True)

doc.add_page_break()
H("The membership ladder",1)
FIG(f"{ASSETS}/fig_tiers.png","Figure 2. Three tiers; each higher tier unlocks value the section catalog already produces.",8.4)
TABLE(["Tier","Price","What you get"],TIERS,widths=[1.6,1.6,6.1],fontsize=9.2,
      caption="Free acquires; Member is the core reader-revenue tier; Founding Insider is built around the print perks (Getaway Field Guide + Almanac).")

doc.add_page_break()
H("Content → revenue map",1)
P("How each monetizable section converts into the three reader-funded lines (products, events, membership/partner) — no advertising anywhere.")
TABLE(["Section","Product","Event","Membership / walled partner"],
      REV, widths=[2.2,2.6,2.2,2.3], fontsize=8.5, first_bold=True)
P("Walled partner lines (Open Door listings, Good Dog pet-partners, Getaway/DMO underwriting, the Insider Card) follow the bright line: a business can buy a booth, a listing, a ticket, or sponsor an experience — never a sentence of editorial.",
  italic=True,color=GREY,size=9)

out="/home/user/prscott/Content_Architecture.docx"
doc.save(out); print("saved",out)
