#!/usr/bin/env python3
"""Famous Faces & Big Ideas — Born Here (Vol. 2). People, inventions, prep-school alumni — 7-scout compile."""
import os, csv
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY="1B3A4B"; TEAL="2A9D8F"; ORANGE="E76F51"; SAND="E9C46A"; SLATE="415A6B"
LIGHT="EEF3F4"; GREY="6B7B85"; GOLD="B8860B"; GREEN="2E7D52"; RED="C0492F"

# category -> list of (name, known_for, tie, status, fact, conf)
PEOPLE={
"Music & the Arts":[
 ("Bessie Smith","'Empress of the Blues'","Born Chattanooga, 1894","Dec.","Highest-paid Black entertainer of her era.","High"),
 ("Clyde Stubblefield","James Brown's 'Funky Drummer'","Born Chattanooga, 1943","Dec.","His break is the most-sampled in recorded music — and he went largely uncredited.","High"),
 ("Jimmy Blanton","Jazz bassist (Duke Ellington)","Born Chattanooga, 1918","Dec.","Reinvented the jazz bass as a solo voice; died of TB at just 23.","High"),
 ("Valaida Snow","'Queen of the Trumpet' jazz star","Born Chattanooga, 1904","Dec.","Played 10+ instruments; a superstar on three continents in the 1930s.","High"),
 ("Usher","Grammy-winning R&B superstar","Raised Chattanooga (b. Dallas '78)","Living","Started in a Chattanooga church choir at age 6.","High"),
 ("Kane Brown","Country/R&B superstar","Born Chattanooga, 1993","Living","Got the key to the city; 'Kane Brown Day' before his Finley Stadium show.","High"),
 ("Isaiah Rashad","Rapper (Top Dawg Entertainment)","From Chattanooga (b. 1991)","Living","Co-founded the local hip-hop collective 'TheHouse.'","High"),
 ("George Ayers Cress","Painter (UTC's Cress Gallery namesake)","Lived/worked Chattanooga (b. AL '21)","Dec.","Led the TN Arts Council and sat on the Hunter Museum board.","High"),
],
"Screen & Stage":[
 ("Samuel L. Jackson","Actor (Pulp Fiction, Marvel, Star Wars)","Raised Chattanooga (b. D.C. '48); Riverside HS","Living","Played French horn, trumpet, flute & piccolo in his school orchestra.","High"),
 ("Leslie Jordan","Emmy-winning actor; viral pandemic star","Born Chattanooga, 1955; attended UTC","Dec.","Stood 4'11\"; became a social-media sensation in his 60s.","High"),
 ("Lori Petty","Actor (A League of Their Own, Tank Girl)","Born Chattanooga, 1963","Living","Was a graphic designer in Omaha before acting.","High"),
 ("Rachel Boston","Actor (Hallmark, Witches of East End)","Born Chattanooga '82; raised Signal Mtn","Living","Was Miss Tennessee Teen USA in 1999.","High"),
 ("Hugh Beaumont","'Ward Cleaver' on Leave It to Beaver","Baylor School grad, 1930","Dec.","Held a master's in theology and was an ordained Methodist lay minister.","High"),
],
"Writers & Thinkers":[
 ("Arthur Golden","Author, Memoirs of a Geisha","Born Chattanooga '56; Baylor grad '74; raised Lookout Mtn","Living","Rewrote the novel three times over six years.","High"),
 ("Jon Meacham","Pulitzer-winning presidential historian","Born Chattanooga; McCallie grad '87","Living","Received the National Humanities Medal.","High"),
 ("Ishmael Reed","Novelist/poet; MacArthur 'genius' Fellow","Born Chattanooga '38 (family left at age 4)","Living","Began as a teenage jazz columnist.","High"),
 ("Carman Barnes","Novelist/screenwriter (Schoolgirl, 1929)","Born Chattanooga '12; attended GPS","Dec.","Wrote a national bestseller at ~16; landed a Paramount contract.","High"),
 ("Bill Dedman","Pulitzer journalist (Empty Mansions)","Chattanooga native; attended Baylor","Living","His Huguette Clark series was the most-read feature in NBC News web history.","High"),
 ("Mary Q. Steele","Newbery Honor children's author","Born Chattanooga, 1922","Dec.","Also wrote as 'Wilson Gage.'","High"),
 ("Albert Hodges Morehead","NYT bridge editor; lexicographer","Baylor grad, 1925","Dec.","The leading authority of his era on contract bridge.","High"),
 ("Emma Bell Miles","Appalachian author/artist (Spirit of the Mountains)","Lived Walden's Ridge (b. Indiana '79)","Dec.","Wrote a column for the Chattanooga News.","High"),
],
"Athletes":[
 ("Reggie White","NFL Hall of Famer, 'Minister of Defense'","Born & raised Chattanooga '61; Howard High","Dec.","Ordained a minister at 17; Reggie White Blvd is named for him. (Note: Howard School, NOT McCallie.)","High"),
 ("Nick Kurtz","MLB (Athletics); 2025 AL Rookie of the Year","Baylor grad, 2021","Living","The freshest name on the list.","High"),
 ("Geoff Gaberino","Olympic gold swimmer (1984)","Baylor grad, 1980","Living","Often called 'Chattanooga's greatest swimmer.'","High"),
 ("Keith Mitchell","PGA Tour (2019 Honda Classic)","Born & raised Chattanooga '92; Baylor","Living","Part of the Baylor PGA 'golf factory' (now lives on the GA coast).","High"),
 ("Stephan Jaeger","PGA Tour; 2024 Houston Open & Paris Olympics","Baylor '08 + UTC + still lives here","Living","The rare star who's truly all-three (schooled, college, resident).","High"),
 ("Harris English","PGA Tour (4 wins, Ryder Cup)","Baylor '03–07 (raised Thomasville GA)","Living","The winningest of the Baylor four; UGA teammate of Harman/Reed.","High"),
 ("Luke List","PGA Tour (2 wins)","Attended Baylor (grew up Ringgold GA)","Living","Won his first Tour title at 37.","High"),
 ("Vonn Bell","NFL safety (Saints/Bengals)","Born Chattanooga, 1994","Living","Now a college safeties coach.","High"),
 ("Herman Hickman","College FB HOF; Yale coach; pro wrestler","Baylor grad, 1928","Dec.","Wrestled as 'The Tennessee Terror' (500+ matches).","High"),
 ("Dave Bristol","MLB manager (helped build the 'Big Red Machine')","Baylor grad, 1951","Dec.","Cincinnati Reds Hall of Fame.","High"),
 ("Jacques McClendon","NFL center; Rams exec (Super Bowl LVI); WME","Attended Baylor (~'06)","Living","Earned a Brown MBA while in the NFL.","High"),
 ("Michael Bingham","Olympic sprinter (GB 4x400 bronze, 2008)","McCallie grad, 2004","Living","Got his medal 8 years late after a doping DQ.","High"),
],
"Public Life — Politics, Civic & History":[
 ("Howard Baker Jr.","U.S. Senate Majority Leader; Reagan WH Chief of Staff","McCallie grad, 1943","Dec.","The 'Great Conciliator.'","High"),
 ("Bill Brock","U.S. Senator; Labor Secretary; RNC chair","McCallie grad '49 (Brock Candy family)","Dec.","Chaired the RNC after Watergate.","High"),
 ("Pat Robertson","Founder, CBN / The 700 Club / Regent University","McCallie grad, 1946","Dec.","His Chattanooga obituary headlined the McCallie tie.","High"),
 ("Estes Kefauver","U.S. Senator; 1956 Democratic VP nominee","Began his law/political career in Chattanooga (b. Madisonville)","Dec.","His 1950–51 televised mob hearings made him an early TV-political star.","High"),
 ("Ed Johnson","Subject of United States v. Shipp (1906)","Chattanooga resident","Dec.","His case is the ONLY criminal trial ever held by the U.S. Supreme Court; exonerated 2000.","High"),
 ("Fob James","Governor of Alabama; founder, Diversified Products","Baylor grad, 1952","Living","Served as governor as both a Democrat and a Republican.","High"),
 ("Carroll Campbell","Governor of South Carolina","McCallie grad","Dec.","A TN prep grad who became a defining modern SC governor.","Med"),
 ("David Abshire","U.S. Ambassador to NATO; co-founder of CSIS","Baylor grad, 1944","Dec.","Co-founded one of the world's top think tanks.","High"),
 ("Zach Wamp","U.S. Congressman (TN-3)","McCallie grad '76; raised East Ridge","Living","His son Weston is now Hamilton County mayor — a dynasty.","High"),
 ("Andy Berke","Chattanooga mayor (2013–21); USDA official","Baylor grad","Living","—","High"),
 ("Robert E. Cooper Jr.","Tennessee Attorney General","Baylor grad, 1975","Living","(Often confused with 'Thomas J. Cooper' in error.)","High"),
 ("Charlie Norwood","U.S. Congressman (GA); dentist","Baylor grad, 1959","Dec.","A practicing dentist and Vietnam combat vet before Congress.","High"),
 ("Frances Zwenig","Senate chief of staff (John Kerry); US-ASEAN trade leader","GPS grad, 1963","Living","Helped run the Senate's POW/MIA investigation.","Med"),
],
"Business & Founders":[
 ("Ted Turner","Founder, CNN/TBS; media mogul; ex-Braves owner","McCallie grad, 1956","Living","Disciplined repeatedly as a student; later donated millions back to McCallie.","High"),
 ("Adolph S. Ochs","Built the modern New York Times","Career Chattanooga (ran the Chattanooga Times from 1878; died here '35)","Dec.","At 19 he borrowed $250 to buy the failing Chattanooga Times.","High"),
 ("George Thomas Hunter","Coca-Cola Bottling magnate; philanthropist","Attended Baylor","Dec.","Founded the Benwood Foundation and the Hunter Museum of American Art.","High"),
 ("Thomas Cartter Lupton","Coca-Cola bottling heir; philanthropist","Baylor (~'17)","Dec.","Founded the Lyndhurst Foundation; ~$200M estate was the largest then probated in the South.","Med"),
 ("John M. Belk","CEO of Belk department stores; mayor of Charlotte","Attended McCallie (~'39)","Dec.","Ran the chain his father founded.","Med"),
 ("Billy Dunavant Jr.","One of the world's largest cotton merchants","Attended McCallie (~'50)","Dec.","Shaped the global cotton futures market from Memphis.","Med"),
 ("Olan Mills II","Chairman, Olan Mills portrait-studio empire","McCallie grad, 1948","Dec.","Ran the ubiquitous mid-century American portrait business.","Med"),
 ("Jo Conn Guild","Electric-utility leader; prominent anti-TVA campaigner","Baylor grad, 1905","Dec.","Fought the federal TVA on behalf of private power.","Med"),
 ("Tim Kelly","Chattanooga mayor; auto-dealer entrepreneur","Born Chattanooga '67; Baylor grad '85","Living","Co-founded the Chattanooga Football Club before politics.","High"),
]}

INVENT=[
 ("The tow truck (wrecker)","Ernest Holmes Sr.","Chattanooga, 1916","Born from a creek rescue of a friend's Model T; Chattanooga is home to the International Towing & Recovery Hall of Fame.","High"),
 ("Miniature golf ('Tom Thumb Golf')","Garnet Carter","Lookout Mtn (GA), ~1927","Built to amuse hotel guests; he patented obstacle putting, then poured the fortune into Rock City.","High"),
 ("Coca-Cola bottling (the system)","Benjamin Thomas & Joseph Whitehead","Chattanooga, 1899","Bought near-nationwide bottling rights from Asa Candler for ~$1 — he thought bottling was a fad.","High"),
 ("The MoonPie","Earl Mitchell / Chattanooga Bakery","Chattanooga, 1917","A coal miner framed the moon and said 'about that big.' ~1M still made daily.","High"),
 ("Double-Cola","Charles Little & Joe Foster","Chattanooga, 1933","The only major soft drink actually invented here (Coke was only bottled here).","High"),
 ("Krystal","Davenport & Sherrill","Chattanooga, 1932","The South's oldest fast-food chain (2nd-oldest in the U.S.).","High"),
 ("The snack 'family pack'","O.D. & Ruth McKee (Little Debbie)","Chattanooga/Collegedale, 1960","Pioneered the 12-count multipack of individually wrapped cakes.","High"),
 ("First community-wide gigabit (EPB)","EPB (municipal utility)","Chattanooga, 2010","First U.S. 'Gig City' — beat Google; later first to 10-Gig and 25-Gig.","High"),
]

PREP=[
 ("McCallie School","Statesmen & media titans","Ted Turner, Howard Baker Jr., Bill Brock, Pat Robertson, Jon Meacham, Zach Wamp, Carroll Campbell — plus business names Belk, Dunavant & Olan Mills II."),
 ("Baylor School","Authors, Olympians & a PGA 'golf factory'","Arthur Golden, Bill Dedman, Hugh Beaumont, Olympic swimmer Geoff Gaberino, MLB's Nick Kurtz & Dave Bristol, governors Fob James & NATO's David Abshire, mayors Kelly & Berke, the Coca-Cola founding families (Hunter, Lupton) — and FOUR PGA Tour pros (English, Mitchell, Jaeger, List) on Tour the same season, a reported first for any U.S. high school (coach King Oehmig)."),
 ("Girls Preparatory School (GPS)","A thinner national bench — be honest","Best-verified: novelist Carman Barnes ('29 bestseller Schoolgirl) and Senate chief-of-staff / trade leader Frances Zwenig ('63). Many 'Distinguished Alumnae' are excellent local/civic figures — a 'local heroes' sidebar, not national names."),
]
DONOTCLAIM=[
 "**Terrell Owens** — Alexander City, AL; the tie is UTC (college) only, never a 'native.' Same for NBA's Gerald Wilkins & Tom Hammonds (UTC only).",
 "**Reggie White did NOT attend McCallie** — he went to Howard School. (Classic local mix-up.)",
 "**Samuel L. Jackson** was raised here but born in Washington, D.C.; **Kefauver** born in Madisonville; **Ishmael Reed**'s family left at age 4.",
 "**Grace Moore** (opera star) — born in Cocke County / only buried here. **John R. Cherry** (Ernest films) = Nashville. **Walter Lang** (director) = Memphis.",
 "**Baylor School ≠ Baylor University (Waco)** — Rand Paul, Thasunda Brown Duckett, etc. are NOT Baylor School. **Davis Tarwater** (swimmer) = Webb School, Knoxville. **Nick Tiano** (QB) = Baylor, not McCallie. **Brandt Snedeker** (golf) = Nashville.",
 "**Cornelia Clark** (TN Chief Justice) — no GPS tie (name collision). A 'dental drill bit invented here' claim could NOT be verified — keep it out unless a named inventor is confirmed.",
]

# ---------- CSV ----------
with open("/home/user/prscott/famous_faces_born_here.csv","w",newline="") as f:
    w=csv.writer(f); w.writerow(["Section","Name/Item","Known for / Inventor","Tie / When","Status","Surprising fact","Confidence"])
    for cat,rows in PEOPLE.items():
        for (n,k,t,s,fa,c) in rows: w.writerow((cat,n,k,t,s,fa,c))
    for (inv,who,when,fact,c) in INVENT: w.writerow(("Invention",inv,who,when,"",fact,c))
ppl=sum(len(v) for v in PEOPLE.values())
print("people:",ppl,"inventions:",len(INVENT))

# ---------- DOCX ----------
doc=Document(); sec=doc.sections[0]; sec.orientation=WD_ORIENT.LANDSCAPE
sec.page_width,sec.page_height=Inches(11),Inches(8.5)
for m in ("top_margin","bottom_margin","left_margin","right_margin"): setattr(sec,m,Inches(0.6))
nm=doc.styles["Normal"]; nm.font.name="Calibri"; nm.font.size=Pt(10); nm.paragraph_format.space_after=Pt(4)
def sca(r,h): r.font.color.rgb=RGBColor.from_string(h)
def shade(c,h):
    p=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd")
    s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); p.append(s)
def ctext(c,t,bold=False,color=None,size=8.4,white=False,align="left"):
    c.text=""; p=c.paragraphs[0]
    p.alignment={"left":WD_ALIGN_PARAGRAPH.LEFT,"center":WD_ALIGN_PARAGRAPH.CENTER}[align]
    r=p.add_run(t); r.font.size=Pt(size); r.font.bold=bold
    if white: r.font.color.rgb=RGBColor.from_string("FFFFFF")
    elif color: r.font.color.rgb=RGBColor.from_string(color)
def H(t,level=1,color=None):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(11 if level==1 else 7); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(t); r.font.bold=True
    if level==1:
        r.font.size=Pt(15); sca(r,color or NAVY)
        pPr=p._p.get_or_add_pPr(); b=OxmlElement("w:pBdr"); bt=OxmlElement("w:bottom")
        bt.set(qn("w:val"),"single"); bt.set(qn("w:sz"),"6"); bt.set(qn("w:space"),"2"); bt.set(qn("w:color"),TEAL)
        b.append(bt); pPr.append(b)
    else: r.font.size=Pt(12); sca(r,color or TEAL)
def P(t,italic=False,color=None,size=10.5):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(5)
    for i,ch in enumerate(t.split("**")):
        if not ch: continue
        r=p.add_run(ch); r.font.size=Pt(size); r.font.italic=italic
        if i%2==1: r.font.bold=True
        if color: sca(r,color)
def BUL(items,size=9.5):
    for it in items:
        p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(3)
        for i,ch in enumerate(it.split("**")):
            if not ch: continue
            r=p.add_run(ch); r.font.size=Pt(size)
            if i%2==1: r.font.bold=True
def PTABLE(rows):
    cols=["Name","Known for","Chattanooga tie","Status","The surprising fact","Conf."]
    widths=[1.5,1.95,1.95,0.6,3.3,0.5]
    t=doc.add_table(rows=1,cols=len(cols)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style="Table Grid"
    for i,h in enumerate(cols):
        shade(t.rows[0].cells[i],NAVY); ctext(t.rows[0].cells[i],h,bold=True,white=True,size=8.6,align="left" if i<5 else "center")
    for r,(n,k,tie,s,fa,c) in enumerate(rows):
        cells=t.add_row().cells; base=LIGHT if r%2==0 else "FFFFFF"
        ctext(cells[0],n,bold=True,color=NAVY,size=8.4); ctext(cells[1],k,size=8.0)
        ctext(cells[2],tie,size=8.0); ctext(cells[3],s,size=7.8,align="center")
        ctext(cells[4],fa,size=8.0); ctext(cells[5],c,size=8.0,align="center")
        for cc in cells: shade(cc,base)
    for i,w in enumerate(widths):
        for row in t.rows: row.cells[i].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)

# Title
doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("THE SCENIC CITY SCOUT"); r.font.size=Pt(23); r.font.bold=True; sca(r,NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Famous Faces & Big Ideas — Born Here"); r.font.size=Pt(17); sca(r,TEAL)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run(f"{ppl} verified people + {len(INVENT)} homegrown inventions + the prep-school powerhouses — companion to the 'Born Here' brands book")
r.font.size=Pt(11); r.font.italic=True; sca(r,SLATE)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Tie is stated precisely: born / raised / schooled / career / lived here. Misattributions parked in the 'do-not-claim' box.")
r.font.size=Pt(9.5); sca(r,GREY)
doc.add_page_break()

for cat,rows in PEOPLE.items():
    H(cat+f"  ({len(rows)})",1); PTABLE(rows)

doc.add_page_break()
H("Chattanooga Invented That",1)
t=doc.add_table(rows=1,cols=5); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style="Table Grid"
for i,h in enumerate(["Invention / first","Inventor(s)","When","The story","Conf."]):
    shade(t.rows[0].cells[i],NAVY); ctext(t.rows[0].cells[i],h,bold=True,white=True,size=8.6,align="left" if i<4 else "center")
for r,(inv,who,when,fact,c) in enumerate(INVENT):
    cells=t.add_row().cells; base=LIGHT if r%2==0 else "FFFFFF"
    ctext(cells[0],inv,bold=True,color=NAVY,size=8.4); ctext(cells[1],who,size=8.0); ctext(cells[2],when,size=8.0); ctext(cells[3],fact,size=8.0); ctext(cells[4],c,size=8.0,align="center")
    for cc in cells: shade(cc,base)
for i,w in enumerate([2.0,1.9,1.0,4.4,0.5]):
    for row in t.rows: row.cells[i].width=Inches(w)
doc.add_paragraph().paragraph_format.space_after=Pt(2)

H("The three prep-school powerhouses",1)
for school,tag,body in PREP:
    P(f"**{school} — {tag}.** {body}")

H("Do not claim (misattributions the Scouts caught)",1,color=ORANGE)
BUL(DONOTCLAIM)

H("How to use it",1)
P("This is the source bench for occasional **'Born Here'** people/ideas features and for **trivia** (a natural membership/engagement hook). Lead with the surprising fact, keep the tie honest (born vs raised vs schooled), and re-verify before publishing. Pairs with the brands book, *Still Standing*, and the Almanac — and supports a heritage trail/map product and a school-rivalry trivia event. Full data is in famous_faces_born_here.csv.")

out="/home/user/prscott/Famous_Faces_Born_Here.docx"
doc.save(out); print("saved",out)
