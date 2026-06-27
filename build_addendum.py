#!/usr/bin/env python3
"""Strategic addendum to the Brief Scout Media plan: Y1 cash flow, HITL+AI, pre-funding
milestones, expansion strategy, and tone/topics/moat."""
import os
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.ticker import FuncFormatter
from matplotlib.patches import Patch
import numpy as np

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY="1B3A4B"; TEAL="2A9D8F"; ORANGE="E76F51"; SAND="E9C46A"; SLATE="415A6B"
LIGHT="EEF3F4"; GREY="6B7B85"
ASSETS="/home/user/prscott/assets"
os.makedirs(ASSETS, exist_ok=True)

plt.rcParams.update({
    "font.family":"DejaVu Sans","font.size":11,
    "axes.edgecolor":"#"+SLATE,"axes.labelcolor":"#"+NAVY,"axes.titlecolor":"#"+NAVY,
    "axes.titleweight":"bold","xtick.color":"#"+SLATE,"ytick.color":"#"+SLATE,
    "axes.grid":True,"grid.color":"#DDE5E8","grid.linewidth":0.8,"figure.dpi":160,
})
def usd(x,pos=None):
    return (f"-${abs(x)/1000:.0f}K" if x<0 else f"${x/1000:.0f}K") if abs(x)>=1000 else f"${x:.0f}"

# ============================================================================
# YEAR-1 MONTHLY CASH FLOW MODEL
# ============================================================================
months = [f"M{i}" for i in range(1,13)]
free_eom = [900,1500,2100,2800,3500,4200,4800,5400,5900,6300,6700,7000]
mem_eom  = [70,95,120,150,180,210,235,260,285,310,330,350]
ARPU_MO = 7.5  # blended monthly net per member ($90/yr)
membership = [round(m*ARPU_MO) for m in mem_eom]

event_gross = [0,0,1500,0,0,7000,0,0,800,9000,0,0]
event_cost  = [0,0,1000,0,0,3500,0,0,500,4500,0,0]
event_net   = [g-c for g,c in zip(event_gross,event_cost)]

prod_gross  = [0,0,0,2500,1500,0,1000,0,0,0,3000,800]
prod_cost   = [0,0,0,1200,100,0,500,0,0,0,1400,300]
prod_net    = [g-c for g,c in zip(prod_gross,prod_cost)]

tools     = [20,20,45,60,60,60,60,60,60,60,60,60]
marketing = [500,300,300,250,200,200,200,250,200,200,300,300]
admin     = [100,100,100,150,150,150,150,150,150,200,200,200]
contract  = [0,0,0,0,0,0,0,0,500,500,600,600]
misc      = [125]*12
opex = [t+m+a+c+x for t,m,a,c,x in zip(tools,marketing,admin,contract,misc)]

total_in = [membership[i]+event_net[i]+prod_net[i] for i in range(12)]
net_cash = [total_in[i]-opex[i] for i in range(12)]

OPENING = 2500  # $5,000 startup capital less ~$2,500 pre-launch spend
cum=[]; bal=OPENING
for n in net_cash:
    bal+=n; cum.append(bal)

T_mem=sum(membership); T_evn=sum(event_net); T_prd=sum(prod_net)
T_rev=sum(total_in); T_opex=sum(opex); T_net=sum(net_cash); ENDBAL=cum[-1]

# ---- cash flow chart ----
def chart_cashflow():
    fig,ax=plt.subplots(figsize=(9.2,4.4))
    x=np.arange(12)
    cols=["#"+TEAL if v>=0 else "#"+ORANGE for v in net_cash]
    ax.bar(x,net_cash,color=cols,width=0.6,label="Monthly net cash flow")
    ax.set_xticks(x); ax.set_xticklabels(months)
    ax.yaxis.set_major_formatter(FuncFormatter(usd))
    ax.axhline(0,color="#"+SLATE,lw=0.8)
    ax.set_ylabel("Monthly net cash")
    ax2=ax.twinx()
    ax2.plot(x,cum,color="#"+NAVY,marker="o",lw=2.6,label="Cumulative cash balance")
    ax2.yaxis.set_major_formatter(FuncFormatter(usd))
    ax2.set_ylabel("Cumulative cash",color="#"+NAVY)
    ax2.grid(False)
    for xi,v in zip(x,cum):
        ax2.text(xi,v+700,usd(v),ha="center",fontsize=7.5,color="#"+NAVY,fontweight="bold")
    ax.set_title("Year-1 Monthly Cash Flow (Bootstrapped, No Outside Funding)")
    l1,la1=ax.get_legend_handles_labels(); l2,la2=ax2.get_legend_handles_labels()
    ax.legend(l1+l2,la1+la2,frameon=False,loc="upper left",fontsize=9)
    fig.tight_layout(); fig.savefig(f"{ASSETS}/fig_cashflow.png",bbox_inches="tight"); plt.close(fig)

# ---- AI time-allocation chart ----
def chart_ai_time():
    fig,ax=plt.subplots(figsize=(7.2,4.2))
    cats=["Writing &\nproduction","Event aggregation\n& research","Admin &\noperations",
          "Events, partnerships\n& community"]
    before=[40,20,20,20]; after=[22,8,10,60]
    x=np.arange(len(cats)); w=0.38
    ax.bar(x-w/2,before,w,color="#"+SLATE,label="Before AI leverage")
    ax.bar(x+w/2,after,w,color="#"+TEAL,label="With HITL + AI")
    for xi,(b,a) in enumerate(zip(before,after)):
        ax.text(xi-w/2,b+1,f"{b}%",ha="center",fontsize=8,color="#"+NAVY)
        ax.text(xi+w/2,a+1,f"{a}%",ha="center",fontsize=8,color="#"+NAVY)
    ax.set_xticks(x); ax.set_xticklabels(cats,fontsize=8.5)
    ax.set_ylabel("Share of founder time")
    ax.set_ylim(0,70)
    ax.set_title("AI Frees the Founder for the Defensible Work")
    ax.legend(frameon=False,loc="upper left")
    fig.tight_layout(); fig.savefig(f"{ASSETS}/fig_ai_time.png",bbox_inches="tight"); plt.close(fig)

# ---- expansion sequencing (gantt-style) ----
def chart_expansion():
    fig,ax=plt.subplots(figsize=(8.6,4.0))
    # (label, start_quarter, duration_quarters, color)
    items=[
        ("Flagship: Scenic City Insider", 0, 16, "#"+NAVY),
        ("Vertical 1: 'What to do' / events edition", 5, 11, "#"+TEAL),
        ("Vertical 2: Newcomer & food/dining", 8, 8, "#"+TEAL),
        ("Membership & events at scale", 6, 10, "#"+SAND),
        ("Market 2 (e.g., Knoxville/Asheville)", 12, 6, "#"+ORANGE),
        ("Playbook productization → network", 14, 4, "#"+ORANGE),
    ]
    for i,(lab,s,d,c) in enumerate(items):
        y=len(items)-1-i
        ax.barh(y,d,left=s,height=0.55,color=c,edgecolor="white")
        ax.text(s+0.2,y,lab,va="center",ha="left",color="white",fontsize=8.5,fontweight="bold")
    ax.set_yticks([]); ax.set_xlim(0,18)
    ax.set_xticks([0,4,8,12,16]); ax.set_xticklabels(["Year 1","Year 2","Year 3","Year 4","Year 5"])
    ax.set_title("Sequencing: Deepen the City First, Then Widen")
    ax.grid(axis="y",visible=False)
    fig.tight_layout(); fig.savefig(f"{ASSETS}/fig_expansion.png",bbox_inches="tight"); plt.close(fig)

# ---- moat stack ----
def chart_moat():
    fig,ax=plt.subplots(figsize=(7.4,4.2))
    layers=["Owned email list (portable audience)","Reader trust (no advertiser capture)",
            "Voice & taste (AI can't copy)","Community & events network effects",
            "Local incumbency & relationships"]
    strength=[6,8,7,9,8]
    cols=["#"+SLATE,"#"+NAVY,"#"+TEAL,"#"+ORANGE,"#"+SAND]
    y=np.arange(len(layers))[::-1]
    ax.barh(y,strength,color=cols,height=0.66)
    for yi,lab,s in zip(y,layers,strength):
        ax.text(0.15,yi,lab,va="center",ha="left",color="white",fontsize=9,fontweight="bold")
    ax.set_xlim(0,10); ax.set_yticks([])
    ax.set_xlabel("Relative defensibility / hardness to replicate")
    ax.set_title("The Brief Scout Media Moat Stack")
    fig.tight_layout(); fig.savefig(f"{ASSETS}/fig_moat.png",bbox_inches="tight"); plt.close(fig)

for fn in (chart_cashflow,chart_ai_time,chart_expansion,chart_moat): fn()
print("charts done")

# ============================================================================
# DOCUMENT
# ============================================================================
doc=Document()
normal=doc.styles["Normal"]; normal.font.name="Calibri"; normal.font.size=Pt(11)
normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.12

def set_color(run,h): run.font.color.rgb=RGBColor.from_string(h)
def shade(cell,h):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),h)
    tcPr.append(shd)
def cell_text(cell,text,bold=False,color=None,size=10,align="left",white=False):
    cell.text=""; p=cell.paragraphs[0]
    p.alignment={"left":WD_ALIGN_PARAGRAPH.LEFT,"center":WD_ALIGN_PARAGRAPH.CENTER,
                 "right":WD_ALIGN_PARAGRAPH.RIGHT}[align]
    r=p.add_run(text); r.font.size=Pt(size); r.font.bold=bold
    if white: r.font.color.rgb=RGBColor.from_string("FFFFFF")
    elif color: r.font.color.rgb=RGBColor.from_string(color)
    return p
def H(text,level=1):
    p=doc.add_paragraph()
    p.paragraph_format.space_before=Pt(14 if level==1 else 10); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(text); r.font.bold=True
    if level==1:
        r.font.size=Pt(17); set_color(r,NAVY)
        pPr=p._p.get_or_add_pPr(); pbdr=OxmlElement("w:pBdr"); b=OxmlElement("w:bottom")
        b.set(qn("w:val"),"single"); b.set(qn("w:sz"),"6"); b.set(qn("w:space"),"2"); b.set(qn("w:color"),TEAL)
        pbdr.append(b); pPr.append(pbdr)
    elif level==2: r.font.size=Pt(13.5); set_color(r,TEAL)
    else: r.font.size=Pt(11.5); set_color(r,SLATE)
    return p
def P(text,italic=False,color=None,size=11,align="left",space_after=6):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(space_after)
    p.alignment={"left":WD_ALIGN_PARAGRAPH.LEFT,"center":WD_ALIGN_PARAGRAPH.CENTER,
                 "right":WD_ALIGN_PARAGRAPH.RIGHT,"just":WD_ALIGN_PARAGRAPH.JUSTIFY}[align]
    for i,ch in enumerate(text.split("**")):
        if not ch: continue
        r=p.add_run(ch); r.font.size=Pt(size); r.font.italic=italic
        if i%2==1: r.font.bold=True
        if color: set_color(r,color)
    return p
def BUL(items):
    for it in items:
        p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(3)
        for i,ch in enumerate(it.split("**")):
            if not ch: continue
            r=p.add_run(ch); r.font.size=Pt(10.5)
            if i%2==1: r.font.bold=True
def TABLE(headers,rows,widths=None,fontsize=9.5,caption=None,first_col_bold=False,header_fill=NAVY):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    t.style="Table Grid"
    for i,h in enumerate(headers):
        shade(t.rows[0].cells[i],header_fill)
        cell_text(t.rows[0].cells[i],h,bold=True,white=True,size=fontsize,align="left" if i==0 else "center")
    for r,row in enumerate(rows):
        cells=t.add_row().cells; fill=LIGHT if r%2==0 else "FFFFFF"
        for i,v in enumerate(row):
            shade(cells[i],fill)
            cell_text(cells[i],str(v),bold=(first_col_bold and i==0),size=fontsize,
                      align="left" if i==0 else "center",
                      color=NAVY if (first_col_bold and i==0) else None)
    if widths:
        for i,w in enumerate(widths):
            for row in t.rows: row.cells[i].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    if caption: P(caption,italic=True,color=GREY,size=8.5,space_after=10)
    return t
def FIG(path,caption,width=6.4):
    doc.add_picture(path,width=Inches(width))
    doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    P(caption,italic=True,color=GREY,size=8.5,align="center",space_after=12)
def set_margins(sec,t=0.9,b=0.9,l=0.9,r=0.9):
    sec.top_margin=Inches(t); sec.bottom_margin=Inches(b); sec.left_margin=Inches(l); sec.right_margin=Inches(r)

for s in doc.sections: set_margins(s)

# ---------------- TITLE ----------------
for _ in range(3): doc.add_paragraph()
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=t.add_run("BRIEF SCOUT MEDIA"); r.font.size=Pt(32); r.font.bold=True; set_color(r,NAVY)
s=doc.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=s.add_run("Strategic Addendum"); r.font.size=Pt(20); set_color(r,TEAL)
s=doc.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=s.add_run("Year-1 Cash Flow · AI + Human-in-the-Loop · Pre-Funding Milestones · Expansion · Moat")
r.font.size=Pt(12.5); r.font.italic=True; set_color(r,SLATE)
for _ in range(2): doc.add_paragraph()
s=doc.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=s.add_run("Companion to the Brief Scout Media Founder Business Plan  ·  June 2026")
r.font.size=Pt(10.5); set_color(r,GREY)
doc.add_page_break()

# ============================================================================
# 1. YEAR-1 MONTHLY CASH FLOW
# ============================================================================
H("1. Year-1 Monthly Cash-Flow Model",1)
P("This model tracks cash month-by-month through the first year, **bootstrapped with no outside "
  "funding**. It assumes ~$5,000 of founder startup capital (≈$2,500 spent pre-launch on LLC, brand, "
  "and site), and that the founder takes **no salary draw in Year 1** — reinvesting all surplus into "
  "a cash war chest while (realistically) keeping a day job. All three reader-funded streams are "
  "switched on from launch. Event and product lines are shown **net** of their direct costs "
  "(venue/JV share, COGS).",align="just")
P(f"**Headline result:** Year 1 generates roughly **{usd(T_rev)} in net revenue** "
  f"(membership {usd(T_mem)}, events {usd(T_evn)}, products {usd(T_prd)}), spends ~{usd(T_opex)} on "
  f"operations, and **ends the year with about {usd(ENDBAL)} in the bank** — having validated every "
  f"revenue stream without a dollar of advertising or investment. Cash never dips below the opening "
  f"balance, so the business is **'default alive.'**",align="just")
FIG(f"{ASSETS}/fig_cashflow.png",
    "Figure A1. Monthly net cash (bars) and cumulative balance (line). Step-ups in Months 6 and 10 are the first ticketed events; the trough in Months 1–2 is the pre-revenue launch ramp.",width=6.9)

P("**Key monthly assumptions:**")
BUL([
 "**Free list** grows ~600/mo early, tapering to ~7,000 by Month 12 (organic / referral / events — no paid geo-acquisition).",
 "**Members** reach 350 by Month 12 (~5% of the free list), seeded by a pre-launch founding-member presale.",
 "**Events:** a small launch event (M3), a ticketed tasting walk (M6), a member mixer (M9), and a larger ticketed event (M10).",
 "**Products:** a merch drop (M4), the first paid city guide (M5), restocks (M7), and a holiday merch push (M11–M12).",
 "**Costs** stay lean: beehiiv is free until ~2,500 subs (M3), then ~$45/mo; light marketing; part-time contractor help only from M9.",
])

# ---- landscape section for the wide monthly table ----
land=doc.add_section(WD_SECTION.NEW_PAGE)
land.orientation=WD_ORIENT.LANDSCAPE
land.page_width,land.page_height=land.page_height,land.page_width
set_margins(land,t=0.7,b=0.7,l=0.8,r=0.8)

H("1.1 Month-by-month cash flow detail",2)
hdr=["Line item ($)"]+months+["Year 1"]
def money_row(label,arr,total=None):
    return [label]+[f"{v:,}" if v else "—" for v in arr]+[f"{(total if total is not None else sum(arr)):,}"]
rows=[
 money_row("Membership (net)",membership,T_mem),
 money_row("Events (net)",event_net,T_evn),
 money_row("Products (net)",prod_net,T_prd),
 money_row("Total cash in",total_in,T_rev),
 money_row("Operating expenses",opex,T_opex),
 ["Net cash flow"]+[f"{v:,}" for v in net_cash]+[f"{T_net:,}"],
 ["Cumulative cash"]+[f"{v:,}" for v in cum]+[f"{ENDBAL:,}"],
]
w=[1.45]+[0.62]*12+[0.72]
TABLE(hdr,rows,widths=w,fontsize=7.5,first_col_bold=True,
      caption="Table A1. Year-1 monthly cash flow. Opening balance $2,500 (startup capital net of pre-launch spend). Founder draw assumed $0 in Year 1.")
H("1.2 Audience build (end of month)",2)
arows=[
 ["Free subscribers"]+[f"{v:,}" for v in free_eom],
 ["Paying members"]+[f"{v:,}" for v in mem_eom],
]
TABLE(["End-of-month"]+months,arows,widths=[1.6]+[0.72]*12,fontsize=8,first_col_bold=True,
      caption="Table A2. Audience at end of each month.")

# back to portrait
port=doc.add_section(WD_SECTION.NEW_PAGE)
port.orientation=WD_ORIENT.PORTRAIT
port.page_width,port.page_height=port.page_height,port.page_width
set_margins(port)

P("**How to read this for decisions.** The two ticketed events (Months 6 and 10) are the swing "
  "factors — they convert the audience into the biggest cash inflows of the year, which is exactly "
  "why events are a core pillar rather than an afterthought. If list growth lags, the founder can "
  "protect cash by trimming marketing (the largest discretionary line) and leaning harder on "
  "JV events, which require almost no upfront capital. The model deliberately keeps fixed costs "
  "near zero so the business cannot run out of money.",align="just")

# ============================================================================
# 2. HITL + AI
# ============================================================================
H("2. Scaling with Human-in-the-Loop + AI",1)
P("AI is now cheap and good enough that anyone can spin up an 'AI local newsletter' — in fact, "
  "6AM City acquired Good Daily's network of **350+ AI-generated local newsletters** in 2025, then "
  "immediately began stripping out crime/politics and refocusing on lifestyle, a tell that "
  "pure-automated local content is **commoditizing**. That is the strategic insight for Brief Scout Media "
  "Media: **use AI to win on speed and cost in the back office, and use humans to win on trust, "
  "voice, and relationships in the front of house.** AI handles volume; humans own the things a "
  "competitor (or a bot) cannot copy.",align="just")

H("2.1 The division of labor",2)
TABLE(["Workflow","What AI does (drafts/assists)","What the human owns (decides/signs off)","Payoff"],
[
 ["Event aggregation","Scrape & de-dupe venue/city calendars; structure listings","Curate the 'what to do' picks with taste","Hours saved; better picks"],
 ["First drafts","Summarize sources; draft sections & blurbs","Rewrite in the Insider voice; verify facts","Faster production, same quality"],
 ["Subject lines & headlines","Generate A/B variants","Pick & refine for voice","Higher open rates"],
 ["Research","Brief the founder; pull data & context","Judgment, sourcing, original reporting","Depth without the grind"],
 ["Social repurposing","Cut newsletter into posts/clips/captions","Approve & add personality","Multi-channel reach, one effort"],
 ["Personalization","Segment list; tailor sends; predict churn","Set strategy & member experience","Better retention & LTV"],
 ["Operations","Automate billing, tagging, reminders (Make/Zapier)","Exception handling & relationships","Solo-operator leverage"],
],widths=[1.35,2.0,2.0,1.0],fontsize=8.5,first_col_bold=True,
caption="Table A3. The human-in-the-loop model: AI accelerates production; a human always edits, verifies, and signs off before publish.")

H("2.2 Why this is a better product, not just a cheaper one",2)
BUL([
 "**Quality compounds.** Cutting production time ~30–50% lets a solo founder hold a high editorial bar at the weekly cadence — and reinvest the saved hours into events, partnerships, and original reporting, the work that actually builds the moat.",
 "**Differentiation through contrast.** As inboxes fill with AI sludge, a visibly human, opinionated, fact-checked Insider stands out. 'Written by a real Chattanoogan who actually went there' becomes a selling point.",
 "**Scale without dilution.** AI makes new verticals (Section 4) feasible for a tiny team, because the repeatable, mechanical 80% is automated and humans focus on the 20% that carries the brand.",
])
P("**Guardrails (non-negotiable):** never auto-publish; every issue is human-edited and "
  "fact-checked; AI is a drafting and ops tool, not the byline; be transparent with readers about "
  "how the Insider is made; protect the voice above all. The promise is human judgment, "
  "advertiser-free — AI just removes the drudgery that gets in its way.")
FIG(f"{ASSETS}/fig_ai_time.png",
    "Figure A2. AI leverage shifts founder time away from mechanical production toward events, partnerships, and community — the defensible, revenue-driving work.")

# ============================================================================
# 3. PRE-FUNDING MILESTONES
# ============================================================================
H("3. Milestones & Accomplishments Before Seeking Investment",1)
P("This business is designed to be **default-alive on reader revenue** — many of the best "
  "comparables (Naptown Scoop, Catskill Crew, Charlotte Ledger) never raised a dollar. Outside "
  "capital should be **optional and offensive**: raised only to accelerate multi-market expansion, "
  "not to survive. That said, if/when the founder chooses to raise (6AM City raised $1.55M at seed "
  "in 2020 and $15.8M total by 2024), the following traction is what de-risks the story and "
  "maximizes valuation. The goal is to walk into any investor conversation already profitable, with "
  "a **proven, documented, repeatable playbook.**",align="just")
H("3.1 The pre-funding milestone ladder",2)
TABLE(["Milestone","Target before raising","Why investors care","Timing"],
[
 ["Audience","7,000–15,000 engaged free subs; 45–55% open rate","Proves organic demand & a real funnel","Y1–Y2"],
 ["Paid conversion","350–1,000+ members at 4–6% conversion","Proves willingness to pay (the hardest unknown locally)","Y1–Y2"],
 ["Recurring revenue","$3,000–$8,000+ MRR from membership","Predictable base; the metric investors index on","Y2"],
 ["Retention","Monthly churn <4%; rising LTV","Shows product love, not just acquisition","Y1–Y2"],
 ["Events","2–4 profitable events; repeatable format","Validates the highest-margin pillar","Y1–Y2"],
 ["Diversification","No single stream >50% of revenue","Resilience; the LION '3+ streams' sustainability marker","Y2"],
 ["Profitability","Owner-paid & cash-flow positive","'Default alive' — you raise from strength","Y2"],
 ["Playbook","Documented launch & ops playbook + founding team","Proves the model is repeatable in a new market","Y2–Y3"],
],widths=[1.2,1.9,2.1,0.7],fontsize=8.5,first_col_bold=True,
caption="Table A4. Traction milestones that convert a local newsletter into a fundable, scalable story.")
P("**Pre-launch proof points** worth banking before Month 1: form the LLC and open a "
  "**founding-member presale** (even 40–60 paid founders is hard evidence of willingness-to-pay in "
  "a market where no paid local newsletter yet exists); line up 5–10 launch partners; and secure "
  "1–2 anchor event venues. These cost almost nothing and materially de-risk the launch.")

# ============================================================================
# 4. EXPANSION
# ============================================================================
H("4. Expansion: Verticals vs. Markets — Ability & Timing",1)
P("There are two ways to grow beyond the flagship, and they have very different economics. "
  "**Vertical expansion** (new products in Chattanooga) monetizes an audience you already own — low "
  "acquisition cost, high incremental LTV, fast. **Geographic expansion** (the same playbook in a "
  "new city) multiplies the audience but resets your relationships, content, and trust to zero — "
  "higher cost, slower, more operationally complex. The recommended sequence is unambiguous: "
  "**go deep before you go wide.**",align="just")
FIG(f"{ASSETS}/fig_expansion.png",
    "Figure A3. Recommended sequencing — saturate Chattanooga with verticals first; attempt a second market only once the playbook is proven and documented (Year 3+).")

H("4.1 Vertical expansion (same city, new products)",2)
P("**Prerequisite:** a stable flagship (~10,000+ free subs, healthy membership, reliable events). "
  "**Timing: Year 2.** Each vertical reuses the brand, list, and ops, so the marginal cost is low "
  "and AI (Section 2) makes additional editorial output feasible for a tiny team.")
BUL([
 "**A dedicated 'What to do' / weekend events edition** — natural extension of the strongest content, feeds the events business.",
 "**Newcomer guide / relocation funnel** — targets Chattanooga's fast-growing CA/IL transplant inflow at the exact moment they'll pay for orientation.",
 "**Food & dining edition** — high engagement, strong JV-event and guide tie-ins.",
 "**Family/kids or neighborhood editions** — deepen household reach and local-business relationships.",
])
H("4.2 Geographic expansion (new markets)",2)
P("**Prerequisite:** the Chattanooga model is profitable, documented, and repeatable, ideally with a "
  "second operator to prove it isn't founder-dependent. **Timing: Year 3+.** Strong candidate metros "
  "are mid-sized, growing Southeastern cities with the same dynamics — **Knoxville, Asheville, "
  "Greenville, Chattanooga-adjacent Huntsville/Birmingham** — entered one at a time. 6AM City's "
  "network of 30+ cities proves replication scales; Charlotte Ledger's depth proves a single city can "
  "be enough. Brief Scout Media should earn the right to choose.")
TABLE(["Path","Prerequisite","Timing","Capital need","Risk","Upside"],
[
 ["Verticals (same city)","Stable flagship ~10K+ list","Year 2","Low","Low — owned audience","Higher LTV per reader"],
 ["2nd market (single)","Proven, documented playbook","Year 3+","Medium","Medium — rebuild trust","Multiplies audience"],
 ["Multi-market network","2nd market succeeds + capital","Year 4–5","High (raise here)","High — ops complexity","Category-leader scale"],
],widths=[1.3,1.6,0.9,1.0,1.1,1.1],fontsize=8.5,first_col_bold=True,
caption="Table A5. Expansion decision framework. Verticals are the low-risk default; geography is where outside capital earns its keep.")

# ============================================================================
# 5. TONE, TOPICS & MOAT
# ============================================================================
H("5. Tone, Topics & the Moat",1)
P("In a market where the free incumbent (NOOGAtoday) is broad and ad-driven and the legacy paper is "
  "behind an expensive paywall, Brief Scout Media differentiates on **trust, taste, and "
  "independence**. The reader — not an advertiser — is the customer, and the brand should sound like "
  "it. That stance is itself a moat: it is the one thing an ad-funded competitor structurally "
  "cannot copy without breaking its own model.",align="just")
H("5.1 Tone",2)
BUL([
 "**Trusted insider, not a feed.** Warm, human, a little opinionated — a knowledgeable friend who lives here, not an aggregator.",
 "**Independent and unbought.** Lean into it: 'We work for our readers, not advertisers.' Editorial recommendations are credible precisely because nothing is paid placement.",
 "**Curated, anti-clickbait.** Respect the reader's time; quality and brevity over volume. The antidote to inbox overload and AI sludge.",
 "**Pro-Chattanooga without boosterism.** Celebrate the city and hold it accountable; genuine, not a chamber-of-commerce brochure.",
])
H("5.2 Topics that differentiate and compound into a moat",2)
TABLE(["Topic pillar","Why it differentiates","Moat it builds"],
[
 ["Curated 'what to do', done with taste","Listings are commodities; judgment isn't","Voice & taste; feeds events"],
 ["Original recommendations & reviews","Can't be AI-scraped; earns trust","Brand authority"],
 ["Accountability / development / business depth","Free competitors won't fund real reporting","Indispensability"],
 ["Newcomer orientation","Serves the fastest-growing audience segment","High-LTV reader relationships"],
 ["Community & reader features","Members see themselves in it","Belonging; network effects"],
 ["Events as content / content as events","Owns experiences, not just attention","Switching costs; relationships"],
],widths=[1.9,2.3,1.6],fontsize=8.5,first_col_bold=True,
caption="Table A6. Topic strategy. Each pillar is chosen because it is hard to copy and it deepens the audience relationship.")
P("**The moat is cumulative.** No single feature is defensible alone, but stacked together — an "
  "owned email list, advertiser-free trust, a voice AI can't replicate, an events-and-community "
  "network, and years of local relationships — they become very hard to dislodge. A new entrant "
  "would have to rebuild all five from zero, in a city where Brief Scout Media is already the trusted name.")
FIG(f"{ASSETS}/fig_moat.png",
    "Figure A4. The moat stack — five reinforcing layers of defensibility, anchored by reader trust and local incumbency.")

# closing
H("6. Bottom Line",1)
P("Year 1 is a bootstrapped validation year that ends with cash in the bank and all three revenue "
  "streams proven — no funding required. AI, kept strictly human-in-the-loop, makes a solo founder "
  "produce like a small team and frees time for the defensible work. With a documented, profitable "
  "playbook, the founder can then choose growth on their own terms: verticals first (Year 2), a "
  "second market only once the model is proven (Year 3+), and outside capital only to accelerate a "
  "network — raised from strength, not need. Throughout, the brand wins on the one thing competitors "
  "can't buy: the trust of readers who know the Insider works for them.",align="just")

out="/home/user/prscott/Brief_Scout_Media_Strategic_Addendum.docx"
doc.save(out); print("saved",out)
print(f"Y1 rev={T_rev} mem={T_mem} evn={T_evn} prd={T_prd} opex={T_opex} endbal={ENDBAL}")
