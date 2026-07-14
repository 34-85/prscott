#!/usr/bin/env python3
"""Insider Insights — the permission-based local research practice (spec / one-pager+)."""
import os, matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY="1B3A4B"; TEAL="2A9D8F"; ORANGE="E76F51"; SAND="E9C46A"; SLATE="415A6B"
LIGHT="EEF3F4"; GREY="6B7B85"; GOLD="B8860B"; GREEN="2E7D52"
ASSETS="/home/user/prscott/assets"; os.makedirs(ASSETS,exist_ok=True)
plt.rcParams.update({"font.family":"DejaVu Sans","figure.dpi":160})

def chart_roadmap():
    fig,ax=plt.subplots(figsize=(7.6,3.9)); ax.set_xlim(0,36); ax.set_ylim(-0.5,5)
    phases=[("Phase 0 · Light reader polls (habit + content + response-rate proof)",0,36,"#"+SLATE,4),
            ("Phase 1 · Contest-based onboarding + first-party appends → early insights",3,12,"#"+NAVY,3),
            ("Phase 2 · Build DEPTH — 'Insiders' panel + methodology",9,24,"#"+TEAL,2),
            ("Phase 3 · 'The Scenic City Index' — house research / authority",18,30,"#"+SAND,1),
            ("Phase 4 · Monetize — custom & sponsored panels (LAST)",24,37,"#"+ORANGE,0)]
    for lab,s,e,c,y in phases:
        ax.barh(y,e-s,left=s,height=0.62,color=c,edgecolor="white")
        ax.text(s+0.3,y,lab,va="center",ha="left",color="white",fontsize=7.6,fontweight="bold")
    ax.axvline(24,ls="--",color="#"+ORANGE,lw=1.4)
    ax.text(24,4.7,"depth & credibility BEFORE sponsored revenue",ha="center",fontsize=8,color="#"+ORANGE,fontweight="bold")
    ax.set_yticks([]); ax.set_xticks([0,6,12,18,24,30,36])
    ax.set_xticklabels(["Launch","6mo","12mo","18mo","24mo","30mo","36mo"],fontsize=8)
    ax.set_title("Insider Insights — build depth first, monetize last",fontweight="bold",color="#"+NAVY)
    for s in ("top","right","left"): ax.spines[s].set_visible(False)
    fig.tight_layout(); fig.savefig(f"{ASSETS}/fig_insights.png",bbox_inches="tight"); plt.close(fig)
chart_roadmap(); print("chart done")

PHASES=[
 ("Phase 0 — Listen","Now (pre-10K)","Light in-newsletter 'Insider asks…' polls. Builds the response habit, makes great content, and proves people answer.","Internal only"),
 ("Phase 1 — Onboard","~3–12 mo","Contest-based 3–5-question micro-surveys at signup (and periodically). Collect consented, first-party profile data — age band, gender, part of town, interests, life stage. Insights form early.","Internal only"),
 ("Phase 2 — Deepen","~9–24 mo","Progressive profiling (one question an issue) grows the opt-in 'Insiders' panel and a real longitudinal dataset. Add methodology rigor (sampling, weighting).","Internal only"),
 ("Phase 3 — Authority","~18–30 mo","Publish 'The Scenic City Index' — house research that makes Brief Scout Media the cited local authority. Feeds editorial; builds the B2B brand.","Brand / PR (low/none)"),
 ("Phase 4 — Monetize","~24 mo+","Only now: custom 'Insider Insights' studies and sponsored/omnibus panels for local orgs — apolitical, aggregate-only, walled from editorial.","B2B research revenue"),
]
QMAP=[
 ("Age band","Tone & reading level; product sizing/merch; event design (nightlife vs. daytime)","Demographic cross-tabs for every study"),
 ("Gender","Balanced content & representation; product/merch mix; event programming","Representative weighting"),
 ("Part of town / area","Local relevance; where to host events; which neighborhoods to feature next","Geographic segmentation (a real local edge)"),
 ("Interests (food/makers/outdoors/arts/family/getaways/civic)","Which sections, products, and events to expand — demand, not guesswork","Interest-based panels for clients"),
 ("Life stage / household","Family vs. date-night vs. retiree programming; product fit","Lifestyle segmentation"),
 ("New to Chattanooga? (under 2 yrs)","Powers the 'Newcomer' Lens + relocation funnel; serves the migration wave","High-value transplant cohort"),
]
GUARD=[
 "**No PII — ever.** Clients receive aggregate findings and cross-tabs, never names, emails, or individual records.",
 "**First-party & consented.** Self-reported at opt-in, built up by progressive profiling. We do NOT silently buy third-party demographic appends and attach them — that's a trust and legal liability that contradicts the brand.",
 "**Opt-in panel.** 'Insiders' is something readers choose to join, incentivized by a contest (prizes are gift cards to the local businesses we cover — the reward itself supports a maker).",
 "**Frequency-capped.** At most one profiling question per issue and one deeper survey a month. Belonging, not extraction — over-surveying churns the list.",
 "**Walled from editorial.** Research never dictates coverage; we decline political, religious, or advocacy work. Same bright line as the partner program.",
 "**Transparent & secure.** Clear privacy policy, easy opt-out honored immediately, sensible data security.",
]
RATE=[
 ("Omnibus question","Add one question to the monthly panel; client gets the cross-tabs","~$250–$750 / question"),
 ("Focused custom study","Concept/pricing/satisfaction test, n≈300–800, report + readout","~$2.5K–$15K"),
 ("Major study","Multi-segment or longitudinal work","~$15K–$50K+"),
 ("Scenic City Index","Subscription / underwriting of the flagship report (clearly labeled)","Subscription or sponsor"),
]
CONTRIB=[
 ("Year 1","~$0 external","The payoff is internal: onboarding data makes content, products & events land — higher conversion & retention from day one."),
 ("Year 2","~$5K–$20K","First Index underwriting + a few omnibus/custom pilots as depth and credibility build."),
 ("Year 3","~$25K–$60K","Custom studies + Index subscriptions mature — plus the larger, unbookable payoff: authority, PR, and product-market fit."),
]

doc=Document()
for s in doc.sections:
    s.top_margin=Inches(0.9); s.bottom_margin=Inches(0.9); s.left_margin=Inches(0.9); s.right_margin=Inches(0.9)
nm=doc.styles["Normal"]; nm.font.name="Calibri"; nm.font.size=Pt(11); nm.paragraph_format.space_after=Pt(6); nm.paragraph_format.line_spacing=1.1
def sca(r,h): r.font.color.rgb=RGBColor.from_string(h)
def shade(c,h):
    p=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd")
    s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); p.append(s)
def ctext(c,t,bold=False,color=None,size=9.2,white=False,align="left"):
    c.text=""; p=c.paragraphs[0]
    p.alignment={"left":WD_ALIGN_PARAGRAPH.LEFT,"center":WD_ALIGN_PARAGRAPH.CENTER}[align]
    r=p.add_run(t); r.font.size=Pt(size); r.font.bold=bold
    if white: r.font.color.rgb=RGBColor.from_string("FFFFFF")
    elif color: r.font.color.rgb=RGBColor.from_string(color)
def H(t,level=1,color=None):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(12 if level==1 else 8); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(t); r.font.bold=True
    if level==1:
        r.font.size=Pt(15); sca(r,color or NAVY)
        pPr=p._p.get_or_add_pPr(); b=OxmlElement("w:pBdr"); bt=OxmlElement("w:bottom")
        bt.set(qn("w:val"),"single"); bt.set(qn("w:sz"),"6"); bt.set(qn("w:space"),"2"); bt.set(qn("w:color"),TEAL)
        b.append(bt); pPr.append(b)
    else: r.font.size=Pt(12.5); sca(r,color or TEAL)
def P(t,italic=False,color=None,size=11):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(5)
    for i,ch in enumerate(t.split("**")):
        if not ch: continue
        r=p.add_run(ch); r.font.size=Pt(size); r.font.italic=italic
        if i%2==1: r.font.bold=True
        if color: sca(r,color)
def BUL(items,size=10):
    for it in items:
        p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(3)
        for i,ch in enumerate(it.split("**")):
            if not ch: continue
            r=p.add_run(ch); r.font.size=Pt(size)
            if i%2==1: r.font.bold=True
def FIG(path,cap,width=6.6):
    doc.add_picture(path,width=Inches(width)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    P(cap,italic=True,color=GREY,size=8.5); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
def TABLE(headers,rows,widths,fontsize=9,caption=None,first_bold=True):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style="Table Grid"
    for i,h in enumerate(headers):
        shade(t.rows[0].cells[i],NAVY); ctext(t.rows[0].cells[i],h,bold=True,white=True,size=fontsize)
    for r,row in enumerate(rows):
        cells=t.add_row().cells; base=LIGHT if r%2==0 else "FFFFFF"
        for i,v in enumerate(row):
            shade(cells[i],base)
            ctext(cells[i],str(v),bold=(first_bold and i==0),size=fontsize,color=NAVY if (first_bold and i==0) else None)
    for i,w in enumerate(widths):
        for row in t.rows: row.cells[i].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    if caption: P(caption,italic=True,color=GREY,size=8.3)

# title
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("SCOUT INSIGHTS"); r.font.size=Pt(24); r.font.bold=True; sca(r,NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Brief Scout Media's permission-based local research practice"); r.font.size=Pt(13); sca(r,TEAL)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("No PII. Opt-in. Aggregate-only. The reader stays the customer — never the product."); r.font.size=Pt(10.5); r.font.italic=True; sca(r,SLATE)

H("The idea in one line",1)
P("Turn a trusted, engaged, *local* audience into the region's best source of consumer and civic insight — sold as **aggregate findings, never identities** — building a high-margin, apolitical B2B revenue line that *deepens* trust instead of spending it, and makes Brief Scout Media the cited local authority.")
P("**Sequencing is the whole strategy: build depth before we monetize.** We start with contest-based onboarding questions that form insights early (and immediately sharpen our own content, products, and events), deepen the dataset over time, establish credibility with house research, and only *then* open sponsored/custom panels.")
FIG(f"{ASSETS}/fig_insights.png","Figure. The five phases — depth and credibility come before any sponsored dollar.")

H("The phased roadmap",1)
TABLE(["Phase","When","What happens","Revenue"],PHASES,widths=[1.5,1.0,3.4,1.3],fontsize=8.8)

H("Phase 1 — the contest-based onboarding mechanic",1)
P("New (and existing) subscribers answer **3–5 quick questions** to enter a contest — the prize is a gift card to a local business we cover, so the reward itself champions a maker. The questions are **first-party and consented**; some are simple demographics (age, gender) so we build the *right* content, products, and events. Insights start forming immediately.")
TABLE(["Onboarding question","What it powers right away","Later research value"],QMAP,widths=[1.7,3.1,2.0],fontsize=8.8,
      caption="Every question earns its place twice — it improves the product now and becomes a research dimension later.")

H("The guardrails (non-negotiable)",1)
BUL(GUARD)

H("The immediate payoff — before a single external dollar",1)
P("Even at $0 of research revenue, this pays for itself. Knowing your readers' age, neighborhood, life stage, and interests lets you **segment sends, tailor sections, size merch, and design events people actually show up to** — lifting free→paid conversion and retention across the core business. The data is a *product-development engine* first and a revenue line second.")

H("The external products (Phase 3–4)",1)
P("**The Scenic City Index** (house research) establishes authority and feeds editorial — the brand-builder. **Insider Insights** (custom + omnibus) is the demand-driven revenue, sold only after depth and credibility exist.")
TABLE(["Product","What it is","Illustrative price"],RATE,widths=[1.6,3.4,1.8],fontsize=9,
      caption="Illustrative; high margin because the audience already exists. Apolitical, aggregate-only, walled from editorial.")

H("Fit in the model & a light 3-year contribution",1)
P("Insider Insights is a **fourth revenue line** alongside membership, products, and events — and a walled B2B practice, like the partner program. It's deliberately back-loaded: the early value is internal (better content/products/events), the later value is revenue *and* the unbookable payoff of authority and brand.")
TABLE(["Year","External revenue (illustrative)","The bigger payoff"],CONTRIB,widths=[0.9,1.9,4.0],fontsize=9)

H("Bottom line",1)
P("A permission-based, no-PII insights practice is the most on-brand non-advertising revenue Brief Scout Media can build: it monetizes trust without selling the reader, sharpens the core product from week one, and — through 'The Scenic City Index' — makes Brief Scout Media the authority on what Chattanooga thinks. Build the depth first; the panels will be worth far more for the wait.")

out="/home/user/prscott/Insider_Insights_Spec.docx"
doc.save(out); print("saved",out)
