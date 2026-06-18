#!/usr/bin/env python3
"""Born Here — Iconic Brands from the Scenic City. Compiled & de-duped from 4 Scout reports."""
import os, csv, matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY="1B3A4B"; TEAL="2A9D8F"; ORANGE="E76F51"; SAND="E9C46A"; SLATE="415A6B"
LIGHT="EEF3F4"; GREY="6B7B85"; GOLD="B8860B"; GREEN="2E7D52"
ASSETS="/home/user/prscott/assets"; os.makedirs(ASSETS,exist_ok=True)
plt.rcParams.update({"font.family":"DejaVu Sans","figure.dpi":160})

def chart_timeline():
    marquee=[(1879,"Chattanooga Medicine\n(Chattem)"),(1887,"Provident\n→ Unum"),(1899,"Coca-Cola\nbottling born"),
     (1902,"MoonPie /\nChatt. Bakery"),(1932,"Krystal"),(1934,"Little Debbie\n(McKee)"),(1945,"BlueCross\nBlueShield TN"),
     (1967,"Shaw\n(Dalton)"),(1972,"Astec"),(1985,"U.S. Xpress"),(1990,"Miller Ind.\n(tow trucks)"),
     (2002,"Access America"),(2011,"Chattanooga\nWhiskey"),(2016,"FreightWaves")]
    fig,ax=plt.subplots(figsize=(9.6,4.0)); ax.set_xlim(1865,2025); ax.set_ylim(-1.6,1.6)
    ax.axhline(0,color="#"+SLATE,lw=2,zorder=1)
    for i,(yr,name) in enumerate(marquee):
        up=1 if i%2==0 else -1; y=up*(0.55 if i%4<2 else 1.05)
        c=["#"+NAVY,"#"+TEAL,"#"+ORANGE,"#"+SAND][i%4]
        ax.plot([yr,yr],[0,y],color=c,lw=1.3,zorder=2)
        ax.scatter([yr],[0],s=42,color=c,zorder=3,edgecolor="white",linewidth=1)
        ax.text(yr,y+(0.12*up),f"{name}\n{yr}",ha="center",va="bottom" if up>0 else "top",fontsize=7,color="#"+NAVY,fontweight="bold")
    ax.set_yticks([]); ax.set_xticks([1880,1900,1920,1940,1960,1980,2000,2020])
    for s in ("top","right","left"): ax.spines[s].set_visible(False)
    ax.set_title("A brand town for 150 years — marquee names founded in the Scenic City region",fontweight="bold",color="#"+NAVY)
    fig.tight_layout(); fig.savefig(f"{ASSETS}/fig_bornhere.png",bbox_inches="tight"); plt.close(fig)
chart_timeline(); print("chart done")

# (brand, iconic_for, connection, surprising_fact, confidence)
STILL=[
 ("MoonPie / Chattanooga Bakery","The marshmallow-graham snack","Founded & HQ Chattanooga, 1902 (MoonPie 1917)","Born when a salesman asked KY coal miners what they wanted and one framed the moon: 'as big as that.' ~1M made/day, still independent.","High"),
 ("Little Debbie / McKee Foods","Snack cakes (Oatmeal Creme Pies)","Founded Chattanooga 1934 → HQ Collegedale 1957","Named for a real granddaughter, age 4 at launch (1960); still family-owned. Also owns Drake's, Sunbelt.","High"),
 ("Double Cola (& Ski)","Regional cola + cult citrus soda 'Ski'","Founded & HQ Chattanooga, 1922/1933","Named 'Double' because the bottle was twice rivals'; pioneered the soda industry's first Applied Color Label.","High"),
 ("Unum (Provident Life)","Disability/benefits insurance leader","Founded Chattanooga 1887; HQ here (NYSE: UNM)","Built insuring railroad workers; offered the first disability benefits (1939); marked 175 years in 2023.","High"),
 ("BlueCross BlueShield of TN","Tennessee's largest health insurer","Founded & HQ Chattanooga, 1945","Started by a newspaper publisher (Roy McDonald); its Cameron Hill campus is the largest LEED-Gold campus in TN.","High"),
 ("Astec Industries","Asphalt/road-building equipment (NASDAQ: ASTE)","Founded & HQ Chattanooga, 1972","Funded by founder Dr. J. Don Brock's royalties from a carpet-dryer invention — now ~$1.5B in sales.","High"),
 ("Miller Industries","World's largest tow-truck/wrecker maker (NYSE: MLR)","Founded & HQ Ooltewah, TN, 1990","A huge share of the planet's tow trucks (Century, Vulcan, Chevron, Holmes) are built in a Chattanooga suburb.","High"),
 ("FreightWaves","'Bloomberg of Freight' data/news","Founded & HQ Chattanooga, 2016","Founder Craig Fuller is the son of U.S. Xpress co-founder Max Fuller — the family that seeded Chattanooga freight-tech.","High"),
 ("Covenant Logistics","Public truckload/logistics carrier (NYSE: CVLG)","Founded & HQ Chattanooga, 1986","Note: NYSE, not NASDAQ. Grew through acquisitions; renamed Covenant Logistics Group in 2020.","High"),
 ("Steam Logistics","Fast-growing international 3PL","Founded & HQ Chattanooga, 2012","Spun out of Access America; among the fastest-growing 3PLs in the U.S.","High"),
 ("Reliance Partners","Freight/trucking insurance agency","Founded & HQ Chattanooga, 2009","The only insurance agency named to the FreightTech 100 every year since its inception.","High"),
 ("Bellhop","Tech-enabled moving service","Founded 2011; HQ Chattanooga (via Lamp Post)","First gig aimed for ~30 dorm moves at Auburn orientation and did 325; now in 65+ cities.","High"),
 ("TransCard","Embedded/B2B payments fintech","Founded & HQ Chattanooga, 2008","Same name lineage as the fleet-payments firm Craig Fuller sold to U.S. Bank — part of the Fuller ecosystem.","Med"),
 ("Dynamo Ventures","Supply-chain / logistics-tech seed VC","Founded & HQ Chattanooga, 2016","Born from the Access America windfall; backed unicorns Sennder and Stord.","High"),
 ("EPB — 'Gig City' fiber","First U.S. community-wide gigabit network","Municipal utility, Chattanooga; fiber 2010","Beat Google to gigabit (and to 10-Gig in 2015, 25-Gig in 2022) — built over a Comcast lawsuit.","High"),
 ("Chattanooga Whiskey","Tennessee high-malt whiskey","Founded & HQ Chattanooga, 2011","Its 'Vote Whiskey' campaign changed TN law (2013) to legalize distilling in Hamilton County after ~100 years.","High"),
 ("Southern Champion Tray","Paperboard food trays / packaging","Founded & HQ Chattanooga, 1927","A ~century-old family business quietly makes much of the bakery/to-go packaging behind American food.","Med"),
 ("Rock/Creek Outfitters","Specialty outdoor-gear retail","Founded Chattanooga 1987; locally rooted","Anchored the city's outdoor identity; ownership has since changed hands (verify current owner).","Med"),
]
REGION_CARPET=[
 ("Shaw Industries","World's largest carpet/flooring maker","Founded & HQ Dalton, GA, 1967","Owned by Warren Buffett's Berkshire Hathaway since 2001 — the world's biggest carpet company sits in a small GA town.","High"),
 ("Mohawk Industries","World's largest flooring manufacturer (NYSE: MHK)","HQ Calhoun, GA (modern entity 1988)","The two largest flooring companies on Earth — Shaw and Mohawk — are headquartered ~30 miles apart.","High"),
 ("Engineered Floors","Major U.S. carpet producer","Founded & HQ Dalton, GA, 2009","Founded by Bob Shaw — who'd already built Shaw Industries — making him a two-time carpet titan after 'retiring.'","High"),
 ("The Dixie Group","High-end carpet (Fabrica, Masland)","Founded Chattanooga 1920 → HQ Dalton","Began as Dixie Mercerizing, treating yarn for ladies' stockings, then pivoted to luxury carpet.","High"),
 ("Mueller Co. (Water Products)","Fire hydrants, valves, water-distribution","Founded Decatur IL 1857; water-products HQ → Chattanooga 2010","A large share of America's fire hydrants now trace through Chattanooga. (HQ here now, not born here.)","High"),
]
MOVED=[
 ("Coca-Cola Bottling — the franchise","The global Coca-Cola bottling system","Chartered in Chattanooga, July 21, 1899","Two local lawyers bought near-nationwide bottling rights from Asa Candler for ~$1 — he thought bottling was a fad. The most lucrative handshake in beverage history.","High"),
 ("Chattem (Chattanooga Medicine Co.)","Gold Bond, Icy Hot, Selsun Blue, Pamprin","Founded Chattanooga 1879; acquired by Sanofi 2010 (now Opella)","Gold Bond & Icy Hot are STILL made in Chattanooga (Opella's only U.S. plant). In WWII it was a top maker of K-rations.","High"),
 ("Krystal","The square steamed slider; South's oldest fast-food chain","Founded Chattanooga 1932; HQ → Atlanta 2013","2nd-oldest fast-food chain in America; named for how 'crystal clean' the first diner looked. Ch. 11 in 2020.","High"),
 ("U.S. Xpress","One of the largest U.S. truckload carriers","Founded Chattanooga 1985; acquired by Knight-Swift 2023","Co-founder Max Fuller's son Craig went on to build FreightWaves — the Fuller family is Chattanooga's freight DNA.","High"),
 ("Access America Transport","Pioneering freight brokerage — the 'Big Bang'","Founded Chattanooga 2002 → Coyote (2014) → UPS (2015) → RXO (2024)","Its ~$260M exit funded Lamp Post, Dynamo, Steam, Reliance & Bellhop — it literally manufactured the 'Silicon Valley of Trucking.'","High"),
 ("Fleetwood Coffee","The Southeast's #1 coffee for ~50 years","Founded Chattanooga 1925; revived 2013","Lost to a corporate roll-up, then brought back by the founders' grandson; now under Vienna Coffee (Maryville).","High"),
 ("Skuid","No-code enterprise app platform","Founded Chattanooga 2013; acquired by Nintex 2023","Raised $24M (2017) for 'codeless' app development.","High"),
 ("Brock Candy","Candy (chocolate-covered cherries, gummies)","Founded Chattanooga 1906; plant closed 2014","Founder W.E. Brock became a U.S. Senator; in the 1994 Brach merger the smaller Chattanooga firm briefly absorbed the Chicago icon Brach's.","High"),
 ("Buster Brown Apparel (United Hosiery Mills)","Children's hosiery & playwear","Founded Chattanooga 1904; closed 1999","The REAL local Buster Brown — apparel, not shoes. (The footwear trademark is Brown Shoe of St. Louis — a common myth.)","High"),
 ("Olan Mills","National portrait-studio chain","Founded Tuscaloosa AL 1933; HQ Chattanooga 1941; → Lifetouch 2011","Ran one of the largest portrait operations in the U.S. (and UK) from Chattanooga for 70 years. (HQ here, not born here.)","High"),
 ("WorkHound","Frontline/driver-retention software","Founded 2015 (Chattanooga + Des Moines); acquired by WorkStep 2025","Dynamo-backed; clients included Covenant.","High"),
 ("Brev.dev","AI/ML GPU dev platform","Incubated at Chattanooga's Brickyard (2021); acquired by NVIDIA 2024","Brickyard's first exit — incubated here, though its founders were not local.","High"),
]
OPS=[
 ("Volkswagen Chattanooga","VW's only U.S. assembly plant","Plant opened 2011 (~$4.2B invested); HQ Germany","The world's first LEED-Platinum auto plant; in 2026 it pivots from the ID.4 EV to the next-gen Atlas.","High"),
 ("Komatsu Chattanooga","Excavators & mining/quarry trucks","Plant opened 1985; HQ Japan","Chattanooga landed Japan's first U.S. heavy-equipment plant in 1985 — and the city's first Japanese plant.","High"),
 ("Mars Wrigley — Cleveland, TN","M&M's, Twix","Plant since ~1979; HQ McLean, VA","The Cleveland plant is the ONLY U.S. facility that makes Peanut Butter M&M's.","High"),
]
GONE=[
 ("Wheland Foundry","Auto brake castings & drums","Chattanooga; closed ~2002–03","Once produced brake parts for nearly HALF of all North American vehicles.","High"),
 ("U.S. Pipe (Chattanooga works)","Cast-iron water pipe & hydrants","Founded Chattanooga 1877; closed 2006","Ran 100+ years; fed the nation's water infrastructure when Chattanooga was the 'Dynamo of Dixie.'","High"),
 ("Chattanooga Glass Co.","Glass bottles — incl. the classic Coke bottle","Founded Chattanooga 1901; closed 1988","A chief maker of the iconic Coca-Cola bottle — tens of millions a year came out of Alton Park.","High"),
 ("Combustion Engineering (Chatt. plant)","Boilers & NUCLEAR reactor vessels","Major Chattanooga plant; defunct (→ ABB)","Reactor pressure vessels for U.S. nuclear plants were forged in Chattanooga.","High"),
 ("Beaulieu of America","Once 3rd-largest U.S. carpet maker","Founded Dalton GA 1978; Ch.11 2017","Founded by a Belgian immigrant; its assets were absorbed by rival Bob Shaw's Engineered Floors.","High"),
]
EDGE=[
 "Lodge Cast Iron — South Pittsburg, TN (1896): iconic regional cookware; include if your geography stretches west.",
 "Mayfield Dairy — Athens, TN (1910): regional edge (~50 mi), not Chattanooga proper.",
 "Cavalier Homes — Addison, AL (1984): manufactured homes; far edge of the region (now under Clayton).",
 "Roper (appliances) — LaFayette, GA plant: tangled GE/Whirlpool brand history; verify before featuring.",
 "Mountain City Flour Mill — Chattanooga (1888, defunct): mainly notable as the parent that spun off Chattanooga Bakery / MoonPie.",
]
MYTHS=[
 "**Buster Brown shoes ≠ Chattanooga.** The footwear trademark belongs to Brown Shoe Co. (St. Louis). Chattanooga's real link is the apparel/hosiery maker (United Hosiery Mills → Buster Brown Apparel). Garanimals/Garan is unrelated.",
 "**Olan Mills was founded in Tuscaloosa, AL** (1933) — it relocated its HQ here in 1941. 'HQ here,' not 'born here.'",
 "**Krystal is no longer HQ'd here** (moved to Atlanta, 2013) though it was founded in Chattanooga in 1932.",
 "**Covenant trades on the NYSE (CVLG)**, not NASDAQ.",
 "**Brev.dev was incubated here, not founded here**; **Cigna** and **First Horizon** (Memphis) are operations only — not Born-Here brands.",
 "**Mars (Cleveland) and VW (Chattanooga) are plants, not headquarters** — iconic 'made here,' not 'born here.'",
]

# ---------- CSV ----------
with open("/home/user/prscott/born_here_brands.csv","w",newline="") as f:
    w=csv.writer(f); w.writerow(["Status group","Brand","Iconic for","Connection","Surprising fact","Confidence"])
    for grp,rows in [("Still here",STILL),("Region — carpet & water",REGION_CARPET),
                     ("Founded here → moved/acquired/revived",MOVED),("Iconic operations (plant, not HQ)",OPS),
                     ("Vanished but legendary",GONE)]:
        for r in rows: w.writerow((grp,)+r)
total=sum(len(x) for x in (STILL,REGION_CARPET,MOVED,OPS,GONE))
print("CSV rows:",total)

# ---------- DOCX ----------
doc=Document(); sec=doc.sections[0]; sec.orientation=WD_ORIENT.LANDSCAPE
sec.page_width,sec.page_height=Inches(11),Inches(8.5)
for m in ("top_margin","bottom_margin","left_margin","right_margin"): setattr(sec,m,Inches(0.6))
nm=doc.styles["Normal"]; nm.font.name="Calibri"; nm.font.size=Pt(10); nm.paragraph_format.space_after=Pt(4)
def sca(r,h): r.font.color.rgb=RGBColor.from_string(h)
def shade(c,h):
    p=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd")
    s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); p.append(s)
def ctext(c,t,bold=False,color=None,size=8.4,white=False,align="left"):
    c.text=""; p=c.paragraphs[0]
    p.alignment={"left":WD_ALIGN_PARAGRAPH.LEFT,"center":WD_ALIGN_PARAGRAPH.CENTER}[align]
    r=p.add_run(t); r.font.size=Pt(size); r.font.bold=bold
    if white: r.font.color.rgb=RGBColor.from_string("FFFFFF")
    elif color: r.font.color.rgb=RGBColor.from_string(color)
def H(t,level=1,color=None):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(12 if level==1 else 8); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(t); r.font.bold=True
    if level==1:
        r.font.size=Pt(15); sca(r,color or NAVY)
        pPr=p._p.get_or_add_pPr(); b=OxmlElement("w:pBdr"); bt=OxmlElement("w:bottom")
        bt.set(qn("w:val"),"single"); bt.set(qn("w:sz"),"6"); bt.set(qn("w:space"),"2"); bt.set(qn("w:color"),TEAL)
        b.append(bt); pPr.append(b)
    else: r.font.size=Pt(12); sca(r,color or TEAL)
def P(t,italic=False,color=None,size=10.5):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(5)
    for i,ch in enumerate(t.split("**")):
        if not ch: continue
        r=p.add_run(ch); r.font.size=Pt(size); r.font.italic=italic
        if i%2==1: r.font.bold=True
        if color: sca(r,color)
def BUL(items,size=9.5):
    for it in items:
        p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(3)
        for i,ch in enumerate(it.split("**")):
            if not ch: continue
            r=p.add_run(ch); r.font.size=Pt(size)
            if i%2==1: r.font.bold=True
def FIG(path,cap,width=9.4):
    doc.add_picture(path,width=Inches(width)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    P(cap,italic=True,color=GREY,size=8.5); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
def BTABLE(rows,caption=None):
    cols=["Brand","Iconic for","Connection (founded / HQ)","The surprising fact","Conf."]
    widths=[1.55,1.7,2.0,4.0,0.55]
    t=doc.add_table(rows=1,cols=len(cols)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style="Table Grid"
    for i,h in enumerate(cols):
        shade(t.rows[0].cells[i],NAVY); ctext(t.rows[0].cells[i],h,bold=True,white=True,size=8.7,align="left" if i<4 else "center")
    for r,(b,ic,con,fact,conf) in enumerate(rows):
        cells=t.add_row().cells; base=LIGHT if r%2==0 else "FFFFFF"
        ctext(cells[0],b,bold=True,color=NAVY,size=8.4); ctext(cells[1],ic,size=8.1)
        ctext(cells[2],con,size=8.0); ctext(cells[3],fact,size=8.1); ctext(cells[4],conf,size=8.0,align="center")
        for c in cells: shade(c,base)
    for i,w in enumerate(widths):
        for row in t.rows: row.cells[i].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    if caption: P(caption,italic=True,color=GREY,size=8.3)

# Title
for _ in range(1): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("THE SCENIC CITY SCOUT"); r.font.size=Pt(23); r.font.bold=True; sca(r,NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Born Here — Iconic Brands from the Scenic City"); r.font.size=Pt(17); sca(r,TEAL)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run(f"{total} verified brands across food, health, carpet, freight, finance, tech & industry — sorted by what happened to them")
r.font.size=Pt(11); r.font.italic=True; sca(r,SLATE)
doc.add_page_break()

H("The thesis: Chattanooga is a brand town",1)
P("Most people never realize how many household names were born within an hour of downtown. Coca-Cola *bottling* itself, MoonPie, Little Debbie, Gold Bond and Icy Hot, Krystal, the world's tow trucks, most of America's carpet, and the 'Bloomberg of freight' all trace here. This is the verified source list for an occasional **'Born Here'** section — a sibling to *Still Standing* — sorted by current status so every entry comes with an honest hook.")
FIG(f"{ASSETS}/fig_bornhere.png","Marquee names by founding year — a brand town for 150 years.")

H("Three narratives worth a whole series",1)
BUL([
 "**The birthplace of bottled Coca-Cola (1899).** Two Chattanooga lawyers bought near-nationwide bottling rights for ~$1; Candler thought bottling was a fad. The global franchise system was invented here.",
 "**The Carpet Capital cluster.** Within ~30 miles — Shaw (world's #1, owned by Buffett), Mohawk (world's largest flooring co.), Engineered Floors, and Dixie — all descended from a 15-year-old's 1895 hand-tufted bedspread. Roughly 70%+ of the world's carpet is made near here.",
 "**The 'Silicon Valley of Trucking.'** The Fuller family (U.S. Xpress → FreightWaves) and the Access America 'Big Bang' (its exit funded Dynamo, Steam, Reliance, Bellhop) turned a mid-size city into a freight-tech capital — plus Astec and Miller Industries (the world's tow trucks, from Ooltewah).",
])

H("Still here — founded, headquartered, or made in the region today",1)
BTABLE(STILL)
H("The region's flooring & water giants",1)
BTABLE(REGION_CARPET)
doc.add_page_break()
H("Founded here → since moved, acquired, or revived",1)
BTABLE(MOVED)
H("Iconic operations (a famous plant, but HQ is elsewhere)",1)
BTABLE(OPS, caption="'Made here,' not 'born here' — include with that caveat.")
H("Vanished but legendary (industrial heritage)",1)
BTABLE(GONE)

doc.add_page_break()
H("Edge / regional — judgment calls",1)
BUL(EDGE)
H("Myths & accuracy traps the Scouts caught",1)
BUL(MYTHS)
H("How to use it",1)
P("Run **'Born Here'** occasionally (monthly/bimonthly), one brand at a time, led by the surprising fact and — where possible — a *living* local thread (the descendant, the plant still running, the building that's left). It pairs with products (a 'Born Here' trail/map, trivia, a heritage page in the Almanac) and events (a centennial toast). Same discipline as everything else: re-verify status before publishing, lead with the human angle, never a dry facts dump.")
P("Full sources for every entry are in born_here_brands.csv.",italic=True,color=GREY,size=9)

out="/home/user/prscott/Born_Here_Brands.docx"
doc.save(out); print("saved",out)
