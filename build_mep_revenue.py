#!/usr/bin/env python3
"""Mise en Place — The Revenue Model. Openly sponsored, events-first, lifestyle-scale.
The deliberate inverse of SCI/Encore's reader-funded model. Outputs .docx."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WINE="6E2B39"; AMBER="B5722E"; HOPS="5E7444"; CREAM="FAF4E9"; INK="2A2420"; SLATE="6b5f53"; GREY="8a7f70"
doc=Document(); sec=doc.sections[0]
for m in ("top_margin","bottom_margin"): setattr(sec,m,Inches(0.75))
for m in ("left_margin","right_margin"): setattr(sec,m,Inches(0.85))
nm=doc.styles["Normal"]; nm.font.name="Georgia"; nm.font.size=Pt(10.3); nm.paragraph_format.space_after=Pt(6); nm.paragraph_format.line_spacing=1.16
def sca(r,h): r.font.color.rgb=RGBColor.from_string(h)
def shade(c,h):
    p=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); p.append(s)
def ctext(c,t,bold=False,color=None,size=8.6,white=False,align="left"):
    c.text=""; p=c.paragraphs[0]; p.alignment={"left":WD_ALIGN_PARAGRAPH.LEFT,"center":WD_ALIGN_PARAGRAPH.CENTER}[align]
    for i,ch in enumerate(t.split("**")):
        if not ch: continue
        r=p.add_run(ch); r.font.size=Pt(size); r.font.bold=bold or (i%2==1)
        if white: r.font.color.rgb=RGBColor.from_string("FFFFFF")
        elif color: r.font.color.rgb=RGBColor.from_string(color)
def H(t,level=1,color=None):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(12 if level==1 else 7); p.paragraph_format.space_after=Pt(3)
    r=p.add_run(t); r.font.bold=True
    if level==1:
        r.font.size=Pt(14.5); sca(r,color or WINE)
        pPr=p._p.get_or_add_pPr(); b=OxmlElement("w:pBdr"); bt=OxmlElement("w:bottom")
        bt.set(qn("w:val"),"single"); bt.set(qn("w:sz"),"6"); bt.set(qn("w:space"),"2"); bt.set(qn("w:color"),AMBER); b.append(bt); pPr.append(b)
    else: r.font.size=Pt(11.5); sca(r,color or INK)
def P(t,italic=False,color=None,size=10.3):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(6)
    for i,ch in enumerate(t.split("**")):
        if not ch: continue
        r=p.add_run(ch); r.font.size=Pt(size); r.font.italic=italic
        if i%2==1: r.font.bold=True
        if color: sca(r,color)
def BUL(items,size=9.9):
    for it in items:
        p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(3)
        for i,ch in enumerate(it.split("**")):
            if not ch: continue
            r=p.add_run(ch); r.font.size=Pt(size)
            if i%2==1: r.font.bold=True
def TABLE(headers,rows,widths,fontsize=8.6,hdr=WINE):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style="Table Grid"
    for i,h in enumerate(headers):
        shade(t.rows[0].cells[i],hdr); ctext(t.rows[0].cells[i],h,bold=True,white=True,size=fontsize+0.2)
    for r,row in enumerate(rows):
        cells=t.add_row().cells; base=CREAM if r%2==0 else "FFFFFF"
        for i,val in enumerate(row):
            ctext(cells[i],val,size=fontsize,color=INK if i==0 else None,bold=(i==0)); shade(cells[i],base)
    for i,w in enumerate(widths):
        for row in t.rows: row.cells[i].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)

doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Mise en Place — The Revenue Model"); r.font.size=Pt(25); r.font.bold=True; sca(r,WINE)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Openly sponsored. Events-first. Lifestyle-scale."); r.font.size=Pt(13); r.font.italic=True; sca(r,INK)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("The sponsored companion to Scenic City Insider  ·  all figures illustrative & conservative  ·  project name “MeP”")
r.font.size=Pt(9); r.font.italic=True; sca(r,GREY)
doc.add_paragraph()

H("The philosophy — the deliberate opposite of SCI",1)
P("Where Scenic City Insider and Encore are reader-funded and ad-free, **MeP is openly, proudly sponsor-supported.** That's the point of the bright line: SCI keeps the company's soul pure, and MeP is the joyful, sponsorable passion title that can earn from issue one. The discipline that keeps it honest is **'the line'**: sponsors are always labeled, the picks are *never for sale*, and celebratory never means paid. The reader trusts MeP because the celebration is earned, not bought.")
P("**Two truths shape the model:** (1) the real margin is in **events**, not banner ads — restaurants are thin ad payers, but tasting dinners and classes are high-margin and on-brand; (2) Tennessee's **no-chain-liquor law** gives MeP a structurally local anchor sponsor base (~20–30 owner-run bottle shops) most cities don't have.")

H("The revenue lines",1)
TABLE(["Line","What it is","Who pays","Role"],
[["Events (the engine)","Supper clubs, tasting dinners, classes, trips","Attendees (+ host partners)","The profit center — highest margin, most on-brand"],
 ["Sponsorship","Presenting + category slots, labeled","Bottle shops, distilleries, restaurants, cookware, roasters","Recurring base; bottle shops anchor it"],
 ["Supper Club membership (light)","First-seat priority + perks (optional)","Regulars","Predictable trickle + locks in the superfans"],
 ["Commerce & the flywheel","The book, merch, partner classes; podcast halo","Readers","Modest margin + authority & reach"]],
[1.55,2.3,1.9,1.5],fontsize=8.5)

H("Events — the engine (illustrative economics)",1)
TABLE(["Format","Shape","Economics"],
[["The Supper Club","20–40 seats, a shop/somm + a kitchen (the Bill Reed / Pedestrian model), monthly-ish","~$100–150/seat; partner supplies wine/food at cost or rev-share; ~$1–1.5K net/night + the week's best content"],
 ["Classes & tastings","12–24 people; wine, whiskey, cheese, pairing","~$50–90/seat; you lead them (WSET/bourbon creds); high margin"],
 ["From-the-Source trips","Small group; a region or distillery weekend","Premium ticket; members pay for access, you host — a clean business event"]],
[1.4,2.7,2.9],fontsize=8.5)

H("Sponsorship — real but shallow; anchor it on bottle shops",1)
P("Realistic pool: ~20–40 viable local cash sponsors, of which maybe 15–30 buy at modest regional rates. Package a **presenting sponsor** (rotating bottle shop) + a few **category slots** (a distillery, a cookware shop, a roaster, a restaurant) + **Supper Club partner** of the month. Keep manufacturer (brewery/distillery) and retailer (bottle shop) slots separate for tied-house compliance.")

H("Illustrative three-year sketch (conservative, lifestyle-scale)",1)
TABLE(["Year","Free list","Sponsorship","Events","Commerce/flywheel","Total (illustrative)"],
[["1 · Build + first events","~2,000","~$10K","~$15K","~$5K","~$30K"],
 ["2 · Cadence","~5,000","~$30K","~$50K","~$10K","~$90K"],
 ["3 · Established","~8,000","~$50K","~$90K","~$20K","~$160K"]],
[1.5,1.0,1.1,1.0,1.4],fontsize=8.4)
P("MeP is a **wonderful small business, not a scale play** — a devoted few-thousand list, a handful of sponsors, and **events as the real income.** It funds the book, throws great parties, deepens the relationships you love, and pays you back in joy as much as dollars. The certifications let *you* lead the rooms, which is what protects the margin.",italic=True,color=SLATE,size=9.8)

H("Guardrails",1)
BUL([
 "**Celebratory ≠ paid.** The one rule that protects the whole thing: you celebrate because it earned it, never because a slot was bought. Label every sponsor.",
 "**Alcohol advertising (TABC)** — email-list promotion by retailers is permitted; separate manufacturer vs. retailer slots; host tastings through the licensed partner. Short TN-attorney check before launch.",
 "**Events live or die on partner math** — price for ~40–50% delivered cost and pre-sell seats; the host kitchen/shop carries the license.",
 "**Name vetting** — 'Mise en Place' is a common culinary term; run the domain + USPTO/TESS pass. Avoid 'The Table' as a feature name (Food as a Verb's brand).",
])

out="/home/user/prscott/Mise_en_Place_Revenue_Model.docx"
doc.save(out); print("saved",out)
