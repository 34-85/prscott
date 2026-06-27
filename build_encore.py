#!/usr/bin/env python3
"""Encore — concept brief (working title). A national membership for the second-act generation.
Positioning, audience, content architecture, community & model. Outputs .docx + a white-space chart."""
import os, matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INK="23282E"; CLARET="6E2A35"; BRASS="B08D57"; IVORY="FBF8F2"; SLATE="5A6470"; GREY="8A8170"; GREEN="2E7D52"
ASSETS="/home/user/prscott/assets"; os.makedirs(ASSETS,exist_ok=True)
plt.rcParams.update({"font.family":"DejaVu Serif","figure.dpi":160})

# ---- white-space chart ----
def chart():
    pts=[("ENCORE",0.85,0.90,CLARET,True),
         ("Retirement / AARP-style media",0.17,0.83,SLATE,False),
         ("Retirement-finance media",0.26,0.68,SLATE,False),
         ("Startup / creator / business media",0.82,0.20,INK,False),
         ("Generic lifestyle media",0.42,0.45,GREY,False)]
    fig,ax=plt.subplots(figsize=(9.2,5.6)); ax.set_xlim(0,1); ax.set_ylim(0,1)
    ax.axhline(0.5,color="#d9d2c4",lw=1.1); ax.axvline(0.5,color="#d9d2c4",lw=1.1)
    ax.add_patch(plt.Rectangle((0.5,0.5),0.5,0.5,facecolor="#"+CLARET,alpha=0.06,zorder=0))
    ax.text(0.985,0.985,"WHITE SPACE\nambitious + built for 55+",ha="right",va="top",fontsize=8.5,
            color="#"+CLARET,fontweight="bold",style="italic")
    for name,x,y,c,big in pts:
        ax.scatter([x],[y],s=300 if big else 95,color="#"+c,edgecolor="white",linewidth=1.5,zorder=3)
        ax.annotate(name,(x,y),xytext=(0,12 if big else 9),textcoords="offset points",ha="center",
                    fontsize=9 if big else 7.4,fontweight="bold" if big else "normal",color="#"+INK)
    ax.set_xticks([0.07,0.93]); ax.set_xticklabels(["Decline &\npreservation","Reinvention &\nambition"],fontsize=8.5,color="#"+SLATE)
    ax.set_yticks([0.07,0.93]); ax.set_yticklabels(["Ignores the\n55+ reader","Built for the\n55+ reader"],fontsize=8.5,color="#"+SLATE,rotation=90,va="center")
    for s in ("top","right"): ax.spines[s].set_visible(False)
    for s in ("left","bottom"): ax.spines[s].set_color("#c9c0b0")
    ax.set_title("Where Encore sits — the second-act white space",fontweight="bold",color="#"+INK,fontsize=12)
    fig.tight_layout(); fig.savefig(f"{ASSETS}/fig_encore.png",bbox_inches="tight"); plt.close(fig)
chart(); print("chart done")

# ---- DOCX ----
doc=Document(); sec=doc.sections[0]
for m in ("top_margin","bottom_margin"): setattr(sec,m,Inches(0.8))
for m in ("left_margin","right_margin"): setattr(sec,m,Inches(0.9))
nm=doc.styles["Normal"]; nm.font.name="Georgia"; nm.font.size=Pt(10.5); nm.paragraph_format.space_after=Pt(6); nm.paragraph_format.line_spacing=1.18
def sca(r,h): r.font.color.rgb=RGBColor.from_string(h)
def shade(c,h):
    p=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); p.append(s)
def ctext(c,t,bold=False,color=None,size=8.6,white=False,align="left"):
    c.text=""; p=c.paragraphs[0]; p.alignment={"left":WD_ALIGN_PARAGRAPH.LEFT,"center":WD_ALIGN_PARAGRAPH.CENTER}[align]
    for i,ch in enumerate(t.split("**")):
        if not ch: continue
        r=p.add_run(ch); r.font.size=Pt(size); r.font.bold=bold or (i%2==1)
        if white: r.font.color.rgb=RGBColor.from_string("FFFFFF")
        elif color: r.font.color.rgb=RGBColor.from_string(color)
def H(t,level=1,color=None):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(13 if level==1 else 8); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(t); r.font.bold=True
    if level==1:
        r.font.size=Pt(15); sca(r,color or CLARET)
        pPr=p._p.get_or_add_pPr(); b=OxmlElement("w:pBdr"); bt=OxmlElement("w:bottom")
        bt.set(qn("w:val"),"single"); bt.set(qn("w:sz"),"6"); bt.set(qn("w:space"),"2"); bt.set(qn("w:color"),BRASS); b.append(bt); pPr.append(b)
    else: r.font.size=Pt(12); sca(r,color or INK)
def P(t,italic=False,color=None,size=10.5):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(6)
    for i,ch in enumerate(t.split("**")):
        if not ch: continue
        r=p.add_run(ch); r.font.size=Pt(size); r.font.italic=italic
        if i%2==1: r.font.bold=True
        if color: sca(r,color)
def BUL(items,size=10):
    for it in items:
        p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(3)
        for i,ch in enumerate(it.split("**")):
            if not ch: continue
            r=p.add_run(ch); r.font.size=Pt(size)
            if i%2==1: r.font.bold=True
def FIG(path,cap,width=6.6):
    doc.add_picture(path,width=Inches(width)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    P(cap,italic=True,color=GREY,size=8.5); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
def TABLE(headers,rows,widths,fontsize=8.6):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style="Table Grid"
    for i,h in enumerate(headers):
        shade(t.rows[0].cells[i],CLARET); ctext(t.rows[0].cells[i],h,bold=True,white=True,size=fontsize+0.2)
    for r,row in enumerate(rows):
        cells=t.add_row().cells; base=IVORY if r%2==0 else "FFFFFF"
        for i,val in enumerate(row):
            ctext(cells[i],val,size=fontsize,color=INK if i==0 else None,bold=(i==0))
            shade(cells[i],base)
    for i,w in enumerate(widths):
        for row in t.rows: row.cells[i].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)

# Title
doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("ENCORE"); r.font.size=Pt(40); r.font.bold=True; sca(r,CLARET)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("For people whose best work is still ahead."); r.font.size=Pt(15); r.font.italic=True; sca(r,INK)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Concept brief — positioning, audience, content, community & model  ·  a Brief Scout Media title  ·  working title, name pending vetting")
r.font.size=Pt(9.5); r.font.italic=True; sca(r,GREY)
doc.add_page_break()

H("The thesis — and the white space",1)
P("Retirement is a 20th-century idea. A growing cohort of accomplished people in their late 50s, 60s, and early 70s are not winding down — they're **reloading**: starting second companies, buying and running small businesses, angel investing, taking board seats, mentoring, writing the book, making the art, traveling with intent. They have capital, time, hard-won judgment, and restlessness. And almost no one is making smart, ambitious media for them.")
P("The market insults this reader from both sides. **Retirement media** frames them as declining — Medicare, downsizing, estate planning, golf. **Startup, creator, and business media** is obsessed with the 25-year-old founder and ignores them entirely. The intersection — *ambitious media built for the 55+ reader* — is open ground.")
FIG(f"{ASSETS}/fig_encore.png","The second-act white space: the upper-right quadrant — reinvention-minded AND built for 55+ — is essentially unoccupied.")

H("Why this is defensible — the founder is the moat",1)
P("Encore's unfair advantage is that the founder **is** the reader: a retired CEO, the author of four books, in the middle of his own creative second act, openly looking for deeper connections with people doing the same. You cannot fake that. The voice has standing the moment it speaks, and the audience trusts a peer in a way it never trusts a 30-year-old editor explaining their lives back to them. Add the reader-funded ethic of the wider company — no one's attention is sold — and a community that compounds, and you have a brand that's hard to copy and easy to love.")

H("Who it's for — five overlapping readers",1)
TABLE(["The reader","What they're doing","What Encore gives them"],
[["The Founder-Again","Starting a second company or buying a small business","Playbooks, peers, and proof it's done after 55"],
 ["The Angel / Investor","Deploying capital, joining boards, syndicates","How-to, deal-adjacent community, judgment"],
 ["The Creator","Writing the book, making the art they deferred","Craft, encouragement, fellow late-bloomers"],
 ["The Mentor / Operator","Advisory, fractional, board & nonprofit work","Ways to contribute that use the whole career"],
 ["The Explorer","Reinventing the life — travel, learning, place","Expeditions, curricula, intent over idleness"]],
[1.5,2.6,2.6],fontsize=9)
P("The common thread isn't age — it's posture. **Experienced, resourced, and not done.**",italic=True,color=SLATE,size=10)

H("The weekly — content architecture",1)
P("Each section is a distinct point of view a competitor can't cheaply copy, and each ties to revenue without selling editorial.",italic=True,color=GREY,size=9.5)
TABLE(["Section","Frequency","The point of view","Monetization"],
[["The Opening","Weekly","The founder's first-person essay on the second-act life — the heart of the thing","Halo; the voice that drives membership"],
 ["The Encore Profile","Weekly","Long-form profile of someone who reinvented after 50 — inspiration with receipts","Flagship shareable; series sponsor; book ties"],
 ["The Playbook","Weekly","Tactical how-to for a second-act move (angel checks, buying a business, board seats, publishing)","Premium deep-dives; cohorts; partner tools"],
 ["The Ledger","Biweekly","Money for **deploying** wealth, not just preserving it — the opposite of retirement-income content","Walled fee-only advisor partners (labeled)"],
 ["The Table","Weekly (members)","The community: member asks & offers, intros, co-founder / advisor / investor matching","THE membership driver; directory; events"],
 ["The Expedition","Monthly","Travel & experience with intent — sabbaticals, learning trips, working from elsewhere","Member retreats & trips (high-margin events)"],
 ["The Long Game","Monthly","Health-span, relationships, meaning, legacy — performance, never decline","Longevity / coaching partners; retreats"],
 ["Last Word","Weekly","A provocation to carry into the week","Shareable; brand"]],
[1.35,1.0,3.05,2.0],fontsize=8.4)

H("The membership & community model",1)
P("This audience values **connection and contribution over price** — which is why community and events monetize far better here than advertising ever could, and why membership can begin earlier than a local newsletter's.")
BUL([
 "**Free** — the weekly Opening, the Profile, and Last Word. The front door.",
 "**Member (~$20–30/mo or ~$200–300/yr)** — the full weekly, the archive, and **The Table**: the member directory, asks-and-offers board, intros, and online gatherings. The network is the product.",
 "**Founding Circle (premium / invite, ~$1–3K/yr)** — curated dinners, small peer cohorts (mastermind-style), priority on retreats, and direct access. This audience pays for the *room*.",
 "**Events & retreats** — the largest line: Encore Retreats, regional dinners, and peer cohorts. A high-margin, high-trust monetizer the ad model can't touch.",
 "**Curated commerce & walled partners** — books, tools, experiences, and clearly-labeled premium partners (fee-only advisors, longevity, education, travel). Same bright line as the rest of the company: a partner can sponsor a room or a guide — never a sentence of editorial.",
])

H("The name — an honest flag",1)
P("**'Encore' is a strong idea but a crowded word.** Encore.org (now CoGenerate) literally pioneered the term 'encore careers'; Wynn's Encore casinos and various Encore media/ticketing brands exist too. So — exactly as we did for Scenic City Insider — the name needs a **domain + USPTO/TESS clearance** before any commitment, and the distinctiveness will likely have to live in a lockup (e.g., *The Encore Letter*, *Encore Society*) or an alternative entirely. Worth vetting alongside: **Second Act, Act Two, The Third Act, Next Chapter, Reload, Vantage, Prime.** Use 'Encore' as the working title; don't print it on anything yet.")

H("How it fits Brief Scout Media",1)
P("Encore is the third title in one portfolio, on shared infrastructure: the same HITL+AI production system, the reader-funded ethic, and the community DNA. It can even borrow a **Signal-style data section** — 'what second-act founders are quietly filing and funding' (their own Form Ds and trademarks) — and, later, **local chapters** that echo Scenic City Insider's place-based community. One company, three audiences: a town (Insider), an industry/market (The Signal), and a generation (Encore).")

H("Launch path",1)
BUL([
 "**Months 1–6:** publish the weekly; establish the founder's voice; grow the free list; open a founding-member waitlist.",
 "**Months 6–12:** turn on membership and The Table; host the first dinners and a pilot retreat; convert founding members (this audience pays sooner than a local list).",
 "**Year 2+:** scale retreats and cohorts; add walled premium partners; consider local chapters and the Signal-style data section.",
])
P("Full sample first issue accompanies this brief.",italic=True,color=GREY,size=9)

out="/home/user/prscott/Encore_Concept_Brief.docx"
doc.save(out); print("saved",out)
