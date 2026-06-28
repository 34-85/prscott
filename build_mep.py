#!/usr/bin/env python3
"""Mise en Place (MeP) — positioning one-pager (v2: food + drink, equally).
Chattanooga's celebration of the table — the openly-sponsored companion to Scenic City Insider."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WINE="6E2B39"; AMBER="B5722E"; HOPS="5E7444"; CREAM="FAF4E9"; INK="2A2420"; SLATE="6b5f53"; GREY="8a7f70"
doc=Document(); sec=doc.sections[0]
for m in ("top_margin","bottom_margin"): setattr(sec,m,Inches(0.7))
for m in ("left_margin","right_margin"): setattr(sec,m,Inches(0.85))
nm=doc.styles["Normal"]; nm.font.name="Georgia"; nm.font.size=Pt(10.3); nm.paragraph_format.space_after=Pt(6); nm.paragraph_format.line_spacing=1.16
def sca(r,h): r.font.color.rgb=RGBColor.from_string(h)
def shade(c,h):
    p=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); p.append(s)
def ctext(c,t,bold=False,color=None,size=8.6,white=False,align="left"):
    c.text=""; p=c.paragraphs[0]; p.alignment={"left":WD_ALIGN_PARAGRAPH.LEFT,"center":WD_ALIGN_PARAGRAPH.CENTER}[align]
    for i,ch in enumerate(t.split("**")):
        if not ch: continue
        r=p.add_run(ch); r.font.size=Pt(size); r.font.bold=bold or (i%2==1)
        if white: r.font.color.rgb=RGBColor.from_string("FFFFFF")
        elif color: r.font.color.rgb=RGBColor.from_string(color)
def H(t,level=1,color=None):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(12 if level==1 else 7); p.paragraph_format.space_after=Pt(3)
    r=p.add_run(t); r.font.bold=True
    if level==1:
        r.font.size=Pt(14.5); sca(r,color or WINE)
        pPr=p._p.get_or_add_pPr(); b=OxmlElement("w:pBdr"); bt=OxmlElement("w:bottom")
        bt.set(qn("w:val"),"single"); bt.set(qn("w:sz"),"6"); bt.set(qn("w:space"),"2"); bt.set(qn("w:color"),AMBER); b.append(bt); pPr.append(b)
    else: r.font.size=Pt(11.5); sca(r,color or INK)
def P(t,italic=False,color=None,size=10.3):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(6)
    for i,ch in enumerate(t.split("**")):
        if not ch: continue
        r=p.add_run(ch); r.font.size=Pt(size); r.font.italic=italic
        if i%2==1: r.font.bold=True
        if color: sca(r,color)
def BUL(items,size=9.8):
    for it in items:
        p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(3)
        for i,ch in enumerate(it.split("**")):
            if not ch: continue
            r=p.add_run(ch); r.font.size=Pt(size)
            if i%2==1: r.font.bold=True
def TABLE(headers,rows,widths,fontsize=8.6,hdr=WINE):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style="Table Grid"
    for i,h in enumerate(headers):
        shade(t.rows[0].cells[i],hdr); ctext(t.rows[0].cells[i],h,bold=True,white=True,size=fontsize+0.2)
    for r,row in enumerate(rows):
        cells=t.add_row().cells; base=CREAM if r%2==0 else "FFFFFF"
        for i,val in enumerate(row):
            ctext(cells[i],val,size=fontsize,color=INK if i==0 else None,bold=(i==0)); shade(cells[i],base)
    for i,w in enumerate(widths):
        for row in t.rows: row.cells[i].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)

# Title
doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Mise en Place"); r.font.size=Pt(34); r.font.bold=True; sca(r,WINE)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Eat well, drink well, Chattanooga — and meet the people who make it."); r.font.size=Pt(13.5); r.font.italic=True; sca(r,INK)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Positioning one-pager (v2: food + drink, equally)  ·  the openly-sponsored companion to Scenic City Insider  ·  project name “MeP” (pending name vetting)")
r.font.size=Pt(9); r.font.italic=True; sca(r,GREY)
doc.add_paragraph()

H("What MeP is",1)
P("**Chattanooga's celebration of the table — food and drink, equally — and the people who make it worth it.** Joyful, useful, free, and unapologetically apolitical. It celebrates the *craft* of cooking, pouring, and serving — the chef, the bartender, the baker, the shop owner — and pairs that with a credentialed drinks beat (wine, beer, bourbon, and the no/low-ABV side) that no one else in town has. The drink is one door into the room; the plate is the other.")
P("**Why food *and* drink, not drink alone:** the audience is bigger and the publication is sturdier. Gen Z drinks less; breakfast and lunch aren't drinking occasions; and many of the best stories are simply a great plate. Going food-equal opens the whole clock — biscuit to nightcap — widens the audience (including the sober-curious), and broadens the sponsor base (roasters, bakeries, daytime spots) well beyond bottle shops.",italic=True,color=SLATE,size=10)

H("The line — MeP's editorial DNA",1)
BUL([
 "**Celebrate the craft.** Cover people for what they make and the joy they create — never as a cause.",
 "**Delight over didactics.** We're here for the pleasure of the table, not the lecture.",
 "**Politics off the plate.** Everyone's welcome at this table. No agenda but a good one — never an activist rant, even a right one.",
 "**Always useful.** Where to eat and drink, what to order, what's new, who to know.",
 "**The plate and the glass are both the door** — and we walk through to meet the people behind the bar and the stove.",
 "**Celebratory ≠ paid.** We celebrate because it earned it, not because someone bought a slot (same bright line as everything we do).",
])

H("Why the lane is open — the competitive wedge",1)
P("A general food newsletter would be a weak bet — that's crowded. But the *celebratory, craft-and-people, genuinely useful, free, apolitical, drinks-expert* lane is open, because each incumbent leaves it empty:")
TABLE(["Incumbent","What they are","MeP's wedge"],
[["Food as a Verb","Food-as-society / agrarian, literary, activist","We celebrate the craft & the joy — the everyday beloved spot, the cook, the bartender — plus a drinks beat & service. Never a lecture. (They'll never spotlight a Choo Choo BBQ; we will.)"],
 ["Times Free Press 'What to Eat Next'","Curated food, but paywalled / limited access","Free — reaches the whole city, not just subscribers."],
 ["NOOGAtoday","Broad daily news; food is one section + ad roundups","Dedicated, celebratory, craft-deep — with a credential-backed drinks beat none of them have."]],
[1.5,2.2,3.3],fontsize=8.4)

H("The bright line with Scenic City Insider",1)
P("MeP sits **next to** SCI, never inside it. **SCI is the reader-funded soul** — no ads, the reader is the customer. **MeP is the openly-sponsored, joyful passion title** — labeled sponsors and ticketed events. SCI may mention a place as part of knowing the city; MeP is the dedicated eat-and-drink-well celebration with events. Two mastheads, two honest models.")

H("The money model — three legs, in priority order",1)
TABLE(["Leg","What it is","Why / how it pays"],
[["1 · Events (the engine)","Intimate wine/whiskey dinners (the Bill Reed / Pedestrian model — 20–40 seats, a shop or somm + a kitchen), tastings, and classes","Highest margin, most on-brand; you host with credentials; scarcity & access are the product. A 24-seat dinner at ~$125 ≈ $3K gross. The night is also the week's best content."],
 ["2 · Sponsorship (recurring 2nd)","Presenting + category slots; bottle shops anchor, now joined by restaurants, roasters, bakeries, cookware","Going food-equal widens the pool well past wine/spirits; ~20–40 realistic local buyers at modest regional rates."],
 ["3 · The flywheel","Book + podcast feeding the newsletter","Authority, reach, and the relationships that become sponsors and event partners. The list is the owned asset."]],
[1.3,2.6,2.8],fontsize=8.5)
P("**Honest scale:** a lifestyle-scale business — a devoted few-thousand list and **events as the real income.** The blend (esp. events) is what makes it work; the certifications (WSET/bourbon) let *you* lead the rooms.",italic=True,color=SLATE,size=9.8)

H("The weekly — content architecture (twin signatures)",1)
TABLE(["Section","The point of view"],
[["The Plate","A dish, a cook, a room worth celebrating — the everyday craft (the BBQ, the biscuit, the taco, the meat-and-three), all dayparts"],
 ["The Pour","The wine, beer, bourbon — or great NA — worth seeking, with a credentialed reason why"],
 ["The Maker / The Counter","The people behind it — chef, baker, roaster, bartender, monger, shop owner — craft & personality, never cause"],
 ["What's New","Openings & what to order now — the freely-accessible service the paywall keeps from people"],
 ["Host This","A pairing or hosting tip — 'mise en place' made practical for friends coming over"],
 ["The Calendar","Tastings, dinners, releases, classes & festivals — the events that also monetize"]],
[1.5,5.2],fontsize=8.6)

H("The sponsor & event-partner map",1)
TABLE(["Tier","Role","Who (real local examples)"],
[["Tier 1","Recurring cash sponsors","Bottle shops (Riverside, Imbibe, Kanku's) · distilleries (Chattanooga Whiskey, Gate 11) · cookware (Good Kinsmen, The Kitchen Collection) · Bleu Fox cheese · coffee roasters (Velo, Goodman) · Niedlov's · restaurants & cafés (now in scope)"],
 ["Tier 2","Paid-event co-hosts","Restaurants for wine dinners (the Pedestrian / Bill Reed model) · breweries · wineries (Lookout, Georgia Winery) · cooking schools (PastaNooga, Sweet & Savory) · chefs/caterers · festivals (SIPTN, Bacon & Barrel)"],
 ["Tier 3","Content / relationship partners","Farms & CSAs, Chattanooga Market & Main Street Farmers Market, micro-makers — stories & credibility, not reliable cash"]],
[0.7,1.7,4.3],fontsize=8.3)

H("Guardrails",1)
BUL([
 "**Stay celebratory, not a listings sheet.** The value is craft, people, and a point of view — not a directory NOOGAtoday already runs.",
 "**Apolitical by design** — the inverse of the activist lane; it's also the more sponsorable and more shareable choice.",
 "**Alcohol advertising (TABC)** still applies to the drinks side & events: email-list promotion by retailers is permitted; keep manufacturer vs. retailer slots separate; host tastings through the licensed partner. A short TN-attorney check before launch.",
 "**Name vetting** — 'Mise en Place' is a common culinary term; run the domain + USPTO/TESS pass we used for SCI. Also avoid 'The Table' as any MeP feature name — it's Food as a Verb's membership brand.",
 "**Small-market ceiling** — lean on the events margin, not premium ad rates.",
])

out="/home/user/prscott/Mise_en_Place_One_Pager.docx"
doc.save(out); print("saved",out)
