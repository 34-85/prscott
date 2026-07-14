#!/usr/bin/env python3
"""Mise en Place — The Credential Roadmap. Sequenced wine/spirits/bourbon certifications
mapped to content, events, and 'from the source' trips. Outputs .docx."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WINE="6E2B39"; AMBER="B5722E"; HOPS="5E7444"; CREAM="FAF4E9"; INK="2A2420"; SLATE="6b5f53"; GREY="8a7f70"
doc=Document(); sec=doc.sections[0]; sec.orientation=WD_ORIENT.LANDSCAPE
sec.page_width,sec.page_height=Inches(11),Inches(8.5)
for m in ("top_margin","bottom_margin"): setattr(sec,m,Inches(0.55))
for m in ("left_margin","right_margin"): setattr(sec,m,Inches(0.7))
nm=doc.styles["Normal"]; nm.font.name="Georgia"; nm.font.size=Pt(10); nm.paragraph_format.space_after=Pt(5); nm.paragraph_format.line_spacing=1.13
def sca(r,h): r.font.color.rgb=RGBColor.from_string(h)
def shade(c,h):
    p=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); p.append(s)
def ctext(c,t,bold=False,color=None,size=8.4,white=False,align="left"):
    c.text=""; p=c.paragraphs[0]; p.alignment={"left":WD_ALIGN_PARAGRAPH.LEFT,"center":WD_ALIGN_PARAGRAPH.CENTER}[align]
    for i,ch in enumerate(t.split("**")):
        if not ch: continue
        r=p.add_run(ch); r.font.size=Pt(size); r.font.bold=bold or (i%2==1)
        if white: r.font.color.rgb=RGBColor.from_string("FFFFFF")
        elif color: r.font.color.rgb=RGBColor.from_string(color)
def H(t,level=1,color=None):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(12 if level==1 else 7); p.paragraph_format.space_after=Pt(3)
    r=p.add_run(t); r.font.bold=True
    if level==1:
        r.font.size=Pt(14.5); sca(r,color or WINE)
        pPr=p._p.get_or_add_pPr(); b=OxmlElement("w:pBdr"); bt=OxmlElement("w:bottom")
        bt.set(qn("w:val"),"single"); bt.set(qn("w:sz"),"6"); bt.set(qn("w:space"),"2"); bt.set(qn("w:color"),AMBER); b.append(bt); pPr.append(b)
    else: r.font.size=Pt(11.5); sca(r,color or INK)
def P(t,italic=False,color=None,size=10):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(5)
    for i,ch in enumerate(t.split("**")):
        if not ch: continue
        r=p.add_run(ch); r.font.size=Pt(size); r.font.italic=italic
        if i%2==1: r.font.bold=True
        if color: sca(r,color)
def BUL(items,size=9.6):
    for it in items:
        p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(3)
        for i,ch in enumerate(it.split("**")):
            if not ch: continue
            r=p.add_run(ch); r.font.size=Pt(size)
            if i%2==1: r.font.bold=True
def TABLE(headers,rows,widths,fontsize=8.4,hdr=WINE):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style="Table Grid"
    for i,h in enumerate(headers):
        shade(t.rows[0].cells[i],hdr); ctext(t.rows[0].cells[i],h,bold=True,white=True,size=fontsize+0.2)
    for r,row in enumerate(rows):
        cells=t.add_row().cells; base=CREAM if r%2==0 else "FFFFFF"
        for i,val in enumerate(row):
            ctext(cells[i],val,size=fontsize,color=INK if i==0 else None,bold=(i==0)); shade(cells[i],base)
    for i,w in enumerate(widths):
        for row in t.rows: row.cells[i].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)

# Title
doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Mise en Place — The Credential Roadmap"); r.font.size=Pt(26); r.font.bold=True; sca(r,WINE)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Earning the authority — in public. A sequenced wine, spirits & bourbon plan, mapped to content, events & trips.")
r.font.size=Pt(12); r.font.italic=True; sca(r,INK)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("All costs are approximate — verify current with each provider before enrolling.")
r.font.size=Pt(9); r.font.italic=True; sca(r,GREY)

H("Why this is a business asset, not a hobby",1)
BUL([
 "**Authority** — a masthead credential (WSET, Certified/Executive Bourbon Steward, Sommelier) separates MeP from every untrained food account and earns the trust of readers and sponsors.",
 "**Content** — earning each cert *in public* (studying, blind-tasting humility, exam day) is a pre-launch audience engine; the journey itself is material for the newsletter, the podcast, and the book.",
 "**Events margin** — once certified, *you* lead the paid tastings, classes, and dinners (MeP's profit center) instead of hiring a sommelier: better margin, higher ticket, you're the draw.",
 "**Relationships** — cert programs put you in rooms with the distillers, shop owners, and trade people who become guests, sponsors, and event partners.",
 "**One honest caution:** don't let the certificate become the starting gun. WSET Level 2 + curiosity is plenty to launch — let credentials *deepen* MeP, not gate it.",
])

H("The recommended sequence",1)
TABLE(["Phase","Earn this","Why now","~Cost (verify)","Format"],
[["1 · Launch-ready","WSET Level 1 Award in Wines  +  Stave & Thief Certified Bourbon Steward","Two fast, cheap credentials to put your name behind day-one content — and two 'I got certified' launch stories","~$210 + ~$60–100","Online (both)"],
 ["2 · Real authority (mo. 1–12)","WSET Level 2 Award in Wines  +  WSET Level 2 Award in Spirits","The credibility sweet spot: a real tasting framework (SAT) for wine, and the production/category framework for explaining bourbon. Lets you host classes with authority","~$500–800 + ~$500–625","SEBEVED (Nashville) or online"],
 ["3 · Signature + the trip (yr 1–2)","Executive Bourbon Steward, in person  +  (optional) CMS Introductory Sommelier","EBS at Moonshine U. in Louisville is your flagship 'from the source' feature AND a credential upgrade; CMS Intro adds pairing/service flair","~$500 + ~$595–695","Louisville (drive) / 2-day"],
 ["4 · Apex (only if MeP earns it, yr 2+)","WSET Level 3 Award in Wines  ·  Council of Whiskey Masters 'Master of Bourbon'","Serious blind-tasting wine authority; the strongest 'Master' title for a byline. Big time/$ — do only when revenue justifies","~$1,000–1,400  ·  ~$8,450","Online + exam / residential"]],
[1.15,2.5,3.5,1.35,1.4],fontsize=8.2)
P("**Recommended path total (Phases 1–3): roughly $2,300–3,000 plus the Louisville trip, ~100–120 study hours spread over a year.** Phase 4 is a separate, later decision.",italic=True,color=SLATE,size=9.5)

doc.add_page_break()
H("Track A — Wine (WSET + Court of Master Sommeliers)",1)
TABLE(["Credential","What it covers","Time / exam","~Cost (verify)","Format / where"],
[["WSET L1 Award in Wines","Beginner: structured tasting, main styles & grapes, basic pairing","~6 hrs · 30 MC, 45 min, 70% pass","~$210–400","Online or in person; SEBEVED, online APPs"],
 ["WSET L2 Award in Wines","Intermediate: key grapes, world regions, why climate/winemaking shape style; the SAT tasting method","~28 hrs · 50 MC, 60 min","~$500–815","SEBEVED (Nashville), Fine Vintage, WineWise (Atlanta), online"],
 ["WSET L3 Award in Wines","Advanced: explain *why* on style/quality/price; blind tasting","~84 hrs · theory paper + blind tasting of 2 wines","~$1,000–1,400","Classroom or online; major self-study"],
 ["CMS Introductory Sommelier","Foundational wine/beverage + intro to deductive tasting (enthusiast-open)","2-day course + exam · MC, 60% pass","~$595–695","In person or online; rotating host cities"],
 ["CMS Certified Sommelier","Deductive tasting + theory + tableside service (trade-oriented)","1 exam day · tasting/theory/service, 60% each","~$300–700","Exam only; requires Intro first"]],
[1.5,3.0,2.1,1.1,2.0],fontsize=8.1)
P("Alternatives: **Society of Wine Educators CSW** (self-study, one MC exam) for self-directed learners; **French/Italian Wine Scholar** to go deep on one country.",size=9.3,color=SLATE)

H("Track B — Spirits & Bourbon",1)
TABLE(["Credential","What it covers","Time / exam","~Cost (verify)","Format / where"],
[["WSET L2 Award in Spirits","How spirits are made (incl. bourbon/whiskey) & how production drives style; SAT for spirits; cocktail basics","~26 hrs · 50 MC, 60 min","~$500–625","SEBEVED or online (L1 Spirits = lighter start)"],
 ["Stave & Thief Certified Bourbon Steward","Bourbon production, history, sensory basics — the KDA's *official* program (instant name recognition)","Self-paced · online exam, 80% pass","~$60–100","Online — from Chattanooga"],
 ["Stave & Thief Executive Bourbon Steward","Advanced: classroom + hands-on distilling + advanced sensory","~1 day (in person) or self-paced online","~$500 in person","Moonshine U., Louisville (a 'from the source' trip) — or online"],
 ["Council of Whiskey Masters — Certified Bourbon Professional","Strong online bourbon credential, broader than Steward","~60–90 hrs · 100 MC, 80% pass","~$395","Online"],
 ["Council of Whiskey Masters — Master of Bourbon","Apex American-bourbon specialization — the byline title","Residential + on-site exam","~$8,450","Travel / residential"]],
[1.7,3.0,1.9,1.1,2.0],fontsize=8.1)
P("**Beer (MeP covers it too):** WSET Level 1/2 Awards in Beer (SEBEVED teaches them), or the beer-world standard **Cicerone** — Certified Beer Server (entry, online, cheap) → Certified Cicerone. Add when you turn to the beer beat.",size=9.3,color=SLATE)

doc.add_page_break()
H("What each credential unlocks for MeP — the ROI",1)
TABLE(["Credential","Content it unlocks","Event it unlocks"],
[["WSET L1 Wines","'Wine 101' explainers; The Pour with real authority","Casual intro tastings"],
 ["Certified Bourbon Steward","'Bourbon basics' & label-law pieces; credibility with distilleries","'Bourbon 101' nights in a whiskey town"],
 ["WSET L2 Wines","'Why this tastes like this' — the SAT lens; region deep-dives","Paid guided wine tastings & classes (you lead)"],
 ["WSET L2 Spirits","Explaining bourbon/whiskey production & categories; allocated-bottle authority","Whiskey flights & comparative classes"],
 ["Executive Bourbon Steward","Flagship 'From the Source: Kentucky' feature; deeper bourbon authority","Premium bourbon dinners; distillery collabs"],
 ["CMS Introductory","Pairing-forward 'drink well' content","Pairing dinners with chef partners"],
 ["WSET L3 / Master of Bourbon","Premium, can't-fake-it authority; book credibility","Member trips; flagship ticketed events"]],
[1.9,3.5,3.0],fontsize=8.4)

H("'From the source' — turning study into trips, content & revenue",1)
P("The certifications justify travel that doubles as premium content and, done right, a clean business activity:")
BUL([
 "**Louisville (Executive Bourbon Steward / Moonshine University)** — your first, easy 'from the source' trip (~5-hr drive): hands-on distilling, sensory lab, and a feature series + the credential in one.",
 "**Vineyard & region trips** (Willamette in pinot season, a Kentucky distillery crawl, sherry in Jerez) — the 'Mise en Place: From the Source' strand: the can't-fake-it material that justifies membership and seeds the chef/wine book.",
 "**MeP member trips** — a curated group trip is the highest-margin event of all: members pay for the access, you host. It's also the *legitimate* version of 'travel for the business' — a real, ticketed business event.",
 "**The tax reality (talk to your CPA):** travel is deductible when it's genuinely for a profit-seeking business — primarily business purpose, personal portion carved out, documented. Running MeP as a real business (revenue, clean books, member trips) is what makes the education and travel legitimately deductible; a money-losing 'hobby' is what the rules push back on.",
])
P("Caveats: every cost above is approximate (snippet-sourced) — confirm with each provider. The two figures hardest to pin were WSET Level 3 full price and the US Certified Sommelier fee. The standout regional provider is **SEBEVED in Nashville** (WSET wine, spirits & beer; in person and online).",italic=True,color=GREY,size=9)

out="/home/user/prscott/Mise_en_Place_Credential_Roadmap.docx"
doc.save(out); print("saved",out)
