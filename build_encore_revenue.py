#!/usr/bin/env python3
"""Encore — The Revenue Model. Reader-funded membership + community + events for the
second-act generation. Connection first, revenue as the byproduct. Outputs .docx."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INK="23282E"; CLARET="6E2A35"; BRASS="B08D57"; IVORY="FBF8F2"; SLATE="5A6470"; GREY="8A8170"; GREEN="2E7D52"
doc=Document(); sec=doc.sections[0]
for m in ("top_margin","bottom_margin"): setattr(sec,m,Inches(0.75))
for m in ("left_margin","right_margin"): setattr(sec,m,Inches(0.85))
nm=doc.styles["Normal"]; nm.font.name="Georgia"; nm.font.size=Pt(10.3); nm.paragraph_format.space_after=Pt(6); nm.paragraph_format.line_spacing=1.16
def sca(r,h): r.font.color.rgb=RGBColor.from_string(h)
def shade(c,h):
    p=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); p.append(s)
def ctext(c,t,bold=False,color=None,size=8.6,white=False,align="left"):
    c.text=""; p=c.paragraphs[0]; p.alignment={"left":WD_ALIGN_PARAGRAPH.LEFT,"center":WD_ALIGN_PARAGRAPH.CENTER}[align]
    for i,ch in enumerate(t.split("**")):
        if not ch: continue
        r=p.add_run(ch); r.font.size=Pt(size); r.font.bold=bold or (i%2==1)
        if white: r.font.color.rgb=RGBColor.from_string("FFFFFF")
        elif color: r.font.color.rgb=RGBColor.from_string(color)
def H(t,level=1,color=None):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(12 if level==1 else 7); p.paragraph_format.space_after=Pt(3)
    r=p.add_run(t); r.font.bold=True
    if level==1:
        r.font.size=Pt(14.5); sca(r,color or CLARET)
        pPr=p._p.get_or_add_pPr(); b=OxmlElement("w:pBdr"); bt=OxmlElement("w:bottom")
        bt.set(qn("w:val"),"single"); bt.set(qn("w:sz"),"6"); bt.set(qn("w:space"),"2"); bt.set(qn("w:color"),BRASS); b.append(bt); pPr.append(b)
    else: r.font.size=Pt(11.5); sca(r,color or INK)
def P(t,italic=False,color=None,size=10.3):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(6)
    for i,ch in enumerate(t.split("**")):
        if not ch: continue
        r=p.add_run(ch); r.font.size=Pt(size); r.font.italic=italic
        if i%2==1: r.font.bold=True
        if color: sca(r,color)
def BUL(items,size=9.9):
    for it in items:
        p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(3)
        for i,ch in enumerate(it.split("**")):
            if not ch: continue
            r=p.add_run(ch); r.font.size=Pt(size)
            if i%2==1: r.font.bold=True
def TABLE(headers,rows,widths,fontsize=8.6,hdr=CLARET):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style="Table Grid"
    for i,h in enumerate(headers):
        shade(t.rows[0].cells[i],hdr); ctext(t.rows[0].cells[i],h,bold=True,white=True,size=fontsize+0.2,align="left" if i==0 else "left")
    for r,row in enumerate(rows):
        cells=t.add_row().cells; base=IVORY if r%2==0 else "FFFFFF"
        for i,val in enumerate(row):
            ctext(cells[i],val,size=fontsize,color=INK if i==0 else None,bold=(i==0)); shade(cells[i],base)
    for i,w in enumerate(widths):
        for row in t.rows: row.cells[i].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)

# Title
doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("ENCORE — The Revenue Model"); r.font.size=Pt(26); r.font.bold=True; sca(r,CLARET)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Reader-funded. Connection first, revenue as the byproduct."); r.font.size=Pt(13); r.font.italic=True; sca(r,INK)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("A Brief Scout Media title  ·  national, for accomplished people 55–75 building what's next  ·  all figures illustrative & conservative")
r.font.size=Pt(9); r.font.italic=True; sca(r,GREY)
doc.add_paragraph()

H("The philosophy",1)
P("Encore is your personal-journey product first, so the model is built so that **money follows connection and never leads it.** It is **reader-funded by design** — no ads, no sponsored content, and (the deliberate inverse of ROAR) **it never sells its members to brands.** The people Encore serves are the only people it answers to. That integrity is both the soul of the thing and, with this audience, the smarter business.")
P("Two facts make the economics work without advertising: the audience is **affluent and time-rich**, and it values **access, peers, and contribution over price.** That means **small numbers produce real revenue**, and the highest-value lines — community and live experiences — are exactly the ones an ad model can't touch. Membership can also begin *earlier* than a local newsletter's, because this reader is used to paying for rooms worth being in.")

H("The revenue lines",1)
TABLE(["Line","What it is","Who pays","Role"],
[["Membership","Paid access to the community + full archive + member events","The reader","The recurring base — predictable, compounding"],
 ["Events & retreats","Working retreats, peer cohorts, regional dinners","Members & guests","The engine — highest margin, deepest trust"],
 ["The Founders' Circle","A small premium inner ring: cohorts, direct access, priority","The most committed readers","The profit core — 'pays for the room' pricing"],
 ["Curated commerce","The book, tools, experiences, member trips","The reader","Halo + modest margin; the book is authority"],
 ["(Deliberately absent)","No ads · no sponsored content · no selling members to brands","—","The integrity moat vs. ROAR"]],
[1.45,2.5,1.6,1.6],fontsize=8.5)

H("Membership tiers",1)
TABLE(["Tier","Price (illustrative)","What you get"],
[["Free","$0","The weekly Opening, the Encore Profile, Last Word — the front door"],
 ["Member","~$15–25/mo or ~$180–250/yr","The full weekly + archive + **The Roundtable** (member directory, asks-and-offers, intros, online gatherings) + member-rate events"],
 ["Founders' Circle","~$3,000–7,500/yr (or invite)","A small peer cohort (mastermind-style), priority/included retreat seats, direct access, the inner room"]],
[1.3,1.8,3.6],fontsize=8.6)
P("Note: 'The Roundtable' is the community feature name (it replaces the earlier 'The Table,' which is Food as a Verb's brand).",italic=True,color=GREY,size=9)

H("Events — the engine",1)
P("For this audience, live experiences are where both the value and the margin concentrate. Three formats, in rising price/intimacy:")
TABLE(["Format","Shape (illustrative)","Economics"],
[["Founders' Dinners","20–30 people, regional, a few times a year","$100–250/seat; modest cash, high relationship value & funnel"],
 ["Working Retreats","25–40 people, multi-day, curated","$2,500–5,000/seat; venue/partner cost ~40–50%; strong net + the year's best content"],
 ["Peer Cohorts","8–12 people, year-long, mastermind-style","$5,000–12,000/yr each; the premium recurring line — small groups, deep value"]],
[1.4,2.7,2.9],fontsize=8.5)
P("Member trips ('From the Source'-style, but for second-act experiences) are a natural extension — members pay for access, the founder hosts. A clean, high-margin business event.",italic=True,color=SLATE,size=9.8)

doc.add_page_break()
H("Illustrative three-year sketch (conservative)",1)
P("National niche + affluent reader. Figures are illustrative and deliberately conservative — to show the *shape*, not to forecast. Assumes membership ARPU ~$200/yr and that events grow as the community and the founder's bandwidth do.",italic=True,color=GREY,size=9.5)
TABLE(["Year","Free list","Paid members","Membership","Events","Book/commerce","Total (illustrative)"],
[["1 · Build","~5,000","~150 founding","~$30K","~$25K (pilot cohort/dinner)","~$10K","~$65K"],
 ["2 · Turn on","~20,000","~800","~$160K","~$180K (retreats + cohorts)","~$20K","~$360K"],
 ["3 · Scale","~50,000","~2,500","~$500K","~$400K","~$50K","~$950K"]],
[1.0,0.95,1.0,1.0,1.25,1.35],fontsize=8.3)
P("Read it as a range, not a promise: the point is that a *reader-funded* model reaches meaningful revenue on modest audience numbers, because the ARPU and event margins are high. A slower 'right-sized' path (smaller list, fewer events) is a perfectly good outcome and stays true to the personal-journey intent.",size=9.8,color=SLATE)

H("The north star — the inner-ring math",1)
P("Scenic City Insider's mountain is **200 members × $1,000** for experiences that make $1,000 feel cheap. Encore's version is the same idea with a bigger, national, affluent base: **a Founders' Circle of ~200 at ~$5,000 ≈ $1M from the inner ring alone** — built on cohorts, retreats, and access that make $5,000 feel inexpensive and indispensable. The broad membership and events then sit *on top* of that core. You don't need a mass audience; you need a few hundred people who'd pay for the room.")

H("What Encore will not do — and why that's the moat",1)
BUL([
 "**No advertising, no sponsored content, no selling members to marketers.** ROAR's revenue comes from the brands it covers; Encore's comes from the people it serves. With a skeptical, accomplished audience, that integrity *is* the product.",
 "**Make the network — not the founder — the durable asset.** Both Encore and ROAR are founder-led; Encore's edge is that The Roundtable and the cohorts create value member-to-member, so the business outlasts any one voice.",
 "**Don't let revenue corrupt the soul.** This is your personal-journey product. Price on value and access, keep it reader-funded, and let the money be the happy byproduct of a community that works.",
])

H("Guardrails",1)
BUL([
 "**Regulatory caution on 'deal flow.'** The Roundtable can host member asks/offers and introductions, but the moment Encore takes a fee or carry on investments/syndicates, it enters securities-regulated territory — keep matchmaking informal, or get counsel before monetizing it.",
 "**Events carry real cost & logistics** — retreats live or die on the partner/venue math; price for ~40–50% delivered cost and pre-sell seats.",
 "**Name vetting** — 'Encore' is crowded (Encore.org/CoGenerate, Wynn); run the domain + USPTO/TESS pass before committing, likely as a lockup. (Working title only.)",
 "**Sequence it:** Year 1 — grow the list, open a founding-member waitlist, run one pilot cohort; Year 2 — turn on membership + The Roundtable, host the first retreats; Year 3 — scale cohorts/retreats and the Founders' Circle.",
])

out="/home/user/prscott/Encore_Revenue_Model.docx"
doc.save(out); print("saved",out)
