#!/usr/bin/env python3
"""New-Openings Watchlist (Word + CSV) from Hamilton County business-license mining,
plus the updated Scout Sourcing Doctrine v2."""
import csv
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY="1B3A4B"; TEAL="2A9D8F"; ORANGE="E76F51"; SAND="E9C46A"; SLATE="415A6B"
LIGHT="EEF3F4"; GREY="6B7B85"; GOLD="B8860B"; GREEN="2E7D52"; RED="C0492F"

# ---- TIER 1: feature now (verified open/real, on-brand, story-rich) ----
TIER1=[
 ("Creature Comforts Beer Garden","Adam Beauchamp (CEO)","Brewery / beer garden · Foundries District (Pipe Way)","Open (Apr 2026)","Beloved Athens, GA brewery plants its first TN flag in a reborn foundry by the Lookouts' new stadium; permanent Riverwalk spot coming.","creaturecomfortsbeer.com","Some"),
 ("Pauly's","Paul Graves","Bar · 1627 Rossville Ave (the former Zarzour's)","Coming (confirm)","Moving into Chattanooga's oldest restaurant address — Zarzour's, est. 1918, closed after 107 years. The marquee heritage story of the batch.","IG @paulysbarroom_","Low"),
 ("The Trinity","Ken Newburger & Carlitos Fernandez (chef Andrew Tucker)","Restaurant (New Orleans Creole) · 1278 Market St","Open (Spring 2026)","Two college friends bring NOLA cooking downtown; named for the Cajun 'holy trinity.'","IG @thetrinitychatt","Some"),
 ("Goonzquad Gear","Simeon & Eleazar Drevenchuk","Apparel / merch · 4108 9th Ave","Open","Two immigrant brothers behind a multi-million-subscriber Chattanooga car-rebuild YouTube empire; merch designed & shipped locally — huge built-in audience.","goonzquad.com / @goonzquad","Some"),
 ("Electric Crocodile","Michelle & Paul Kuffrey","Vintage resale collective · new 2nd loc 1791 Reggie White Blvd","Open (expanding)","Neon-pink vintage collective (40+ local vendors) that survived a storm tearing through its stock — now opening a second location near the breweries.","FB The Electric Crocodile","Some"),
 ("Petal and Pour","Candace Barber & Keegan Chitwood","Floral atelier + café · 800 Broad St","Open (2026)","Billed as Chattanooga's first immersive floral atelier + café; founders travel to NYC yearly for inspiration.","@thepetalandpour","Some"),
 ("C. Winston Smith Bespoke Tailoring","C. Winston Smith","Bespoke tailor / clothing maker · 5020 University Dr, Collegedale","Open","A master tailor, 46+ years in the trade — started in 1972 Brooklyn, also a 20-yr fashion-design professor — who relocated his craft to Collegedale.","Yelp / FB","Some"),
 ("The Quilty Pleasure","Stella Rankin","Maker — clothing from vintage quilts · Soddy Daisy","Open","Hoodies and jackets sewn from tattered antique quilts others would toss — 'each piece stitched with a little history.'","thequiltypleasure.com","Low"),
 ("Smoking Crow Creations","Matthew Johns","Maker — artisan briar tobacco pipes · Hixson","Open","A solo briar-pipe carver working a niche heritage craft (handle hints at a fire/sideshow-performer crossover).","smokingcrowcreations.com","Low"),
 ("Broken & Bottled","Tiffany Reid","Maker — custom art · Chattanooga","Open","Chattanooga Market artist (Strawberry Festival vendor): 'it becomes their piece of art — I'm just making it for them.'","@tiffanyreidart","Some"),
 ("Offbeat Eats","Amanda Ertley","Microbakery (sourdough) · home-based, sells at markets","Open","Husband-and-wife cottage bakery; chef-driven, stone-milled organic flour, baguette preorders at the Ooltewah/St. Albans/Farmers-on-First markets.","@offbeat.eats","Low"),
 ("Good News Roasters","Kent & Lisa Smith","Coffee roaster + shop + mobile · 9710 Dayton Pike, Soddy Daisy","Open (est. 2021)","Mission-driven micro-roaster; profits support the poor, homeless, widowed, orphaned. Signature drinks 'Amazing Grace,' 'Raise a Hallelujah.'","goodnewsroasters.com","Low"),
 ("Loma's Ja-Maken Mi Kraaven","Loma Fisher Swaby","Jamaican meal-prep / pop-ups · Kitchen Incubator (5704 Marlin Rd)","Open","A Jamaican home cook building a brand out of the food incubator; punny patois name ('Jamaican me cravin'').","none found","Low"),
 ("Shucos Bites and Tacos","Brenda Ramos","Guatemalan shucos (street hot dogs) + tacos · 951 S Watkins St + mobile","Open","One of the only Guatemalan shuco spots in town — distinctive cuisine, family-run, strong Spanish-language following.","FB Shucos Bites And Chattanooga","Low"),
 ("The Southern Sazon Sizzle","Elaina Laureano","Puerto Rican food truck · Hixson/Middle Valley","Open","A home-grown Puerto Rican truck bringing sazón and lechón to Hixson (scored 100 on its health inspection).","none found","Low"),
]
# ---- TIER 2: watch & verify (promising; confirm open/owner before featuring) ----
TIER2=[
 ("Signal Mountain Cookie Lady","Heather Allison (license)","Cookie/cake shop · 1217 Taft Hwy, Signal Mtn","Open — ownership change?","The famous shop was Sandie Benson's for ~13 yrs; the license now lists Heather Allison — a likely new-owner story. Confirm before writing.","signalmountaincookielady.com","Some"),
 ("Conga (Latin Food)","Jimmy Obando (license)","Latin American / pupusas · 26 E Main St","Open — ownership change?","Established family spot ('the woman who brought pupusas to Chattanooga'); new license may signal a handoff — confirm.","congalatinfood.com","Some"),
 ("Riverside Beverage","David Powers","THC/hemp beverages (NOT a beer brewery) · old Hutton & Smith bldg, 3108 Riverside Dr","Open (2025)","Old brewery building reborn for the THC-drink boom — and squarely in the path of TN's July 1, 2026 hemp ban. Frame with care.","nationbeer.com","Low"),
 ("Raccoon Mountain Caverns","Chase Becker (Open Road Resorts)","Cavern attraction + campground (new owner) · 319 W Hills Dr","Open (acquired 2026)","New stewardship of a classic 5.5-mile cave system — Open Road Resorts' Tennessee debut.","raccoonmountain.com","Some"),
 ("The Chattery","Jennifer Holder","Nonprofit community classroom + art-supply thrift · 1800 Rossville Ave","Open (reopened Jan 2026)","Beloved nonprofit returned to its Rossville Ave home after a 2024 fire — comeback story (it's a nonprofit, not a typical thrift).","thechattery.org","Some"),
 ("Ace Pickleball Club","Jacobson (license: George; press: Don III)","Indoor pickleball (46k sq ft) · 5450 Hwy 153, Hixson","Coming (2026)","Local 4th-generation manufacturer brings a national franchise to a former Stockdale's. Verify the owner's first name before printing.","acepickleballclub.com","Some"),
 ("Wander Craft Co.","Jonathan Mellas (verify)","Maker — custom camper-van builds · Soddy Daisy","Active","Hands-on van-life builder converting vans (and vintage coffee trailers) into homes.","wandercraftcompany.com","Low"),
 ("Signal Modern Studio","Travis Randolph (verify)","Maker — custom modern furniture · Signal Mtn","Active","A mountain-town woodworker doing modern design.","@signalmodernstudio","Low"),
 ("Lookout Bricks","Samuel Mindeman","Maker — toy building blocks · Lookout Mtn (home)","Coming/likely","Custom building-brick maker; IG exists, content unconfirmed.","@lookout_bricks","Low"),
 ("Frog & Fern","Reagan Carson","Maker — 3D-printed sensory/fidget toys · markets/pop-ups, NorthShore","Open","Whimsical 'enchanted forest' 3D-printed toys; pickup at Ship Happens.","frogandfern.com","Low"),
 ("Tressa's Baking Company","Tressa Scott","Cottage bakery — custom cookies · Ooltewah (home)","Likely open","Custom-designed cookies with online ordering (Bakesy). Confirm identity (possible Ringgold rebrand).","tressasbakingcompany.com","Low"),
 ("Delhi Cafe","Robin Benjamin (license)","Indian restaurant · 6940 Lee Hwy","Open","Taking another swing at an 'Indian-restaurant graveyard' suite (Dosa Hut, Taj India both closed here) — a resilience angle.","none found","Low"),
 ("One Corner Bakery","Shan Sun","Coffee & bakery · 417 Georgia Ave (downtown)","Coming","Brand-new downtown coffee+bakery at a real storefront; no presence yet.","none found","None"),
 ("Chattanoodle","Brianne Larkins","Maker — fresh pasta · Signal Mtn (home)","Can't confirm","Appealing fresh-pasta concept; no business footprint yet — reach out.","none found","None"),
 ("Pan Y Mojo","Yanet Sosa Perez","Cuban restaurant · 5214 Dayton Blvd, Red Bank","Coming","A new Cuban spot at a real commercial address — adds to CHA's small Cuban scene.","none found","None"),
 ("Chattbout Jerk","Antole Thelwell","Jamaican catering · Ooltewah","Can't confirm","Clever name ('chat 'bout jerk'); no footprint — confirm before featuring.","none found","None"),
 ("Athara Grill","Muhammed Abazid","Food truck (cuisine unconfirmed) · Ooltewah","Can't confirm","Possible Middle Eastern/Syrian truck — verify the cuisine directly; don't assert.","none found","None"),
]
# ---- WHAT'S COMING: notable arrivals (kept separate from features) ----
COMING=[
 ("Hattie B's Hot Chicken","205 Manufacturers Rd (Signal Mill, NorthShore)","Nashville hot-chicken arrival; license issued June 2026 (moved from 'exploring' to real)."),
 ("Southern Steer Butcher","20 Cherokee Blvd (NorthShore)","National craft-butcher franchise. NOTE: the listed 'Kevin Gillespie' is almost certainly NOT the Atlanta celebrity chef — common name."),
 ("Bean Sprouts","321 Chestnut St","National kids/family café concept downtown."),
 ("Seven Brew / Scooter's Coffee / Dunkin / Starbucks NorthShore","various","Drive-thru/chain coffee build-out across the metro (context, not features)."),
 ("Hotels: Caption by Hyatt, Sonesta Select, Moxy, Home2/Tru","downtown / Hamilton Place / Lookout","Continued hotel growth — useful 'what's changing' context, not member features."),
 ("The Landing Social House (Common House)","1517 Mitchell Ave (former 1929 YMCA)","Established members' club, not a new independent — skip unless covering the rooftop."),
]
# ---- COLD LEADS / VERIFY-FIRST (licensed, little/no footprint or home-based) ----
COLD=[
 "Quirky watch: The Little Nursery (Harrison) — a 'reborn doll adoption experience'; potentially a great oddball feature IF it has a public component (home address — verify).",
 "Food/drink to confirm: Binder Tobacconist (514 E Main, cigars), Townsend St Wine (E Main), Cochise Coffee (Ooltewah), Vintage Coffee (Soddy Daisy), Shore Brew (Apison), Little Adobe Botanical Bakehouse (Signal Mtn, home).",
 "Makers to confirm: Nordic Metalworx (welding), New Forge Dog Biscuit Co (no footprint), Caelum Workshop (volumetric displays — intriguing, no footprint), Signal-area furniture/brick makers.",
 "Retail to confirm: Benched Bros Breakroom (hobby/cards), Sweetheart Yarns, Queen of Circa (vintage furniture, Lookout Mtn), The Green Flamingo Boutique (online/Ooltewah).",
 "Privacy-flagged (home addresses — do NOT publish, verify intent first): Chattanooga Beverage Co. (apartment), Dogleg LLC (Signal Mtn), plus the home-based makers/cottage bakers above.",
]
# ---- STORY THREADS (bigger than any one shop) ----
THREADS=[
 ("The Foundries District rises","The old U.S. Pipe/Wheland Foundry is becoming a food-and-drink district (adjacent to 'The Bend' riverfront redevelopment). Creature Comforts' beer garden is the first business at the Lookouts' new Erlanger Park stadium — a multi-part series as the district fills in."),
 ("The Kitchen Incubator of Chattanooga","5704 Marlin Rd keeps appearing because it's KIC — a 10,000 sq ft shared commercial kitchen and the region's largest food-truck commissary, launching immigrant and first-time food entrepreneurs. 'Food Truck Friday' the first Friday monthly. Profile the launchpad, then follow its graduates."),
 ("Zarzour's, 107 years later","Chattanooga's oldest restaurant (est. 1918) closed in 2025; Pauly's is taking the address. A then-and-now on the city's most storied room."),
 ("The trading-card boom","Roughly seven new sports-card/'breaks' shops licensed this year (Benched Bros, Cards2theMax, Anesthesia Baseball Cards, StoneGuys, WhatUpBaseball, Harjes CG, Freedom Cards). A trend feature on a hobby's local resurgence."),
 ("The July 1 hemp reckoning","TN's hemp ban (effective July 1, 2026) threatens a wave of just-licensed local makers — Riverside Beverage, Dogleg, Snapdragon Cannabis Co. 'What the ban means for Chattanooga's hemp-beverage startups.'"),
]
# ---- DOCTRINE v2 additions ----
DOCTRINE=[
 "**Mine the license list every month.** The county's new-business licenses (via the Times Free Press weekly roundup) are the earliest possible signal — by definition press footprint = None. This is the highest-scoop pipeline we have. Triage it on a set cadence.",
 "**Triage hard.** KEEP consumer-facing independents — food/drink, makers/artisans, retail/boutiques, experiences, and genuinely quirky concepts. SKIP the noise — contractors/trades, cleaning/janitorial, lawn care, logistics/trucking, vending, parking, real-estate/STVR, finance/consulting/notary, and national chains (route chains to 'What's Coming,' not features).",
 "**A license is intent, not a grand opening.** Verify a business is actually real and open before featuring; reach out directly (DM / call / visit). Expect ~zero web presence — that's the point.",
 "**Trust the name, not the form fields.** Product descriptions can be wrong (Riverside's 'beer brewery' is actually THC beverages). The owner field may reflect a registration or ownership change (Cookie Lady, Conga) — which is often the better story.",
 "**Protect privacy.** A huge share are home-based; never publish a home address. Confirm a public storefront or contact, and respect cottage/maker operators' privacy.",
 "**Dodge name-collision traps.** Don't attach an unrelated website or a celebrity to a licensee (Windy Hill Pottery = a Cleveland studio; Skyforge Robotics ≠ the drone company; 'Kevin Gillespie' ≠ the Top Chef; Queen of Circa ≠ Circa Chattanooga).",
 "**Hunt for threads, not just dots.** Recurring addresses (the Kitchen Incubator), new districts (Foundries), category booms (card shops), regulatory shifts (the hemp ban), and ownership changes become features bigger than any single shop.",
 "**Run recurring segments off the pipeline:** 'Just Licensed / Coming Soon' (early scoop), 'Now Open' (verified), 'What's Coming' (arrivals), and the occasional trend/thread feature.",
]

# ===================== CSV =====================
def all_rows():
    rows=[]
    for r in TIER1: rows.append(("1 - Feature now",)+r)
    for r in TIER2: rows.append(("2 - Watch/verify",)+r)
    return rows
with open("/home/user/prscott/new_openings_watchlist.csv","w",newline="") as f:
    w=csv.writer(f)
    w.writerow(["Tier","Business","Owner","What & where","Status","Story angle","Social/Web","Press footprint"])
    for r in all_rows(): w.writerow(r)
    for (biz,where,note) in COMING: w.writerow(["What's coming",biz,"",where,"Arrival",note,"",""])
print("CSV rows:",len(all_rows())+len(COMING))

# ===================== DOCX =====================
doc=Document()
sec=doc.sections[0]; sec.orientation=WD_ORIENT.LANDSCAPE
sec.page_width,sec.page_height=Inches(11),Inches(8.5)
for m in ("top_margin","bottom_margin","left_margin","right_margin"): setattr(sec,m,Inches(0.6))
nm=doc.styles["Normal"]; nm.font.name="Calibri"; nm.font.size=Pt(10); nm.paragraph_format.space_after=Pt(4)

def scolor(r,h): r.font.color.rgb=RGBColor.from_string(h)
def shade(c,h):
    p=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd")
    s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); p.append(s)
def ctext(c,t,bold=False,color=None,size=8.3,white=False,align="left"):
    c.text=""; p=c.paragraphs[0]
    p.alignment={"left":WD_ALIGN_PARAGRAPH.LEFT,"center":WD_ALIGN_PARAGRAPH.CENTER}[align]
    r=p.add_run(t); r.font.size=Pt(size); r.font.bold=bold
    if white: r.font.color.rgb=RGBColor.from_string("FFFFFF")
    elif color: r.font.color.rgb=RGBColor.from_string(color)
def H(t,level=1,color=None):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(12 if level==1 else 8); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(t); r.font.bold=True
    if level==1:
        r.font.size=Pt(15); scolor(r,color or NAVY)
        pPr=p._p.get_or_add_pPr(); b=OxmlElement("w:pBdr"); bt=OxmlElement("w:bottom")
        bt.set(qn("w:val"),"single"); bt.set(qn("w:sz"),"6"); bt.set(qn("w:space"),"2"); bt.set(qn("w:color"),TEAL)
        b.append(bt); pPr.append(b)
    else: r.font.size=Pt(12); scolor(r,color or TEAL)
def P(t,italic=False,color=None,size=10):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(5)
    for i,ch in enumerate(t.split("**")):
        if not ch: continue
        r=p.add_run(ch); r.font.size=Pt(size); r.font.italic=italic
        if i%2==1: r.font.bold=True
        if color: scolor(r,color)
def BUL(items,size=9.5):
    for it in items:
        p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(3)
        for i,ch in enumerate(it.split("**")):
            if not ch: continue
            r=p.add_run(ch); r.font.size=Pt(size)
            if i%2==1: r.font.bold=True

def wtable(rows, fill_first=SAND):
    cols=["Business","Owner","What · Where","Status","Story angle","Press"]
    widths=[1.35,1.45,1.95,1.15,3.55,0.55]
    t=doc.add_table(rows=1,cols=len(cols)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style="Table Grid"
    for i,h in enumerate(cols):
        shade(t.rows[0].cells[i],NAVY); ctext(t.rows[0].cells[i],h,bold=True,white=True,size=8.7,align="left" if i<5 else "center")
    for r,(biz,owner,where,status,story,social,press) in enumerate(rows):
        cells=t.add_row().cells; base=LIGHT if r%2==0 else "FFFFFF"
        ctext(cells[0],biz,bold=True,color=NAVY,size=8.3)
        ctext(cells[1],owner,size=8.0)
        ctext(cells[2],where,size=7.9)
        sc_={"Open":GREEN}.get(status.split()[0],SLATE)
        ctext(cells[3],status,size=7.9,color=(GREEN if status.startswith("Open") else (RED if "Can't" in status else SLATE)),bold=True)
        ctext(cells[4],story,size=8.1)
        pc={"None":GREEN,"Low":GOLD,"Some":ORANGE}.get(press,SLATE)
        ctext(cells[5],press,size=8.0,align="center",color=pc,bold=True)
        for c in cells: shade(c,base)
    for i,wd in enumerate(widths):
        for row in t.rows: row.cells[i].width=Inches(wd)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)

# Title
for _ in range(2): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("THE SCENIC CITY SCOUT"); r.font.size=Pt(24); r.font.bold=True; scolor(r,NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("New-Openings Watchlist"); r.font.size=Pt(18); scolor(r,TEAL)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Mined from Hamilton County business licenses — the earliest, highest-scoop pipeline (press footprint = None by design)")
r.font.size=Pt(11.5); r.font.italic=True; scolor(r,SLATE)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Triaged from ~600 new licenses · verified by scouts · June 2026 · Press footprint: None (biggest scoop) / Low / Some")
r.font.size=Pt(9.5); scolor(r,GREY)
doc.add_page_break()

H("How this list was built",1)
P("We pulled this year's Hamilton County new-business licenses (~600), threw out the ~85% that don't matter to members (contractors, cleaning, logistics, vending, parking, real-estate/STVR, consulting, national chains), and sent the on-brand independents to scouts to verify which are actually open and find the person + story. A license is an **intent** signal, not a grand opening — so each entry is marked Open / Coming / Can't confirm, and home-based operations are flagged (we never publish a home address).")

H("★ Tier 1 — Feature now (verified, on-brand, story-rich)",1,color=GOLD)
wtable(TIER1)

H("Tier 2 — Watch & verify",1)
wtable(TIER2)

H("What's Coming — notable arrivals (context, not features)",1)
t=doc.add_table(rows=1,cols=3); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style="Table Grid"
for i,h in enumerate(["Arrival","Where","Note"]):
    shade(t.rows[0].cells[i],NAVY); ctext(t.rows[0].cells[i],h,bold=True,white=True,size=8.7)
for r,(biz,where,note) in enumerate(COMING):
    cells=t.add_row().cells; base=LIGHT if r%2==0 else "FFFFFF"
    ctext(cells[0],biz,bold=True,color=NAVY,size=8.3); ctext(cells[1],where,size=8.1); ctext(cells[2],note,size=8.1)
    for c in cells: shade(c,base)
for i,wd in enumerate([2.0,2.0,5.5]):
    for row in t.rows: row.cells[i].width=Inches(wd)
doc.add_paragraph().paragraph_format.space_after=Pt(2)

H("Cold leads & verify-first (licensed, little/no footprint)",1,color=ORANGE)
BUL(COLD)

H("Story threads — bigger than any one shop",1)
for title,body in THREADS:
    P(f"**{title}.** {body}")

doc.add_page_break()
H("Updated Scout Sourcing Doctrine v2 — Mining the License List",1)
P("The original doctrine (Addendum: Hidden Gems) said avoid the obvious, prize a low press footprint, find the person and the secret, and verify in person. License-mining extends it with eight rules:")
BUL(DOCTRINE)
P("Net effect: the license list becomes a renewable, defensible discovery engine. We meet new independents the week they file — before any competitor knows they exist — verify, and tell their story first. That is the press-footprint-of-None scoop the whole brand is built on.",
  italic=True,color=SLATE,size=9.5)

out="/home/user/prscott/New_Openings_Watchlist.docx"
doc.save(out); print("saved",out)
