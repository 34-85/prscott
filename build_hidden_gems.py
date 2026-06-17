#!/usr/bin/env python3
"""Build the 'Hidden Gems & Wider Territory' scout source book (Word) + CSV."""
import csv
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY="1B3A4B"; TEAL="2A9D8F"; ORANGE="E76F51"; SAND="E9C46A"; SLATE="415A6B"
LIGHT="EEF3F4"; GREY="6B7B85"; GOLD="B8860B"; GREEN="2E7D52"

# (business, person, category, town, secret, source, press, conf, pick)
DATA = {
"The Ridge & Red Bank — Signal Mtn · Lookout Mtn · Walden · Red Bank":[
 ("Mountain Memories","Jeanne Crawford — owner","Souvenir/history shop","Lookout Mtn, TN","Granddaughter of Leo Lambert, who discovered Ruby Falls in 1928; she runs a 1928 'halfway' horse-watering shop and tells carriage-era stories tourists speeding to Rock City never hear.","local3news / noogatoday","Low","High",True),
 ("Automotive Artifex","Douglas Swafford Jr. — owner","Classic VW restoration","Red Bank","A full-time schoolteacher who restores air-cooled Beetles and dune buggies on the side — and in 2023 stopped taking new work to finish the cars already in his shop rather than turn customers away.","facebook / bbb.org","Low","Med",True),
 ("Red Bank Barber Shop","Pamela Sanders & Michelle Fitzsimmons — sisters/owners","Barbershop","Red Bank","Two sisters (one barbered 13 years on a U.S. Air Force base) whose walls hold a collection of 400+ hats.","timesfreepress","Low","High",False),
 ("Mayfly Coffee","Atley Davidson — founder/roaster","Coffee roaster/café","Signal Mountain","Named for his love of fly fishing; he started at the Signal Mtn farmers market and names menu items after local hiking trails.","mayflycoffee.com","Low","High",False),
 ("Walden Peak Farm","Andrea Zoppo ('Ms. Ladybug') & Ryan Welch","U-pick berry farm","Walden","A 100+ year-old Welch family farm; co-owner Andrea is a children's entertainer named a PBS Early Learning Champion — one of just 17 in the U.S.","waldenpeakfarm.com","Low","High",False),
 ("Terra Flower Farm","Kelly Garcia — farmer-florist","Cut-flower farm","Signal Mountain","Trained under legendary UGA horticulturist Dr. Allan Armitage before going solo with daily direct-to-florist deliveries.","terraflowerfarm.com","Low","Med",False),
 ("Will Morrison Construction","Will & Kate Morrison — owners","Custom home builder","Signal Mountain","A third-generation, lifelong Signal Mtn builder running a two-person family shop that specializes in the tricky craft of building on steep mountain terrain.","willmorrisonconstruction.com","Low","High",False),
 ("Fairyland Pharmacy","The Voges family (Alan & Gail; now Dusty & Jessica)","Independent pharmacy","Lookout Mtn, GA","Founder Alan was known to unlock the doors at night and trudge through snow to fill emergency prescriptions; now second-generation family-run.","fairylandpharmacy.com","Low","High",False),
 ("Rallo's Bar & Grill","Ron Potter — founder/owner","Locals' bar/diner","Signal Mountain","A low-key locals' bar whose real secret is the Saturday-morning brunch spread that starts at 9 a.m.","rallosbar.com","Low","Med",False),
],
"Cleveland, Bradley & Polk County":[
 ("Smalley's Shoe & Boot Repair","Tim Smalley — owner","Cobbler / leather repair","Cleveland","A one-man cobbler quietly resoling boots, belts and purses since 1990 — no website, a near-extinct trade.","bbb.org","None","High",True),
 ("The Village Bake Shop","Teresa Gilbert — owner/baker","Bakery","Cleveland","Founded by her parents in 1961; their hand-indented 'thumbprint' cookies now ship to 500+ addresses worldwide from a strip-center bakery.","thevillagebakeshop.com","Low","High",False),
 ("Presswood's Vintage Antiques","Walter & Sheila Presswood — owners","Antiques","Cleveland","A couple filling two buildings on N. Ocoee St with 30+ years of their own personally-collected finds — not mall-booth consignment.","yelp","None","Med",False),
 ("Grindstone Gardens","David Rogers — manager","Farm / produce co-op","McDonald","Began as an honor-system self-serve produce stand by the road in 2008; now a no-spray farmer cooperative stocking year-round local goods.","knoxvillevoyager","Low","High",False),
 ("Mash & Hops","'Rob' — owner","Craft beer café","Cleveland","Dreamed up while backpacking past a Seattle world-beer bar; opened 2014 as downtown Cleveland's first craft beer.","mashandhops.com","Low","Med",False),
 ("Burgess Hardware & Feed","Stephanie Everett — store mgr (4th gen)","Feed & hardware","Benton","Four generations deep (founded 1968); the only STIHL Elite Dealer in SE Tennessee, with the great-grandkids now on the sales floor.","burgesshardware.com","None","High",True),
 ("Webb Brothers Float Service & General Store","The Webb family (since 1936)","Outfitter / general store","Reliance","A 1936 Depression-era general store; the owner became postmaster so the store became the post office, then pivoted to river rafting in 1969. On the National Register.","webbbros.com","Some","High",True),
 ("Wicked Wood Shop","Joshua Morrison — owner","Custom woodworker","Cleveland","A small, eccentrically-named shop (est. 2019) turning out mantels, range hoods and one-off furniture — found mainly via Facebook.","facebook","None","Med",False),
 ("El Dorado Pupuseria","owner unverified","Salvadoran diner","Cleveland","A humble from-scratch pupusa spot on Spring Place Rd, ~4.8 stars on a tiny review count — the classic word-of-mouth immigrant kitchen.","tripadvisor","None","Med",False),
],
"Dalton & Whitfield County, GA":[
 ("La Esperanza Bakery","Delia & Jorge Lara — owners","Panaderia / Mexican bakery","Dalton","Delia knew nothing about baking when she co-founded this 2000 panaderia 'on faith it would grow'; now the go-to for quinceañera cakes, with piñatas in back.","dailycitizen.news","Low","High",True),
 ("Garmony House Coffee & Cocktails","Jillian Bearden — owner","Coffee + cocktails","Dalton","Named for her family's home on Scotland's Isle of Mull; a one-woman downtown café built around the Scottish coffee-and-tea ritual she grew up with.","garmonyhousecoffee","Low","High",True),
 ("Arcadia Coffee Roasters","Win & Asher — brothers/roasters","Coffee roaster","Dalton","NW Georgia's first local roaster, run by two schoolteacher brothers who learned sourcing during time in Kenya and El Salvador.","arcadia.coffee","Low","High",False),
 ("Cohutta Auctions","Tim Farner — owner/auctioneer","Auction house","Tunnel Hill","A 20-year Air Force vet who fell for auctions while stationed in England, got licensed after retiring, and runs antiques-and-'rusty gold' auctions twice a month.","chamberofcommerce","None","Med",False),
 ("Mi Cafecito","owner unverified","Latino café","Dalton","The secret isn't lattes — it's traditional champurrado and atole and homemade syrups, an under-the-radar morning ritual on E. Morris St.","joe.coffee","Low","Med",False),
 ("La Bendicion Pupuseria","owner unverified","Salvadoran restaurant","Dalton","An unassuming spot serving hand-made pupusas and yuca frita — Dalton's lesser-known Salvadoran thread, not just Mexican.","restaurantji","Low","Med",False),
 ("Tacos Tu Go","owner unverified","Taqueria","Dalton","A 20-year hole-in-the-wall where you watch gorditas hand-pressed; the off-menu draw is barbacoa and lengua gorditas for those in the know.","tripadvisor","Low","Med",False),
 ("Hammontree Tufters","family — name unverified","Hand-tufting textile craft","Cohutta","A rare surviving hand-tufter — a living link to Dalton's pre-carpet chenille-bedspread craft ('Peacock Alley'). Verify still active.","yelp","None","Low",False),
],
"Walker & Catoosa County, GA":[
 ("Droop Scoops","Marla Prince — owner","Ice cream / hot dog shop","Chickamauga","Named for her dad's school nickname 'Droop'; third generation in town, she ships in hand-dipped ice cream from a Michigan creamery.","facebook / wdef","Low","High",True),
 ("Warehouse 71","Jason & Michelle Hughett — owners","Antique / vintage shop","Ringgold","A couple married 25+ years set up shop inside an old Sinclair gas station and document their thrift hunts on a YouTube channel.","warehouse71","Low","High",False),
 ("The Great Awakening Coffee & Doughnuts","Dane & Vanessa Maddux — owners","Coffee shop","LaFayette","Sold coffee out of a truck across the country before landing a brick-and-mortar on the LaFayette square.","thegreatawakeningcoffeeco","Low","High",True),
 ("House of Harlow Antiques","Lorie McCamy Headrick — owner","Antiques","Ringgold","A small-town antique shop themed entirely around 1930s screen siren Jean Harlow — old-Hollywood glamour in North Georgia.","timesfreepress","Low","High",False),
 ("Susan's Sweets","Susan — surname unverified","Bakery","Fort Oglethorpe","A tiny shop on handed-down recipes that locals swear makes 'the best fried pies in the Southeast,' open only a few afternoons a week.","facebook / yelp","None","Med",False),
 ("High Point Farms","The Gardner family (Margie & David; daughter Jill Baldschun)","U-pick flower farm","Flintstone","A 75-acre family farm (50+ years); daughter Jill grows 200+ flower varieties and runs 'Pick-N-Sip' days with baby goats under Lookout Mtn.","highpointfarms.co","Some","High",False),
 ("Backyard Shack Creations","'Anja' — owner, unverified","Ceramics studio","LaFayette","A walk-in hand-built pottery studio tucked on a rural road — the kind of maker space you'd never find without a tip.","facebook","None","Med",False),
 ("The Tap Room","owner unverified","Beer bar","LaFayette","A county-seat square bar pouring international rarities (Belgian Lindemans, Athens' Terrapin) plus artisan flatbreads and gourmet dogs.","lafayettetravel","Low","Med",False),
],
"Sequatchie Valley & Northern Exurbs":[
 ("Cookie Jar Cafe","Sue Ann & Debra Lockhart — owners","Country diner / bakery","Dunlap","Started in 2002 as a cookie-and-pie stand to save the family's ~150-year-old dairy farm; now a from-scratch farmhouse restaurant with a petting zoo.","tennesseecrossroads","Some","High",True),
 ("Sequatchie Valley Institute / Moonshadow","The Kimmons family (Johnny & Carol; sons Joel & Patrick)","Off-grid arts & education homestead","Whitwell","A back-to-the-land family homestead since 1971, fully off-grid solar with hand-built cob houses and a hidden artists' gallery deep in the woods.","svionline.org","Some","High",True),
 ("Yum Yum's Takeout Cafe","Brittany Garrett — founder","Soul food","South Pittsburg","A 'struggle to success' founder running a soul-food café with a team of six women and a community-giving mission.","mariontribune","Low","High",True),
 ("Floyd's Jamaican Jerk Chicken","owner from St. Catherine, Jamaica — name unverified","Jamaican restaurant","South Pittsburg","Authentic island jerk in a cast-iron town best known for Lodge — cooked by a St. Catherine, Jamaica native.","floydsjamaicanjerk","Low","Med",False),
 ("Castle's Grocery","Castle family — names unverified","General store","Whitwell","A family-run Highway 28 grocery since 1974 — a genuine 'everybody shops here' institution with almost no press.","yelp","None","Med",False),
 ("Frazier Produce","Robert Leo Frazier — owner","Roadside produce stand","Sale Creek","A two-person Dayton Pike stand (since ~2005) selling produce, regional honey, Pepsi and Little Debbies — $3 strawberries and small-town charm.","yelp","None","Med",False),
 ("The Dunlap Mercantile","owner unverified","Vintage boutique","Dunlap","Housed in the oldest building in town (1894), with original 'petticoat' stairs and a working soda fountain.","dunlapmercantile","Low","Med",False),
 ("Glen Brooks Custom Knives","Glen Brooks — bladesmith","Bladesmith","Soddy-Daisy (confirm)","Got back into knives as a hobby on his father-in-law's ranch, was mentored by veteran makers, and went full-time hand-grinding blades.","glenbrooksknives","None","Med",False),
 ("South Pittsburg Antiques","owner unverified","Antiques curator","South Pittsburg","A single-owner trove of museum-scale oddities — 12-foot stained-glass windows, bronze chandeliers — quietly in the shadow of Lodge.","yelp","Low","Med",False),
],
"Chattanooga's Overlooked Neighborhoods":[
 ("6 Ear Knife & Tool","Mak Kelsay — bladesmith","Knifemaker","Ooltewah","A full-time Chattanooga firefighter who hand-forges high-end chef's knives in his garage on his days off — and competed on TV's 'Forged in Fire.'","local3news / 6earknives","Low","High",True),
 ("Tienda Maya #5","owner unverified","Guatemalan grocery / butcher","Brainerd","Out front, a vendor sells homemade Guatemalan-style ice cream from a miniature fire-truck freezer; inside, a butcher counter and banana leaves bundled for tamales.","chattanoogapulse","Low","High",True),
 ("Tienda y Carnicería Sontay","Wilfredo Sontay Chavez — owner","Guatemalan grocery / butcher","East Ridge","A carnicería that doubles as the neighborhood's Central-American hub — cut-to-order meat, money transfers, and bill pay under one roof.","chamberofcommerce","None","Med",False),
 ("Pupuseria Marelyn","'Marelyn' — family, unverified","Salvadoran restaurant","East Ridge","A tiny family spot hand-patting pupusas and cow-tongue tacos on the Ringgold Rd strip — loyal regulars, almost no press.","restaurantji","Low","Med",False),
 ("World Tour Market","Hamid Merabti — owner","International market","Brainerd","A small Middle-Eastern/Central-European market — fresh halal meat, Turkish coffee, baklava — found only by word of mouth.","yelp","None","Med",False),
 ("Federal Bake Shop","The Ainsworth family (since 1974)","Bakery","Hixson","The last surviving outpost of a California bakery chain founded in 1921; the Ainsworths still bake from the original 1920s recipes.","federalbakeshop","Some","High",True),
 ("The Farm at the Beth","Dr. Reginald Floyd Smith II — exec director","Urban farm","Alton Park","Built on a former basketball court at a faith-based center; ~60% of the harvest is given away free to families in a food desert.","chattanoogan / wutc","Some","Med",False),
 ("Grow Hope Farm","Joel Tippens — founder","Urban youth farm","South Chattanooga","A youth-development gardener teaching inner-city kids self-sufficiency through crops on Hope for the Inner City land.","chattanoogapulse","Low","Med",False),
 ("Grindhead Coffee Co.","'Olivia' / crew — unverified","Coffee roaster / café","East Brainerd","A Caribbean-and-pirate-themed roastery where the staff are 'the crew,' the owner is 'Harbor Master,' and the island music never stops.","visitchattanooga","Low","Med",False),
 ("Irie Jamaican Cuisine","Robert & Nadine Officer — owners","Jamaican restaurant","Red Bank","A couple who immigrated from Jamaica and opened up only because friends kept begging after home-cooked meals; 'Irie' = everything's alright.","local3news / tfp","Some","High",False),
],
}

FAMOUS = [
 ("Calliope","Khaled Albanna — chef/co-owner","Downtown",
  "Drop the 'NYT Top 50' line everyone uses. The Scout angle: the civil-engineering student who fell for cooking and now bridges Amman and Appalachia on one plate — built around his family's recipes, the local farms he sources from, and the line cooks he's training."),
 ("Little Coyote","Erik & Amanda Niel — owners","St. Elmo",
  "Skip the Michelin Bib. The Scout angle: spotlight Amanda Niel (the business partner usually in the chef's shadow) and the masa — the daily nixtamal-and-tortilla ritual and the person at the comal — as a Texas 'love letter' to St. Elmo, not an award."),
]

RESERVE = [
 "Dalton Brewing Co. (engineers-turned-brewers) & Prater's Mill (1855 gristmill, closed 2023) — more press / less 'hidden.'",
 "Little Diner on 1st (Cleveland), Creekside Flower Farm (Chickamauga), Brown Dirt Farm (Whitwell) — owner names unverified.",
 "A Catoosa/Walker woodworking artist profiled by NW Georgia News — chase the name by phone.",
 "Out-of-territory backups if you widen the map: Oren Wooden's Apple House (Pikeville), White Oak Harness (Spencer, plain-community), Dixie Soaps / Dixie-Does Alpines goat-milk soap (Dayton), Ocoee Outdoors & Hiwassee River Guides (Reliance/Benton).",
]
FLAGS = [
 "Confirm owner names/spelling by phone before printing: Mi Cafecito, La Bendicion, Tacos Tu Go, El Dorado, Floyd's, Castle's, Dunlap Mercantile, South Pittsburg Antiques, Pupuseria Marelyn, Tienda Maya #5, Grindhead, Backyard Shack, Susan's Sweets, The Tap Room, Mash & Hops ('Rob').",
 "Confirm still operating / location: Hammontree Tufters, Automotive Artifex (wound down new work in 2023), Glen Brooks Knives (town), Presswood's.",
 "Scouting hazard caught & excluded: Cleveland, OHIO impostors; Red Bank, NEW JERSEY impostors; a Connecticut 'Highland Park Market'; an Ottawa market. Always confirm it's actually in our footprint.",
 "No names were invented; every 'unverified' is labeled. Thin web presence is expected (and good) — verify in person.",
]
DOCTRINE = [
 "**Avoid the obvious.** If it has national press or everyone already covers it, it isn't a Scout story — unless we find a genuinely new angle.",
 "**Press footprint is a feature, not a bug.** Prize None/Low. Use a 'Some' pick only when the *secret* is still untold.",
 "**Go where others won't.** The satellite towns and the overlooked neighborhoods — not just Southside/NorthShore/St. Elmo.",
 "**Lead with the person and the secret.** A surprising, specific, human detail — never 'award-winning.'",
 "**Earn the reader a discovery.** They should feel let in on something and want to forward it.",
 "**Verify in person.** Hidden gems have thin web footprints; confirm names by phone or visit. Never invent a name.",
 "**Even the famous get the Scout treatment.** Keep a known name only with an untold angle (see Calliope & Little Coyote).",
]

# ---------- CSV ----------
with open("/home/user/prscott/hidden_gems_source_list.csv","w",newline="") as f:
    w=csv.writer(f)
    w.writerow(["Territory","Business","Person","Category","Town","Hidden-gem secret","Source","Press footprint","Confidence","Scout pick"])
    for terr,rows in DATA.items():
        for (biz,person,cat,town,secret,src,press,conf,pick) in rows:
            w.writerow([terr.split(" — ")[0],biz,person,cat,town,secret,src,press,conf,"Yes" if pick else ""])
total=sum(len(r) for r in DATA.values())
print("CSV rows:",total)

# ---------- DOCX ----------
doc=Document()
sec=doc.sections[0]; sec.orientation=WD_ORIENT.LANDSCAPE
sec.page_width,sec.page_height=Inches(11),Inches(8.5)
for m in ("top_margin","bottom_margin","left_margin","right_margin"): setattr(sec,m,Inches(0.6))
nm=doc.styles["Normal"]; nm.font.name="Calibri"; nm.font.size=Pt(10); nm.paragraph_format.space_after=Pt(4)

def sc(r,h): r.font.color.rgb=RGBColor.from_string(h)
def shade(c,h):
    p=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd")
    s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); p.append(s)
def ctext(c,t,bold=False,color=None,size=8.4,white=False,align="left"):
    c.text=""; p=c.paragraphs[0]
    p.alignment={"left":WD_ALIGN_PARAGRAPH.LEFT,"center":WD_ALIGN_PARAGRAPH.CENTER}[align]
    r=p.add_run(t); r.font.size=Pt(size); r.font.bold=bold
    if white: r.font.color.rgb=RGBColor.from_string("FFFFFF")
    elif color: r.font.color.rgb=RGBColor.from_string(color)
def H(t,level=1,color=None):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(12 if level==1 else 8); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(t); r.font.bold=True
    if level==1:
        r.font.size=Pt(15); sc(r,color or NAVY)
        pPr=p._p.get_or_add_pPr(); b=OxmlElement("w:pBdr"); bt=OxmlElement("w:bottom")
        bt.set(qn("w:val"),"single"); bt.set(qn("w:sz"),"6"); bt.set(qn("w:space"),"2"); bt.set(qn("w:color"),TEAL)
        b.append(bt); pPr.append(b)
    else: r.font.size=Pt(12); sc(r,color or TEAL)
def P(t,italic=False,color=None,size=10):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(5)
    for i,ch in enumerate(t.split("**")):
        if not ch: continue
        r=p.add_run(ch); r.font.size=Pt(size); r.font.italic=italic
        if i%2==1: r.font.bold=True
        if color: sc(r,color)
def BUL(items,size=9.5):
    for it in items:
        p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(2)
        for i,ch in enumerate(it.split("**")):
            if not ch: continue
            r=p.add_run(ch); r.font.size=Pt(size)
            if i%2==1: r.font.bold=True

def gem_table(rows):
    cols=["Business","Person & role","What · Where","The secret (the hidden-gem hook)","Press","Conf"]
    widths=[1.25,1.55,1.6,4.1,0.6,0.5]
    t=doc.add_table(rows=1,cols=len(cols)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style="Table Grid"
    for i,h in enumerate(cols):
        shade(t.rows[0].cells[i],NAVY); ctext(t.rows[0].cells[i],h,bold=True,white=True,size=8.8,align="left" if i<4 else "center")
    for r,(biz,person,cat,town,secret,src,press,conf,pick) in enumerate(rows):
        cells=t.add_row().cells; fill=SAND if pick else (LIGHT if r%2==0 else "FFFFFF")
        ctext(cells[0],biz+("  ★" if pick else ""),bold=True,color=NAVY,size=8.4)
        ctext(cells[1],person,size=8.2)
        ctext(cells[2],f"{cat} · {town}",size=8.0)
        ctext(cells[3],secret,size=8.2)
        pc={"None":GREEN,"Low":GOLD,"Some":ORANGE}.get(press,SLATE)
        ctext(cells[4],press,size=8.2,align="center",color=pc,bold=True)
        ctext(cells[5],conf,size=8.2,align="center")
        for c in cells: shade(c,fill)
    for i,wd in enumerate(widths):
        for row in t.rows: row.cells[i].width=Inches(wd)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)

# Title
for _ in range(2): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("THE SCENIC CITY SCOUT"); r.font.size=Pt(24); r.font.bold=True; sc(r,NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Hidden Gems & Wider Territory"); r.font.size=Pt(18); sc(r,TEAL)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run(f"{total} under-the-radar businesses and the people behind them — from the ridges to Cleveland, Dalton, LaFayette & the valley"); r.font.size=Pt(11.5); r.font.italic=True; sc(r,SLATE)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Scouted June 2026  ·  ★ = Scout's Pick  ·  Press footprint: None (biggest scoop) / Low / Some  ·  Full sources in the CSV"); r.font.size=Pt(9.5); sc(r,GREY)
doc.add_page_break()

# Doctrine
H("The Scout's Sourcing Doctrine",1)
P("A Scout's job is **discovery**, not coverage. Anyone can profile the award-winners; our moat is finding the diner in LaFayette, the cobbler in Cleveland, the potter on Signal Mountain that nobody else has written up. The rules:")
BUL(DOCTRINE)
P("**Press footprint** is rated for every find — None means a true scoop, Some means it's been lightly covered and we'd need a fresh angle. Lower is better.", color=SLATE)

# Scout's Picks shortlist (no duplicate rows)
picks=[(b,p2,terr.split(' — ')[0]) for terr,rows in DATA.items() for (b,p2,c,tw,se,sr,pr,cf,pk) in rows if pk]
H("★ Scout's Picks — start with these",1,color=GOLD)
P("The richest, most verifiable human hooks across the territory — ideal for the first features. Each is marked ★ in its territory table below.")
BUL([f"{b} — {p2}  ·  {terr}" for (b,p2,terr) in picks])

# Famous, re-angled
H("Giving the famous the Scout treatment",1)
P("Two nationally-recognized names worth keeping — but only with an untold angle, never the press-release accolade:")
t=doc.add_table(rows=1,cols=3); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style="Table Grid"
for i,h in enumerate(["Business","Where","The fresh, non-award angle"]):
    shade(t.rows[0].cells[i],NAVY); ctext(t.rows[0].cells[i],h,bold=True,white=True,size=9)
for r,(biz,person,town,angle) in enumerate(FAMOUS):
    cells=t.add_row().cells; fill=LIGHT if r%2==0 else "FFFFFF"
    ctext(cells[0],f"{biz}\n{person}",bold=True,color=NAVY,size=8.6); ctext(cells[1],town,size=8.6); ctext(cells[2],angle,size=8.6)
    for c in cells: shade(c,fill)
for i,wd in enumerate([1.6,1.0,7.0]):
    for row in t.rows: row.cells[i].width=Inches(wd)
doc.add_paragraph().paragraph_format.space_after=Pt(4)

# Territory sections
for terr,rows in DATA.items():
    H(terr+f"  ({len(rows)})",1)
    gem_table(rows)

# Reserve / flags
H("Reserve bench & out-of-territory backups",1)
BUL(RESERVE)
H("Verify before print (the Scout's discipline)",1,color=ORANGE)
BUL(FLAGS)
P("Full source for every entry is in hidden_gems_source_list.csv. Hidden gems shift — confirm hours, ownership, and the human hook in person before you publish.",
  italic=True,color=GREY,size=9)

out="/home/user/prscott/Hidden_Gems_Source_Book.docx"
doc.save(out); print("saved",out,"| picks:",len(picks))
