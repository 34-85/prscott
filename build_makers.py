#!/usr/bin/env python3
"""Build the 'Meet a Maker' Source Book (Word) + CSV from the research."""
import csv
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY="1B3A4B"; TEAL="2A9D8F"; ORANGE="E76F51"; SAND="E9C46A"; SLATE="415A6B"
LIGHT="EEF3F4"; GREY="6B7B85"; GOLD="B8860B"

# Each: (business, person, area, angle, source, confidence, flagship)
DATA = {
"Restaurants & Chef-Owners":[
 ("Calliope","Khaled Albanna — chef & co-owner","Downtown/Southside","Amman-born, came to UTC for civil engineering, switched to cooking; modern Levantine-meets-South named among the NYT's top 50 U.S. restaurants.","https://www.callioperestaurant.com/about","High",True),
 ("Alleia","Daniel Lindley — chef-owner","Southside (Main St)","Native Chattanoogan who started at 17 as a dishwasher, cooked around the world, opened Alleia in 2009; multiple James Beard nominee.","https://alleiarestaurant.com/about-2/","High",True),
 ("Little Coyote","Erik & Amanda Niel — owners","St. Elmo","Chef Niel's Texas 'love letter' — wood-pit smoke and fresh tortillas; Michelin Bib Gourmand. (Same owners as Easy Bistro & Main Street Meats.)","https://littlecoyote.com/st-elmo-restaurant-chattanooga-tennessee","High",False),
 ("Hello Monty","Rob & Clay Gentry — owners (brothers)","Southside (Main St)","A veteran restaurateur and an award-winning brewmaster, lifelong locals, opened a wood-fired spot named for a beloved local booster.","https://www.hellomontyonmain.com/our-story","High",False),
 ("Boathouse Rotisserie & Raw Bar","Lawton & Karen Haygood — owners","Northshore (Riverside Dr)","Found a riverfront lot, tore it down, and 'spent all the money they had' to build a Tennessee River institution.","https://www.boathousechattanooga.com/about-us/","High",False),
 ("1885 Grill","Miguel Morales — owner","St. Elmo / multiple","SC-raised son of a Puerto Rican father; coastal-Southern-meets-Spanish; known for saving ~460 jobs through the pandemic. (Verify current address.)","https://1885grill.com/our-story","High",False),
 ("The Tap House","Chris Calhoun — owner","St. Elmo","A 20-year Marine Corps veteran who came home to open a veteran-owned beer bar and kitchen (2017).","https://taphousechatt.com/chattanooga-chattanooga-the-tap-house-about","Med",False),
 ("Elsie's Daughter","Chloe Wright & Ryan Smith — owners","Southside (Choo Choo)","Wright's father ran a restaurant on the Choo Choo campus in the '80s; she performed there as a teen — now her own bistro/cocktail bar (Smith built the bar program after NYC's The Wren).","https://www.elsiesdaughters.com/","High",False),
],
"Bars, Breweries & Distilleries":[
 ("Chattanooga Whiskey","Tim Piersant — co-founder/CEO (w/ Joe Ledbetter)","Downtown","Led the grassroots 'Vote Whiskey' campaign that overturned a century-old local distilling ban; built the first legal Chattanooga distillery in ~100 years. 'Whiskey to the People.'","https://chattanoogawhiskey.com/history/","High",True),
 ("Gate 11 Distillery","Bill & Wanda Lee — founders","Southside (Choo Choo)","A Sequatchie Valley native and former chemical engineer turned distiller; his retired-schoolteacher wife formulates their award-winning gin, inside historic Terminal Station.","https://www.gate11distillery.com/story","High",False),
 ("Hutton & Smith Brewing","Joel & Melanie Krautstrunk — owners","Southside (E. MLK)","Husband-and-wife who relocated from Las Vegas to open a brewery (2015); outdoor-adventure branding tied to local climbing/hiking.","https://huttonandsmithbrewing.com/about-us/our-team/","High",False),
 ("Naked River Brewing","Jake Raulston & Nathan Woods — co-founders","Southside","Hometown engineer-turned-brewer built the kid-/dog-friendly brewery-plus-BBQ he felt the city lacked, in an 1875 trolley-car factory.","https://noogatoday.6amcity.com/culture/the-history-of-naked-river-brewing-co","High",False),
 ("WanderLinger Brewing","Mike & Chris Dial — brothers","Downtown (Station St)","Built around a backpacking philosophy — 'wander some, linger some' — pairing beer with live music and local art.","https://www.wanderlinger.com/","High",False),
 ("OddStory Brewing","Jay Boyd — owner/head brewer","Highland Park / Central Ave","Alabama transplant who built a true family business (father and wife work alongside him); 'Drink Ours, Tell Yours.'","https://cityscopemag.com/city-scope/behind-the-brew/","High",False),
 ("Chattanooga Brewing Co.","Mark Marcum & Jonathan Clark — co-founders","Downtown (Chestnut St)","Two homebrewers revived the name of the original 1890 brewery killed by Prohibition — reclaiming the city's lost brewing heritage.","https://www.chattabrew.com/history","High",False),
 ("Five Wits Brewing","Nate Michaels, Elliot Kehoe & Bryan Harris","Southside (Long St)","A Chattanooga native lured two veteran Great Divide (Denver) brewers back home to start a local brewery.","https://fivewitsbrewing.com/","High",False),
 ("Matilda Midnight (The Dwell Hotel)","Seija Ojanpera — owner","Downtown (E. 10th St)","Bought a 110-year-old building in 2015 and reimagined it as a mid-century-modern boutique hotel with a dark, tiki-leaning cocktail lounge.","https://matildamidnight.com/","High",False),
],
"Coffee, Bakeries & Food Makers":[
 ("Mad Priest Coffee","Michael & Cherita Rice — co-founders","Downtown (McCallie)","Founded 2015 after meeting displaced refugees abroad; mission to 'champion the displaced' — a co-owner, Tarig Idris, was a Sudanese refugee and their first employee.","https://madpriestcoffee.com/pages/about-us","High",True),
 ("Velo Coffee Roasters","Andrew & Jessica Gage — owners","Downtown/Southside","Named for the French word for 'bike' — Andrew delivered fresh-roasted beans by bicycle in the early days (2010); college sweethearts.","https://dailycoffeenews.com/2014/09/08/inside-velo-coffee-roasters-growing-specialty-coffee-in-chattanooga/","High",False),
 ("Goodman Coffee Roasters","Ian Goodman — founder","Warehouse Row / St. Elmo","Opened his first shop at 22 in 1996, sold it, then returned to roasting in 2016 — a third-wave pioneer's comeback.","https://www.goodmancoffeeroasters.com/about","High",False),
 ("Niedlov's Bakery & Café","Erik & Lauren Zilen — owners","Southside (E. Main)","The city's cornerstone sourdough bakery; the Zilens hire people coming out of prison, compost ~1,300 lbs/month, and give away thousands of loaves.","https://niedlovs.com/about-us","High",True),
 ("Chattanooga Bakery (MoonPie)","Family-owned (heritage; no single owner)","Downtown","The 100+ year home of the MoonPie, invented 1917 when a miner asked for a snack 'as big as the moon.' Best as a heritage/centennial piece.","https://www.moonpie.com/about/index.html","High",False),
 ("Alchemy Spice Co.","Henry & John Oehmig — owner-brothers","Chattanooga","Longtime fans who bought the small-batch spice-blend brand and now run it together — 'two brothers, one passion.'","https://www.alchemyspicecompany.com/pages/about-us","High",False),
 ("The Hot Chocolatier","Wendy & Brandon Buckner — owners","Chattanooga","Husband-and-wife shop (since 2008); Wendy trained at Chicago's French Pastry School and makes handmade chocolates in small batches.","https://thehotchocolatier.com/pages/about-us","High",False),
 ("Cocoa Asante","Ella Livingston — founder","Chattanooga","Self-taught chocolatier and former math teacher; a viral 2018 TikTok launched her bonbons — now working toward sourcing cacao from her family's farm in Ghana.","https://www.cocoaasante.com/our-story-2","High",True),
 ("Noke's Granola","Luther Cutchins — owner","Chattanooga Market","Brothers started making granola for their catering business; after Ryan 'Noke' died, Luther kept it going in his memory. (Verify current ops.)","https://www.instagram.com/nokesgranola/","Med",False),
 ("Belle Chocolates","Brendan Patrick — owner","Northshore","Crafts chocolates from his Northshore kitchen — a solo 'candy man' kitchen-to-counter story. (Verify still operating.)","https://www.timesfreepress.com/news/2021/feb/06/candy-man-founder-belle-chocolates-puts-hhear/","Med",False),
],
"Retail, Boutiques & Creative Services":[
 ("Embellish Collection","Ann Trammell Newton — owner","Downtown (Warehouse Row)","A longtime customer who'd worked in NYC with von Furstenberg, de la Renta and Vince Camuto bought her favorite Chattanooga boutique in 2021 — a full-circle clothing story.","https://www.timesfreepress.com/news/2021/jun/01/embellish-boutique/","High",True),
 ("Anna Ball White","Trish Foy — owner","Hamilton Place area","A second-generation boutique founded in 1955; the founder's daughter now runs it, blending an interior-design eye into personal styling.","https://www.annaballwhite.com/about","High",False),
 ("Belle Rive Boutique","Silvinia Peralta-Ramos — owner","Northshore","Owner-curated womenswear known for warm, personal service. (Owner name from a listing site — verify.)","https://evendo.com/locations/tennessee/chattanooga/northshore/shop/belle-rive-boutique","Med",False),
 ("Locals Only Gifts & Goods","Danielle Landrum — founder (w/ Eric)","Northshore","Began assembling baskets of local products for her realtor husband's clients; now stocks 130+ TN makers and won a 2026 Shopify Milestone Award.","https://www.timesfreepress.com/news/community/story/2018/mar/07/locals-only/465199/","High",True),
 ("I Go Tokyo","Margaret Armour — owner (w/ Etsuko Lammon)","Northshore (Frazier Ave)","Born from a 2016 trip to Japan; the pair fly there 2–3x a year to hand-source linens, ceramics, and stationery unavailable in the U.S.","https://www.chattanoogan.com/2020/10/1/416078/Shopping-Trips-To-Japan-Provide-Unique.aspx","High",False),
 ("Sophie's Shoppe","Tamara Dillard & Veronica Deck — co-owners","Northshore & Downtown","Grew from pop-ups and flower arrangements into a 20-year institution championing local makers across home, art, and gifts.","https://sophiesshoppe.com/pages/our-shoppe","High",False),
 ("The Book & Cover","Emily Lilley, Blaes Green & Sarah Jackson","North Chattanooga","Three best friends and natives crowd-built their dream indie bookshop — a friendship-to-storefront origin.","https://www.timesfreepress.com/news/2023/feb/01/north-chattanooga-independent-book-shop-the-book/","High",False),
 ("Yellow Racket Records","Ben Vanderhart — founder","Chattanooga","Started as a record label in 2018 (named for his collie), opened the shop in 2020, now stocking 15,000+ records.","https://yellowracketrecords.com/","High",False),
 ("Dallos Vinyl Love","Marlo White — owner","Downtown (Market St)","A transplant who moved to Chattanooga in 2020; a curated soul/funk/jazz/hip-hop vinyl passion project, opened 2025.","https://www.timesfreepress.com/news/2025/jun/11/dallos-vinyl-love-brings-record-store-to-market/","High",False),
 ("Joli Jardin","Becca Coleman & Erin Leonard — owners","Signal Mtn (farm) + Downtown","Two women run a sustainable cut-flower farm supplying their own downtown shop — 'farm-to-vase,' scaled up from 100+ pop-ups.","https://jolijardin.co/","High",False),
 ("INK MXR Tattoo Studio","Josh Hansen — owner/artist","Chattanooga","Founded his custom-tattoo studio in 2022 and began mentoring the next generation (took on an apprentice in 2025).","https://inkmxr.com/","Med",False),
 ("All City Barber Co.","Kristin Gonzalez — owner","Southside","Moved from Sacramento at 19 and opened her own unisex barbershop by 23 — a young female owner in a male-dominated trade. (Story from 2016; verify.)","https://www.timesfreepress.com/news/edge/story/2016/sep/01/new-all-city-barber-co-southside-unisex-cater/383797/","Med",False),
],
"Home Builders, Makers & Craftspeople":[
 ("Collier Construction","Ethan Collier — founder/CEO","Chattanooga","Grew up here; a LEED-accredited builder who pioneered green building and urban infill before it was fashionable, with multiple LEED Platinum projects.","https://www.collierbuild.com/our-story","High",True),
 ("River Stone Construction","Dustin Wong — president/co-owner","Chattanooga","Named HBA of Greater Chattanooga 'Home Builder of the Year' (2019); partnership pairs a builder with a broker/investor. (Verify details.)","https://riverstone-homes.com/about/","Med",False),
 ("Tristar Custom Homes","Collette Branham — owner","Chattanooga","A second-generation builder who learned the trade from her father — a woman-owned custom builder, uncommon in the field. (Verify name.)","https://www.tristarcustomhomes.com/","Med",False),
 ("Boreal Woodworks","Nathan Kolb — founder (w/ Nate Penman)","North Chattanooga","Chattanooga Woodworking Academy grad who apprenticed under a master; builds heirloom furniture from Appalachian hardwoods (TFP profile, 2026).","https://www.borealwoodworks.com/about","High",True),
 ("Forman Pottery","John-Michael Forman — owner/potter","St. Elmo","Snuck into the college art studio to teach himself; grew a basement hobby into a full-time studio making thousands of functional pieces a year.","https://www.formanpottery.com/about/","High",False),
 ("Bean and Bailey Ceramics","Anderson Bailey & Jessie Bean Goodman","Highland Park","Met in the glass studio at the Appalachian Center for Craft and launched in 2013 on a CreateHere ArtsMove grant — a craft-school love-and-business story.","https://www.beanandbaileyceramics.com/the-makers","High",False),
 ("Nooga Leatherworks","Stan Daniel — owner","Chattanooga","Left an engineering career to handcraft one-of-a-kind leather goods — a classic maker career-pivot (NewsChannel 9 profile).","https://newschannel9.com/features/made-in-our-hometown/nooga-leatherworks-owner-makes-one-of-a-kind-items-from-leather","Med",False),
],
"Farms & Farmers":[
 ("Sequatchie Cove Farm","The Keener family (Bill & Miriam; Kelsey & Ashley)","Sequatchie Cove (Marion Co.)","Four generations on 300 regenerative acres; the creamery's raw-milk 'Dancing Fern' is a nationally recognized Tennessee farmstead cheese.","https://www.sequatchiecovefarm.com/philosophy","High",True),
 ("Riverview Farms","The Swancy family (Wes, Brad & Drew)","NW Georgia (Main St Mkt)","Sons converted the family hog-and-corn farm to certified-organic in 2000 to save it; Brad stone-mills heirloom grits from farm-grown corn.","https://grassfedcow.com/our-story","High",False),
 ("Southerly Flower Farm","Matthew & Sarah Ervin — owners","Bledsoe County","A husband-and-wife team who built a sustainable cut-flower farm from scratch 'with no experience,' now supplying markets, CSAs, and weddings. (Verify names.)","https://southerlyflowerfarm.com/about/","Med",False),
],
"Outdoor, Adventure & Founders":[
 ("Synergy Climbing and Ninja","Lisa Rands & Isaac Caldiero — co-founders","Southside (E. Main)","A world-champion boulderer and the first million-dollar 'American Ninja Warrior' winner built a community gym in their adopted hometown.","https://synergyclimbingandninja.com/our-story/","High",True),
 ("High Point Climbing & Fitness","Johnny O'Brien & John Wiygul — founders","Downtown","Friends who met after a 2005 triathlon built the gym Climbing Magazine called the 'coolest in the country' — glass walls over Broad St — then a regional brand.","https://climbingbusinessjournal.com/southern-climbing-reaches-its-high-point/","High",False),
 ("Suck Creek Cycle","Mike Skiles — owner","Northshore","A '70s BMX racer who 'swept floors' into the trade and has run the city's premier mountain-bike shop for nearly 30 years.","https://www.wdef.com/driving-our-economy-forward-suck-creek-cycle/","High",False),
 ("Two Bikes Chattanooga","Mitchell Connell — founder/ED","Chattanooga","A nonprofit bike shop that donates a bike for every one sold and runs paid internships for at-risk youth — retail meets mission.","https://www.dayfirepodcast.com/ep-246-mitchell-connell-/-two-bikes-chattanooga/","High",True),
 ("Brev.dev (now NVIDIA Brev)","Nader Khalil — co-founder/CEO","Brickyard","First Brickyard exit: came because 'Brickyard was the only investor that wanted to meet in person'; built a GPU dev-tools startup NVIDIA acquired in 2024.","https://chattanoogachamber.com/from-brickyard-to-the-bay-brev-dev-acquired-by-nvidia/","High",False),
 ("Brickyard","Cameron Doody & Matt Patterson — co-founders","Southside (rug warehouse)","Bellhop co-founders who require startups to relocate to Chattanooga and work in person until $1M revenue — the people building the scene itself.","https://hypepotamus.com/startup-news/chattanooga-brickyard-international-founders/","High",False),
 ("FreightWaves","Craig Fuller — founder/CEO","Chattanooga","Third-generation trucking family member who built the 'Bloomberg of freight,' cementing Chattanooga as the 'Silicon Valley of trucking.'","https://www.thelogisticsoflogistics.com/repost-why-chattanooga-is-the-silicon-valley-of-trucking-with-craig-fuller/","High",False),
],
}

RESERVE = [
 "Easy Bistro & Main Street Meats (Erik & Amanda Niel) — same owners as Little Coyote; alt concepts.",
 "SideTrack / Sugar's Ribs (Lawton & Karen Haygood) — same owners as Boathouse.",
 "Whiskey Thief, atop The Edwin (bartending team — name unverified); STIR (beverage lead unverified).",
 "Jason Bowers — Clever Ale House / Civil Provisions / The Bitter Alibi (Red Bank multi-concept).",
 "Hemline; Monkee's of Chattanooga; Electric Blue Collection (Anita Oaks Headrick) — more apparel.",
 "GreenTech Homes (Paul Terura & Jim Storey); Cuthbertson Homes; Pfeffer Torode Architecture.",
 "Chattanooga Leather Works (Ben Beasley); Fox & Forest Handcrafted Leather.",
 "Owen Cyclery (Erik Hunt & Brian Lowery); L2 Outside (Deb & Steve Clanin) — verify owners.",
 "Ada's Heritage Farm; Chattanooga Olive Oil Co. — owner names unverified.",
 "Legacy/exit founder profiles: Dawson Wheeler (Rock/Creek), Ken McElrath (Skuid), Cam Doody (Bellhop).",
]

FLAGS = [
 "1885 Grill — St. Elmo location reported closing in 2026; brand/owner intact. Confirm current address.",
 "Noke's Granola, Belle Chocolates, All City Barber, Señor Shan's — confirm 2026 operating status before pitching.",
 "Belle Rive, Southerly Flower Farm, River Stone, Tristar — owner names from secondary/aggregator sources; verify spelling.",
 "DO NOT feature as active: Hoff & Pepper (hot sauce, closed 2025); Brash Coffee (moved to Atlanta).",
 "MoonPie/Chattanooga Bakery — family-owned, no single public 'maker'; treat as a heritage piece.",
 "Confirm any single-sourced name in the interview itself — no names here were invented; unverified ones are labeled.",
]

GAPS = [
 "No verified Chattanooga-metro cidery surfaced — a category gap to fill.",
 "Ooltewah, Signal Mountain, and East Ridge are thin on chef-owner/indie coverage — source locally for geographic spread.",
 "Photographers: strong verified owner stories were scarce — ask readers for nominations.",
 "Ongoing pipelines: Main Street Farmers Market (Wed) & Chattanooga Market (Sun) for farmers/makers; CO.LAB & the Chamber for founders.",
]

# ---------------- CSV ----------------
with open("/home/user/prscott/meet_a_maker_source_list.csv","w",newline="") as f:
    w=csv.writer(f)
    w.writerow(["Category","Business","Person","Area","Story angle","Source","Confidence","Flagship"])
    for cat,rows in DATA.items():
        for (biz,person,area,angle,src,conf,flag) in rows:
            w.writerow([cat,biz,person,area,angle,src,conf,"Yes" if flag else ""])
total=sum(len(r) for r in DATA.values())
print("CSV rows:",total)

# ---------------- DOCX ----------------
doc=Document()
sec=doc.sections[0]
sec.orientation=WD_ORIENT.LANDSCAPE
sec.page_width,sec.page_height=Inches(11),Inches(8.5)
sec.top_margin=Inches(0.6); sec.bottom_margin=Inches(0.6); sec.left_margin=Inches(0.6); sec.right_margin=Inches(0.6)
normal=doc.styles["Normal"]; normal.font.name="Calibri"; normal.font.size=Pt(10)
normal.paragraph_format.space_after=Pt(4)

def sc(r,h): r.font.color.rgb=RGBColor.from_string(h)
def shade(cell,h):
    p=cell._tc.get_or_add_tcPr(); s=OxmlElement("w:shd")
    s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); p.append(s)
def ctext(cell,t,bold=False,color=None,size=8.5,white=False,align="left"):
    cell.text=""; p=cell.paragraphs[0]
    p.alignment={"left":WD_ALIGN_PARAGRAPH.LEFT,"center":WD_ALIGN_PARAGRAPH.CENTER}[align]
    r=p.add_run(t); r.font.size=Pt(size); r.font.bold=bold
    if white: r.font.color.rgb=RGBColor.from_string("FFFFFF")
    elif color: r.font.color.rgb=RGBColor.from_string(color)
def H(t,level=1,color=None):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(12 if level==1 else 8); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(t); r.font.bold=True
    if level==1:
        r.font.size=Pt(16); sc(r,color or NAVY)
        pPr=p._p.get_or_add_pPr(); b=OxmlElement("w:pBdr"); bt=OxmlElement("w:bottom")
        bt.set(qn("w:val"),"single"); bt.set(qn("w:sz"),"6"); bt.set(qn("w:space"),"2"); bt.set(qn("w:color"),TEAL)
        b.append(bt); pPr.append(b)
    else: r.font.size=Pt(12.5); sc(r,color or TEAL)
def P(t,italic=False,color=None,size=10):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(5)
    for i,ch in enumerate(t.split("**")):
        if not ch: continue
        r=p.add_run(ch); r.font.size=Pt(size); r.font.italic=italic
        if i%2==1: r.font.bold=True
        if color: sc(r,color)
def BUL(items,size=9.5):
    for it in items:
        p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(2)
        r=p.add_run(it); r.font.size=Pt(size)

def maker_table(rows, with_conf=True):
    cols=["Business","Person & role","Area","Story angle (the human hook)","Conf."]
    widths=[1.5,2.0,1.15,5.0,0.55]
    t=doc.add_table(rows=1,cols=len(cols)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style="Table Grid"
    for i,h in enumerate(cols):
        shade(t.rows[0].cells[i],NAVY); ctext(t.rows[0].cells[i],h,bold=True,white=True,size=9,align="left" if i!=4 else "center")
    for r,(biz,person,area,angle,src,conf,flag) in enumerate(rows):
        cells=t.add_row().cells; fill=SAND if flag else (LIGHT if r%2==0 else "FFFFFF")
        ctext(cells[0],biz+("  ★" if flag else ""),bold=True,color=NAVY,size=8.5)
        ctext(cells[1],person,size=8.5)
        ctext(cells[2],area,size=8.5)
        ctext(cells[3],angle,size=8.5)
        ctext(cells[4],conf,size=8.5,align="center")
        for c in cells:
            shade(c,fill)
    for i,wd in enumerate(widths):
        for row in t.rows: row.cells[i].width=Inches(wd)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)

# Title
for _ in range(2): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("THE SCENIC CITY SCOUT"); r.font.size=Pt(24); r.font.bold=True; sc(r,NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("“Meet a Maker” Source Book"); r.font.size=Pt(18); sc(r,TEAL)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run(f"{total} real Chattanooga small businesses and the people behind them — an editorial pipeline"); r.font.size=Pt(12); r.font.italic=True; sc(r,SLATE)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Researched June 2026  ·  ★ = flagship-ready profile  ·  Full sources & confidence in the accompanying CSV"); r.font.size=Pt(9.5); sc(r,GREY)
doc.add_page_break()

# How to use
H("How to use this source book",1)
P("This is a working pipeline of **verified, real** Chattanooga-area businesses and the named people behind them — built for the weekly “Meet a Maker” feature and the monthly long-form flagship profile. Every entry has a person, a neighborhood, a one-line human hook, a source, and a confidence rating (in the CSV). **No names were invented**; anything unverified is labeled and listed under “Verify before print.”")
P("**Editorial reminder:** these are *editorial* picks, never paid. A business appearing here must be chosen on merit — that independence is the whole moat. (See the bright-line rule in Addendum 2.)", color=SLATE)
BUL([
 "★ marks the strongest flagship-ready stories (clean origin + high confidence) — start here.",
 "Confidence: High = multi-source / primary; Med = single or secondary source; verify before print.",
 "Use the CSV to drop the whole list into a spreadsheet or Airtable and track outreach, status, and run dates.",
])

# Flagship section
flagship=[(b,p2,a,an,s,c,f) for cat,rows in DATA.items() for (b,p2,a,an,s,c,f) in rows if f]
H("★ Flagship-ready: start with these",1,color=GOLD)
P("The cleanest origin stories with the best human hooks and high confidence — ideal for the first run of features and the monthly long-form profile.")
maker_table(flagship)

# Category sections
for cat,rows in DATA.items():
    H(cat+f"  ({len(rows)})",1)
    maker_table(rows)

# Reserve + flags + gaps
H("Reserve bench & additional leads",1)
P("Strong leads held back to avoid double-profiling the same owners, or pending verification:")
BUL(RESERVE)
H("Verify before print",1,color=ORANGE)
BUL(FLAGS)
H("Category gaps & ongoing sourcing pipelines",1)
BUL(GAPS)

P("Sources for every entry are in meet_a_maker_source_list.csv. Confirm any single-sourced name during the interview itself.",
  italic=True,color=GREY,size=9)

out="/home/user/prscott/Meet_a_Maker_Source_Book.docx"
doc.save(out); print("saved",out)
print("flagship count:",len(flagship))
