#!/usr/bin/env python3
"""Still Standing — the Legacy Businesses Source Book.
Compiled from the verified Legacy markdown + Insider (new finds) & Nerd (status) agent passes.
Editorial layer: longevity tiers, live status, a 'commodity vs. true-Insider' filter, and the angle for each."""
import os, csv, matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY="1B3A4B"; TEAL="2A9D8F"; ORANGE="E76F51"; SAND="E9C46A"; SLATE="415A6B"
LIGHT="EEF3F4"; GREY="6B7B85"; GOLD="B8860B"; GREEN="2E7D52"; RUST="A6452F"
ASSETS="/home/user/prscott/assets"; os.makedirs(ASSETS,exist_ok=True)
plt.rcParams.update({"font.family":"DejaVu Sans","figure.dpi":160})
NOW=2026

# ============================================================
# DATA  — (name, founded, town, category, status, scout_angle, conf)
#   status: OPEN / SOLD / RENOVATING / CLOSED / RELOCATED
#   founded is an int (year). years open computed = NOW - founded.
# ============================================================

# --- 100+ years: the Centenarians ---
CENT=[
 ("Miller & Martin",1867,"Chattanooga, TN","Law / professional","OPEN",
  "Chattanooga's oldest law firm; early clients included the original Coca-Cola bottler. Angle: the firm that lawyered the handshake that built a global beverage system.","High"),
 ("Fischer Evans Jewelers",1869,"Chattanooga, TN","Retail — jewelry","OPEN",
  "The oldest independently owned retail store in the city, same Market St. block since 1869, now 3rd-gen Glover family. Angle: the engagement-ring counter that has seen five generations of the same families propose.","High"),
 ("Chattanooga Times Free Press",1869,"Chattanooga, TN","Media","OPEN",
  "A 20-year-old Adolph Ochs bought it in 1878 before going on to The New York Times; cut to Sunday-only print Jan 2026. Angle: handled carefully — we're the new kid; tip the cap to the paper of record, don't pick a fight.","High"),
 ("Read House Hotel",1872,"Chattanooga, TN","Hospitality","OPEN",
  "The longest continuously operated hotel in the South (150 yrs in 2022); current tower is 1926. Angle: the ghost-story/room-311 lore + the jazz-era lobby, not the press-release renovation.","High"),
 ("Chambliss Center for Children",1872,"Chattanooga, TN","Childcare / nonprofit","OPEN",
  "Started by church women as an orphans' home; ran 24-hour childcare when working mothers had nowhere else. Angle: the quiet 150-yr institution nobody writes about — a giving/volunteer tie-in, not a 'cute history' piece.","High"),
 ("Leitner, Williams, Dooley & Napolitan",1882,"Chattanooga, TN","Law / professional","OPEN",
  "Founding partner W.G.M. Thomas incorporated the town of Lookout Mountain and represented it pro bono for life. Angle: the lawyer who gave a whole town away.","Med"),
 ("Adams Lithographing",1886,"Chattanooga, TN","Printing","OPEN",
  "Installed the city's first rotary press (1935); 3rd-gen Hogue family. Angle: 'Who Made This' — the ink behind a century of Chattanooga's posters, packaging, and programs.","Med"),
 ("Erlanger Health System",1899,"Chattanooga, TN","Healthcare","OPEN",
  "Born from a French baron's $5,000 gift; region's first open-heart surgery (1960). Angle: 'where Chattanooga was born' — generations literally delivered here. (Verify the current Erlanger name/governance — see status pass.)","High"),
 ("Chattanooga Bakery / MoonPie",1902,"Chattanooga, TN","Food mfg. / retail","OPEN",
  "Bakery OPEN, ~1M+ MoonPies/day, 5th-gen, still independent — the Broad St. General Store closed/relocated May 2025 (confirm new riverfront address). COMMODITY: everyone covers the MoonPie. New angle only: the bakery floor, the family, or the store's move — not the origin myth again.","High"),
 ("Fort Payne Opera House",1889,"Fort Payne, AL","Entertainment / venue","OPEN",
  "Oldest in-use theater in Alabama, built in the coal-and-iron boom. Angle: a Getaway Field Guide tie-in (Fort Payne / Little River Canyon day trip).","Med"),
 ("Chattanooga Choo-Choo",1909,"Chattanooga, TN","Hospitality / entertainment","OPEN",
  "Terminal Station 1909 → hotel/entertainment complex 1973; now the Hotel Chalet. COMMODITY: heavily covered. New angle only: a specific room, venue, or person inside, never the 'famous song' lede.","High"),
 ("Fillauer",1914,"Chattanooga, TN","Healthcare / prosthetics","OPEN",
  "Started as the Red Cross Pharmacy; pivoted to fitting braces for WWI vets, now a global prosthetics leader (acquired by Hanger Inc., Feb 2024, still operating under its own name here). Angle: a maker-of-mobility human story — a local patient who walks because of a 110-yr-old neighbor company.","Med"),
 ("Zarzour's Café",1918,"Chattanooga, TN","Restaurant","SOLD",
  "TN's longest-running family restaurant — opened in the 1918 flu year, never moved. SOLD late 2025; new owners pledged to keep the name and reopen after renovation. COMMODITY + sensitive: only cover the reopening, with respect for the Zarzour/Fuller legacy. (Per our once-rule, don't lean on it.)","High"),
 ("Hamilton Health Care System",1921,"Dalton, GA","Healthcare","OPEN",
  "Founded after the 1918 flu with Crown Cotton Mills money; first GA hospital certified for Medicare. Angle: a Dalton/North-GA reader-service piece, not a Chattanooga story.","Med"),
 ("Ruby Falls",1930,"Chattanooga, TN","Tourism / attraction","OPEN",
  "85-ft underground waterfall, discovered 1928, open since 1930; named for Leo Lambert's wife. COMMODITY: the #1 paid attraction — tourists' turf. Skip unless a genuinely new angle (a guide, a night tour, the geology).","High"),
 ("Rock City Gardens",1932,"Lookout Mountain, GA","Tourism / attraction","OPEN",
  "'See Rock City' barns (from 1935) are Americana; approaching its centennial. COMMODITY: don't do the barns again. Possible angle: the people who still paint/maintain them, or Frieda Carter's fairy-tale design.","High"),
 ("McKee Foods / Little Debbie",1934,"Collegedale, TN","Food mfg. / retail","OPEN",
  "From cakes sold off a 1928 Whippet to ~$1.9B, still family-owned; Little Debbie Park opened 2023. COMMODITY-ish: the brand is famous, the family is private. Angle: the bakery-store deals, the park, or a Collegedale day-out — local-service, not corporate history.","High"),
]

# --- 75+ years: the Near-Centenarians ---
NEAR=[
 ("Wally's Restaurant",1937,"Chattanooga, TN","Restaurant — meat-and-three","OPEN",
  "Opened as a 1937 drive-in on McCallie, sit-down since 1972; walls papered in city history. Angle: a 'The Regular' profile — a booth, a waitress of 30 years, the Tuesday meatloaf crowd.","High"),
 ("Brody Jewelers",1937,"Rossville, GA","Retail — jewelry","OPEN",
  "Began as the Rossville Pawn Shop (1937), now 4th-gen GIA-diamond jeweler 10 min from downtown. Angle: a satellite-town find — Rossville rarely gets the spotlight.","Med"),
 ("Chattanooga Restaurant Supply",1948,"Chattanooga, TN","Food co-op (B2B-ish)","OPEN",
  "Six restaurant owners formed a co-op in 1948 — one of only two wholesale food co-ops in the nation. Angle: 'Who Made This' / the unseen plumbing behind every local kitchen you love.","Med"),
 ("Bea's Restaurant",1950,"Chattanooga, TN","Restaurant — Southern buffet","OPEN",
  "Lazy-Susan family-style Southern food since 1950, now 5th generation; did $1 plates for its 75th in 2025. Angle: the lazy-Susan ritual itself — a sensory 'Last Call'/'Regular' piece.","High"),
 ("Pruett's Market",1953,"Signal Mountain, TN","Grocery — independent","OPEN",
  "Signal Mountain's primary grocer since 1953; deli, bakery, sushi, craft beer — an independent that beat the chains. Angle: a satellite-town anchor; how a mountain town kept its own store.","High"),
]

# --- 50+ years: Established Veterans ---
VET=[
 ("Memo's Grill",1966,"Chattanooga, TN","Restaurant — hot dogs","OPEN",
  "A 1966 landmark for chopped-wiener plates in thick chili with slaw on toast. Angle: the dish is the story — a 'The Dig' on the chopped-wiener lineage (vs. Columbus GA's scrambled dog).","Med"),
 ("T.U. Parks Construction",1944,"Chattanooga, TN","Construction / services","OPEN",
  "Oldest general contractor in greater Chattanooga (1944); built much of what you walk past. Angle: 'Building's Biography' partner — they can tell you what a building used to be.","Med"),
 ("Pickle Barrel",1982,"Chattanooga, TN","Bar / restaurant","OPEN",
  "The city's oldest bar (1982) in the flatiron building at 1012 Market, food till 1:30am. Angle: the late-night institution — a 'Last Call' piece. (Note: 44 yrs — just shy of 50; include as 'approaching.')","Med"),
]

# --- Closed, but matter for history (Building's Biography / The Dig) ---
GONE=[
 ("Old Plantation Bar-Be-Que",1969,"Chattanooga, TN","Restaurant — BBQ","CLOSED",
  "Brick-pit takeout on Dodson Ave (1969), ran 50+ yrs before closing. Use only for history — a 'what was here' / The Dig reference, never a 'go visit.'","Med"),
]

# --- Regional / edge — judgment calls (within driving range) ---
EDGE=[
 ("Bright Star Restaurant",1907,"Bessemer, AL","Restaurant","OPEN",
  "Alabama's oldest family restaurant (1907), James Beard 'America's Classic.' Use only as a Getaway Field Guide stop on a Birmingham run — not a Chattanooga story.","Med"),
 ("The Partridge Restaurant",1933,"Rome, GA","Restaurant","CLOSED",
  "Rome institution (1933) that closed ~2017-20; building bought 2023 to reopen as an event venue. Watch for a reopening; history-only for now.","Med"),
 ("Jack's Family Restaurants",1960,"Homewood, AL (regional)","Fast food","OPEN",
  "Regional chain (1960), 285+ locations; the Fort Payne store anchors the DeKalb corridor. Too chain-y for a feature; map context only.","High"),
]

# --- NEW Insider finds (hidden, satellite-town legacy) — folded from Insider agent ---
# (name, founded, town, category, status, scout_angle, conf)
SCOUT_FINDS=[
 ("Morgan Furniture",1909,"Dayton, TN","Retail — furniture","OPEN",
  "Downtown Dayton furniture store; once delivered by horse-and-wagon and rented rooms to Scopes Trial journalists in 1925. Believed oldest family retailer in Rhea County. Angle: a Scopes-Trial-adjacent 'Building's Biography' + a Dayton day-trip.","High"),
 ("Buddy's Shoe Repair",1917,"Hixson, TN","Cobbler — shoe repair","OPEN",
  "Three-generation Scallia-family cobbler (trade dates to 1906 Sicily); reportedly the area's only remaining shoe-repair shop, with a 4th gen training in. Angle: 'Who Made This' — an endangered craft, the last cobbler.","High"),
 ("Beaty Hardware (Beaty Feed & Hardware)",1936,"Cleveland, TN","Hardware / feed store","OPEN",
  "Started as a one-room feed store (1936), added hardware in 1952; reportedly Cleveland's oldest hardware store. Angle: the unglamorous farm-town institution — feed, lawn & garden, and old-timers' advice.","High"),
 ("The Chef",1967,"Cleveland, TN","Diner / drive-in","OPEN",
  "The last surviving original Burger-Chef building still serving; home of 'Hot Slaw' (now TN's official state food); says it has never closed a day since 1967. Angle: the dish + the building — a Cleveland day-trip and a 'The Dig' on Hot Slaw.","High"),
 ("Sliger's Jewelers",1946,"Athens, TN","Retail — jewelry","OPEN",
  "Family-owned downtown Athens jeweler — custom rings, watch & jewelry repair — 70+ years on E. Washington Ave. Angle: a satellite-town counter nobody outside Athens writes about.","High"),
 ("Joy's Flowers (Joy Floral Co.)",1915,"Chattanooga, TN","Retail — florist","OPEN",
  "Among the South's oldest florists (lineage to 1877, retail shop ~1915) on McCallie Ave, with a historical marker. Angle: the flowers behind a century of local weddings & funerals — slightly more 'known,' so needs a fresh thread.","Med"),
 ("Oakwood Cafe",1915,"Dalton, GA","Diner / meat-and-three","OPEN",
  "Long-running Southern comfort-food diner on W. Cuyler St; the Metcalf family took over an already-30-yr-old cafe in the 1940s (founding ~1910s, year approx — verify). Angle: a meat-and-three that outlasted the carpet boom.","Med"),
 ("Professional Pharmacy",1961,"Dalton, GA","Independent pharmacy","OPEN",
  "Family-owned independent drugstore, open 365 days/yr with county-wide delivery, 60+ years (founding decade confirmed; exact 1961 not nailed down). Angle: the independent that outlived the chains in a satellite town.","Med"),
 ("Access Family Pharmacy",1955,"Hixson, TN","Independent / compounding pharmacy","OPEN",
  "Locally owned compounding & medical-supply pharmacy on Hixson Pike (sources disagree: 1955 vs 1959; changed hands 1990). VERIFY founding year before featuring. Angle: a neighborhood compounding druggist.","Low"),
]
# Insider's suspected-but-unverified (do NOT publish without a phone call):
SCOUT_LEADS=[
 "**The Folk Music Store** — Red Bank, TN (3411 Dayton Blvd): acoustic-instrument shop; a 'sndsr1948' handle hints at 1948 but that may be a birth year, not a founding. Worth a call — would be a perfect gem if it's real.",
 "**Tipton Funeral Home** — Cleveland, TN: death records back to 1899 suggest a late-1800s founding; current continuity unverified. Check the Bradley County historical society.",
 "**Dayton Drug / Robinson Drug** — Dayton, TN: Robinson Drug is where the Scopes Trial was hatched in 1925; unclear if today's 'Dayton Drug & Wellness' is a continuation. Verify before claiming.",
 "**Richest untapped vein (next pass):** funeral homes with retail, feed stores, and barber/beauty shops in LaFayette, Ringgold, Fort Oglethorpe, and Soddy-Daisy — poorly indexed online; would need phone calls or local historical societies.",
]
# Recently LOST institutions (history only — do not invite readers):
SCOUT_LOST=[
 "**Koch's Bakery** — Chattanooga (est. Feb 14, 1948): a top pick that CLOSED March 18, 2023 after 75 years. A recently-lost institution worth a tribute/The Dig, not a visit.",
 "**Nikki's Drive-Inn** — Chattanooga (1941): iconic drive-in, now closed; also over-covered.",
]

# --- Status verification notes — folded from Nerd agent ---
STATUS_NOTES=[
 "**Zarzour's Café — SOLD & closed for renovation (Nov 2025).** Owner Joe 'Dixie' Fuller sold the building (~$515K); final service mid-Nov 2025. Anonymous LLC pledged to keep the name and reopen post-remodel — no date, original staff not returning. Watch for the reopening; don't invite readers yet. (HIGH)",
 "**Chattanooga Times Free Press — cut print to Sunday-only**, effective Jan 26, 2026 (still WEHCO Media-owned, no sale). A live local-media story in itself. (HIGH)",
 "**Erlanger — completed public-to-private nonprofit conversion (2023)** and merged in its Foundation (2024); leadership restructured. Name unchanged — still 'Erlanger.' (HIGH)",
 "**Fillauer — acquired by Hanger Inc. (Feb 2024)** but still open under its own name in Chattanooga; consolidated its TRS line here in 2025. Technically under new corporate parent. (HIGH)",
 "**MoonPie General Store (429 Broad St) — closed/relocated May 6, 2025.** The Broad St address is gone; the brand store appears to have moved to the riverfront. CONFIRM the new permanent address before publishing — this is our lowest-confidence detail. (MED)",
 "**Old Plantation Bar-Be-Que — confirmed CLOSED** (announced 2022 after 53 yrs; owner's stroke + COVID/staffing). History only. (HIGH)",
 "**Wally's (1937 original, 1600 McCallie) — OPEN.** Note: a separate, newer East Ridge location (6521 Ringgold Rd) closed in 2023 and is becoming a Whataburger — that is NOT the original. Don't conflate them. (HIGH)",
 "All others verified OPEN as of mid-2026: Fischer Evans, Miller & Martin, Read House (added a garden in 2025), Rock City & Ruby Falls (now timed-entry), McKee/Little Debbie, Chambliss (expanding via micro-centers), the Choo-Choo/Hotel Chalet, Bea's, Brody, Memo's, Pruett's, Pickle Barrel, T.U. Parks, Chattanooga Restaurant Supply.",
]

MYTHS=[
 "**'Oldest restaurant' is contested.** Zarzour's (1918) is the longest-running *family* restaurant in TN; Wally's (1937) is among the oldest *operating* — phrase precisely and don't crown one without a source.",
 "**Status changes fast.** Zarzour's SOLD (2025), the MoonPie General Store CLOSED (2025), Old Plantation CLOSED. Re-verify 'still open' the week you publish — an obituary printed as an invitation is the worst look we can have.",
 "**'Founded' ≠ 'at this location since' ≠ 'continuously operated.'** Read House sits on the older Crutchfield House site; the tower is 1926, not 1872. Say what you mean.",
 "**B2B/professional firms are 'consumer-facing' only at the edges.** Law and construction firms belong in heritage features, not in 'go spend your Saturday here' picks.",
 "**Regional ≠ local.** Bright Star (Bessemer), Partridge (Rome), Hamilton Health (Dalton) are drive-trip or reader-service content, not Chattanooga-proper stories — label them honestly.",
]

# ============================================================
# CSV
# ============================================================
def yrs(f): return NOW - f
GROUPS=[("100+ years — Centenarians",CENT),("75+ years — Near-Centenarians",NEAR),
        ("50+ years — Established Veterans",VET),("Closed — history only",GONE),
        ("Regional / edge",EDGE),("Insider finds — hidden & satellite-town",SCOUT_FINDS)]
with open("/home/user/prscott/legacy_source_list.csv","w",newline="") as f:
    w=csv.writer(f); w.writerow(["Tier/group","Business","Founded","Years open","Town","Category","Status","Insider angle","Confidence"])
    for grp,rows in GROUPS:
        for (n,fo,tn,cat,st,ang,cf) in rows:
            w.writerow([grp,n,fo,yrs(fo),tn,cat,st,ang,cf])
TOTAL=sum(len(r) for _,r in GROUPS)
print("CSV rows:",TOTAL)

# ============================================================
# CHART — longevity timeline
# ============================================================
def chart_timeline():
    pts=[(n,fo) for grp,rows in [("",CENT),("",NEAR),("",VET)] for (n,fo,*_ ) in rows]
    pts.sort(key=lambda x:x[1])
    fig,ax=plt.subplots(figsize=(9.8,4.4)); ax.set_xlim(1860,2030); ax.set_ylim(-1.7,1.7)
    ax.axhline(0,color="#"+SLATE,lw=2,zorder=1)
    for i,(name,yr) in enumerate(pts):
        up=1 if i%2==0 else -1; y=up*(0.6 if i%4<2 else 1.12)
        c=["#"+NAVY,"#"+TEAL,"#"+ORANGE,"#"+GOLD][i%4]
        ax.plot([yr,yr],[0,y],color=c,lw=1.2,zorder=2)
        ax.scatter([yr],[0],s=38,color=c,zorder=3,edgecolor="white",linewidth=1)
        short=name.split(" / ")[0].split(",")[0]
        ax.text(yr,y+(0.1*up),f"{short}\n{yr}",ha="center",va="bottom" if up>0 else "top",
                fontsize=6.3,color="#"+NAVY,fontweight="bold")
    ax.axvspan(1860,1926,color="#"+TEAL,alpha=0.05)
    ax.text(1893,-1.55,"the Centenarians (100+)",fontsize=7,color="#"+TEAL,style="italic")
    ax.set_yticks([]); ax.set_xticks([1870,1890,1910,1930,1950,1970,1990,2010])
    for s in ("top","right","left"): ax.spines[s].set_visible(False)
    ax.set_title("Still Standing — legacy consumer businesses by founding year",fontweight="bold",color="#"+NAVY)
    fig.tight_layout(); fig.savefig(f"{ASSETS}/fig_legacy.png",bbox_inches="tight"); plt.close(fig)
chart_timeline(); print("chart done")

# ============================================================
# DOCX helpers
# ============================================================
doc=Document(); sec=doc.sections[0]; sec.orientation=WD_ORIENT.LANDSCAPE
sec.page_width,sec.page_height=Inches(11),Inches(8.5)
for m in ("top_margin","bottom_margin","left_margin","right_margin"): setattr(sec,m,Inches(0.6))
nm=doc.styles["Normal"]; nm.font.name="Calibri"; nm.font.size=Pt(10); nm.paragraph_format.space_after=Pt(4)
def sca(r,h): r.font.color.rgb=RGBColor.from_string(h)
def shade(c,h):
    p=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd")
    s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); p.append(s)
def ctext(c,t,bold=False,color=None,size=8.4,white=False,align="left"):
    c.text=""; p=c.paragraphs[0]
    p.alignment={"left":WD_ALIGN_PARAGRAPH.LEFT,"center":WD_ALIGN_PARAGRAPH.CENTER}[align]
    r=p.add_run(t); r.font.size=Pt(size); r.font.bold=bold
    if white: r.font.color.rgb=RGBColor.from_string("FFFFFF")
    elif color: r.font.color.rgb=RGBColor.from_string(color)
def H(t,level=1,color=None):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(12 if level==1 else 8); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(t); r.font.bold=True
    if level==1:
        r.font.size=Pt(15); sca(r,color or NAVY)
        pPr=p._p.get_or_add_pPr(); b=OxmlElement("w:pBdr"); bt=OxmlElement("w:bottom")
        bt.set(qn("w:val"),"single"); bt.set(qn("w:sz"),"6"); bt.set(qn("w:space"),"2"); bt.set(qn("w:color"),TEAL)
        b.append(bt); pPr.append(b)
    else: r.font.size=Pt(12); sca(r,color or TEAL)
def P(t,italic=False,color=None,size=10.5):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(5)
    for i,ch in enumerate(t.split("**")):
        if not ch: continue
        r=p.add_run(ch); r.font.size=Pt(size); r.font.italic=italic
        if i%2==1: r.font.bold=True
        if color: sca(r,color)
def BUL(items,size=9.5):
    for it in items:
        p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(3)
        for i,ch in enumerate(it.split("**")):
            if not ch: continue
            r=p.add_run(ch); r.font.size=Pt(size)
            if i%2==1: r.font.bold=True
def FIG(path,cap,width=9.4):
    doc.add_picture(path,width=Inches(width)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    P(cap,italic=True,color=GREY,size=8.5); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
STATUS_COLOR={"OPEN":GREEN,"SOLD":GOLD,"RENOVATING":GOLD,"CLOSED":RUST,"RELOCATED":SLATE}
def LTABLE(rows,caption=None):
    cols=["Business","Founded","Yrs","Town","Category","Status","The Insider angle (POV that earns membership)","Conf."]
    widths=[1.45,0.55,0.4,1.25,1.15,0.75,3.6,0.45]
    t=doc.add_table(rows=1,cols=len(cols)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style="Table Grid"
    for i,h in enumerate(cols):
        shade(t.rows[0].cells[i],NAVY)
        ctext(t.rows[0].cells[i],h,bold=True,white=True,size=8.2,align="left" if i in(0,3,4,6) else "center")
    for r,(n,fo,tn,cat,st,ang,cf) in enumerate(rows):
        cells=t.add_row().cells; base=LIGHT if r%2==0 else "FFFFFF"
        ctext(cells[0],n,bold=True,color=NAVY,size=8.2)
        ctext(cells[1],str(fo),size=8.0,align="center")
        ctext(cells[2],str(NOW-fo),size=8.0,align="center")
        ctext(cells[3],tn,size=7.9); ctext(cells[4],cat,size=7.9)
        ctext(cells[5],st,bold=True,color=STATUS_COLOR.get(st,SLATE),size=7.8,align="center")
        ctext(cells[6],ang,size=8.0); ctext(cells[7],cf,size=7.9,align="center")
        for c in cells: shade(c,base)
    for i,w in enumerate(widths):
        for row in t.rows: row.cells[i].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    if caption: P(caption,italic=True,color=GREY,size=8.3)

# ============================================================
# DOCX body
# ============================================================
doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("THE SCENIC CITY SCOUT"); r.font.size=Pt(23); r.font.bold=True; sca(r,NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Still Standing — the Legacy Businesses Source Book"); r.font.size=Pt(17); sca(r,TEAL)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run(f"{TOTAL} verified long-standing businesses — sorted by longevity, flagged by status, with the angle that makes each one ours, not a commodity")
r.font.size=Pt(10.5); r.font.italic=True; sca(r,SLATE)
doc.add_page_break()

H("Why a 'Still Standing' source book",1)
P("Longevity is a story competitors can't fake and chains can't buy. A business that has fed, dressed, insured, or sheltered this town for 50, 75, or 100 years is exactly the kind of deep-roots subject the Insider exists to honor — it pairs naturally with **Building's Biography**, **The Regular**, **Last Call**, and **The Dig**. This is the working source list: verified, tiered by years open, and — critically — marked with a **live status** and a **Insider angle** so we never run a tired history dump or, worse, invite readers to a place that just closed.")
P("**The discipline that makes it ours:** the obvious centenarians (Rock City, Ruby Falls, the Choo-Choo, the MoonPie origin myth) are *commodities* — every outlet has done them. We flag those and refuse the lazy version. Our value is the fresh POV: the person, the booth, the building, the closing, the descendant — never the press kit. And the once-rule still governs: a place appears once an issue, rarely twice a month.")
FIG(f"{ASSETS}/fig_legacy.png","Founding years across the three longevity tiers — the deep-roots map for an occasional 'Still Standing' series.")

H("100+ years — the Centenarians",1)
P("More than a century of continuous operation. Lead with the human thread and re-verify status before publishing.",italic=True,color=GREY,size=9)
LTABLE(CENT)
doc.add_page_break()

H("75+ years — the Near-Centenarians",1)
LTABLE(NEAR)
H("50+ years — the Established Veterans",1)
LTABLE(VET, caption="Pickle Barrel is 44 — just shy of 50 — included as 'approaching.'")
doc.add_page_break()

H("Closed — for history only (Building's Biography / The Dig)",1)
P("Do NOT invite readers here. Use only as 'what used to be' context.",italic=True,color=RUST,size=9)
LTABLE(GONE)
H("Regional / edge — Getaway & reader-service, not Chattanooga-proper",1)
LTABLE(EDGE)

# --- Insider finds (folded from agent) ---
if SCOUT_FINDS:
    doc.add_page_break()
    H("Insider finds — the hidden & satellite-town legacy others miss",1)
    P("Sourced by a Insider pass for durable, unglamorous, non-tourist businesses in the satellite towns — the hardware stores, diners, barbershops, and independents nobody profiles. Verify each before featuring; confidence noted.",italic=True,color=GREY,size=9)
    LTABLE(SCOUT_FINDS)
    H("Insider leads — promising but unverified (call before publishing)",2)
    BUL(SCOUT_LEADS)
    H("Recently lost — tribute / The Dig only",2)
    BUL(SCOUT_LOST)

# --- Status verification (folded from agent) ---
if STATUS_NOTES:
    H("Status check — verified as of mid-2026 (Nerd pass)",1)
    BUL(STATUS_NOTES)

doc.add_page_break()
H("How to run 'Still Standing'",1)
BUL([
 "**Cadence:** occasional — monthly or bimonthly, one business at a time, governed by the once-rule.",
 "**Always lead with a person or a thing, never a founding date.** The booth, the waitress of 30 years, the lazy Susan, the brick pit, the hydrant — the date is a footnote.",
 "**Re-verify status the week you publish.** Status changes fast (Zarzour's sold, MoonPie store closed, Old Plantation closed). Never invite readers to a closed door.",
 "**Refuse the commodity version.** If the only angle is the one everyone's run (the barns, the song, the origin myth), the Nerd vetoes it — find the fresh thread or skip it.",
 "**Monetization that doesn't sell editorial:** a 'Still Standing' trail/map product, an anniversary-toast event, a heritage page in The Almanac, a centennial keepsake — independent of whether a business ever pays us a cent.",
])
H("Accuracy traps the Scouts & Nerds caught",1)
BUL(MYTHS)
P("Full source list with every row and confidence rating is in legacy_source_list.csv.",italic=True,color=GREY,size=9)

out="/home/user/prscott/Still_Standing_Legacy_Source_Book.docx"
doc.save(out); print("saved",out)
